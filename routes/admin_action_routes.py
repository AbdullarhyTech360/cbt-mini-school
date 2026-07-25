from flask import render_template, session, request, jsonify, current_app, Response, stream_with_context
from models import db, User, Section, School, Permission, TraitDefinition, StudentTrait
from models.school_term import SchoolTerm
from models.subject import Subject
from models.exam import Exam
from models.question import Question
from models.assessment_type import AssessmentType
from routes.dashboard import admin_required
from models.class_room import ClassRoom
from models.student import Student
from models.teacher import Teacher
from datetime import datetime, timedelta, date
import calendar
from werkzeug.utils import secure_filename
import os
import random
from models.associations import teacher_classroom
from models.grade import Grade

from typing import List
import uuid


def initialize_default_permissions():
    """Initialize default permissions for the system. Must be called within an app context."""
    default_permissions = [
        {
            "permission_name": "users_can_register",
            "permission_description": "Allow users to register for accounts",
            "is_active": True,
            "created_for": "system"
        },
        {
            "permission_name": "teachers_create_exams",
            "permission_description": "Allow teachers to create exams",
            "is_active": True,
            "created_for": "system"
        },
        {
            "permission_name": "students_view_results",
            "permission_description": "Allow students to view their results immediately after exams",
            "is_active": False,
            "created_for": "system"
        },
        {
            "permission_name": "students_can_write_exam",
            "permission_description": "Allow students to write exams",
            "is_active": False,
            "created_for": "system"
        },
        {
            "permission_name": "admins_can_upload_questions",
            "permission_description": "Allow admins to upload questions",
            "is_active": True,
            "created_for": "system"
        },
        {
            "permission_name": "teachers_can_upload_questions",
            "permission_description": "Allow teachers to upload questions",
            "is_active": True,
            "created_for": "system"
        },
        {
            "permission_name": "students_can_view_dashboard",
            "permission_description": "Allow students to view dashboard",
            "is_active": True,
            "created_for": "system"
        },
        {
            "permission_name": "teachers_can_view_dashboard",
            "permission_description": "Allow teachers to view dashboard",
            "is_active": True,
            "created_for": "system"
        },
        {
            "permission_name": "staff_can_view_dashboard",
            "permission_description": "Allow staff to view dashboard",
            "is_active": True,
            "created_for": "system"
        },
        {
            "permission_name": "demo_question_bank",
            "permission_description": "Allow students to take demo tests",
            "is_active": False,
            "created_for": "system"
        }
    ]

    for perm_data in default_permissions:
        existing = Permission.query.filter_by(
            permission_name=perm_data["permission_name"]
        ).first()

        if not existing:
            perm = Permission(
                permission_name=perm_data["permission_name"],
                permission_description=perm_data["permission_description"],
                is_active=perm_data["is_active"],
                created_for=perm_data["created_for"]
            )
            db.session.add(perm)

    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()


def admin_action_route(app):
    # Helper function to check allowed file extensions
    def allowed_file(filename):
        return (
            "." in filename
            and filename.rsplit(".", 1)[1].lower()
            in current_app.config["ALLOWED_EXTENSIONS"]
        )

    @app.route("/admin/user_management", methods=["GET", "POST"])
    @admin_required
    def user_management():
        users = db.session.query(User).all()
        class_rooms = db.session.query(ClassRoom).all()
        current_user = User.query.get(session["user_id"])
        return render_template(
            "admin/user_management.html",
            users=users,
            class_rooms=class_rooms,
            current_user=current_user
        )
    
    @app.route("/admin/delete/user/<user_id>", methods=["DELETE"])
    @admin_required
    def delete_user(user_id):
        try:
            user = User.query.get(user_id)
            if not user:
                return jsonify({"success": False, "message": "User not found"}), 404

            db.session.delete(user)
            db.session.commit()

            return (
                jsonify({"success": True, "message": "User deleted successfully"}),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error deleting user: {str(e)}")
            return (
                jsonify(
                    {"success": False,
                        "message": f"Error deleting user: {str(e)}"}
                ),
                500,
            )

    @app.route("/admin/add/user", methods=["POST"])
    @admin_required
    def add_user():
        try:
            data = request.get_json()
            if not data:
                return jsonify({"success": False, "message": "No data provided"}), 400

            # Extract user data
            first_name = data.get("first_name", "").strip()
            last_name = data.get("last_name", "").strip()
            email = data.get("email", "").strip()
            gender = data.get("gender", "Not specified")
            dob_str = data.get("dob")
            register_number = data.get("register_number")
            class_room_id = data.get("class_room_id")
            role = data.get("role", "").strip()
            password = data.get("password", "Pa$$w0rd!")

            # Validation
            if not all([first_name, last_name, dob_str, role]):
                return jsonify({"success": False, "message": "All required fields must be filled"}), 400

            # Parse DOB
            try:
                dob = datetime.strptime(dob_str, "%Y-%m-%d").date()
            except Exception:
                return jsonify({"success": False, "message": "Invalid date format. Use YYYY-MM-DD"}), 400

            # Role-specific validation
            if role == "student":
                if not class_room_id or not register_number:
                    return jsonify({"success": False, "message": "Class and register number are required for students"}), 400

                # Check duplicate register_number within class
                existing = User.query.filter_by(
                    register_number=str(register_number),
                    class_room_id=class_room_id,
                    role="student"
                ).first()
                if existing:
                    return jsonify({"success": False, "message": "Register number already used in this class"}), 409
            elif role in ["staff", "admin"]:
                if not email:
                    return jsonify({"success": False, "message": "Email is required for staff/admin"}), 400

                import re
                if not re.match(r'^[^\s@]+@[^\s@]+\.[^\s@]+$', email):
                    return jsonify({"success": False, "message": "Please enter a valid email address"}), 400

                if not class_room_id:
                    return jsonify({"success": False, "message": "Class is required for staff/admin"}), 400

                # Check duplicate email
                existing = User.query.filter_by(email=email.lower()).first()
                if existing:
                    return jsonify({"success": False, "message": "Email already exists"}), 409

            # Generate unique username
            from models.user import generate_uuid
            username = User.generate_username(role=role)
            base_username = username
            counter = 1
            while User.query.filter_by(username=username).first():
                username = f"{base_username}{counter}"
                counter += 1

            # Create new user
            user = User(
                id=generate_uuid(),
                username=username,
                first_name=first_name,
                last_name=last_name,
                email=email.lower() if email else None,
                gender=gender,
                dob=dob,
                register_number=str(
                    register_number) if register_number else None,
                class_room_id=class_room_id,
                role=role,
            )
            user.set_password(password)
            db.session.add(user)
            db.session.flush()

            # Create role-specific profile
            if role == "student":
                student = Student(id=user.id, user_id=user.id)
                student.admission_number = username
                student.admission_date = dob
                student.parent_name = "Not specified"
                student.parent_phone = "Not specified"
                student.parent_email = "Not specified"
                student.blood_group = "Not specified"
                student.address = "Not specified"
                db.session.add(student)

                # Auto-enroll in class subjects
                from models.associations import class_subject, student_subject
                class_subjects = db.session.execute(
                    db.select(class_subject.c.subject_id).where(
                        class_subject.c.class_room_id == user.class_room_id
                    )
                ).scalars().all()

                for subject_id in class_subjects:
                    enrollment = student_subject.insert().values(
                        student_id=user.id,
                        subject_id=subject_id
                    )
                    db.session.execute(enrollment)
            elif role in ["staff", "admin"]:
                from models.teacher import Teacher
                teacher = Teacher(id=user.id, user_id=user.id)
                db.session.add(teacher)

                # Assign teacher to the selected class
                if class_room_id:
                    from models.associations import teacher_classroom
                    assignment = teacher_classroom.insert().values(
                        teacher_id=user.id,
                        class_room_id=class_room_id
                    )
                    db.session.execute(assignment)

            db.session.commit()

            return jsonify({
                "success": True,
                "message": "User created successfully",
                "username": username
            }), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error creating user: {str(e)}")
            return jsonify({"success": False, "message": f"Error creating user: {str(e)}"}), 500

    @app.route("/admin/update/user/<user_id>", methods=["PUT"])
    @admin_required
    def update_user(user_id):
        try:
            user = User.query.get(user_id)
            if not user:
                return jsonify({"success": False, "message": "User not found"}), 404

            data = request.get_json()
            if not data:
                return jsonify({"success": False, "message": "No data provided"}), 400

            # Update basic fields
            if "first_name" in data:
                user.first_name = data["first_name"].strip()
            if "last_name" in data:
                user.last_name = data["last_name"].strip()
            if "email" in data and data["email"]:
                # Check if email is already used by another user
                existing = User.query.filter(
                    User.email == data["email"].lower(),
                    User.id != user_id
                ).first()
                if existing:
                    return jsonify({"success": False, "message": "Email already exists"}), 409
                user.email = data["email"].lower()
            if "gender" in data and data["gender"]:
                gender_value = data["gender"].strip()
                normalized_gender = gender_value.title()
                if normalized_gender.lower() in ["male", "female", "other"]:
                    user.gender = normalized_gender
                else:
                    user.gender = gender_value
            if "dob" in data:
                try:
                    user.dob = datetime.strptime(
                        data["dob"], "%Y-%m-%d").date()
                except Exception:
                    return jsonify({"success": False, "message": "Invalid date format"}), 400
            if "class_room_id" in data:
                user.class_room_id = data["class_room_id"]
            if "register_number" in data and user.role == "student":
                # Check if register number is already used in the same class
                existing = User.query.filter(
                    User.register_number == str(data["register_number"]),
                    User.class_room_id == user.class_room_id,
                    User.role == "student",
                    User.id != user_id
                ).first()
                if existing:
                    return jsonify({"success": False, "message": "Register number already used in this class"}), 409
                user.register_number = str(data["register_number"])

            db.session.commit()

            return jsonify({"success": True, "message": "User updated successfully"}), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error updating user: {str(e)}")
            return jsonify({"success": False, "message": f"Error updating user: {str(e)}"}), 500

    @app.route("/admin/publish_all_scores", methods=["POST"])
    @admin_required
    def publish_all_scores():
        """Publish all scores in the system so they appear in reports"""
        try:
            # Get all unpublished grades
            unpublished_grades = Grade.query.filter_by(is_published=False).all()
            
            # Count how many we are updating
            count = 0
            
            # Update each grade to be published
            for grade in unpublished_grades:
                grade.is_published = True
                count += 1
            
            # Commit all changes
            db.session.commit()
            
            return jsonify({
                "success": True,
                "message": f"Successfully published {count} scores",
                "count": count
            })
            
        except Exception as e:
            db.session.rollback()
            # print(f"Error publishing scores: {str(e)}")
            return jsonify({
                "success": False,
                "message": f"Error publishing scores: {str(e)}"
            }), 500

    @app.route("/admin/toggle/user/<user_id>", methods=["PUT"])
    @admin_required
    def toggle_user_status(user_id):
        try:
            user = User.query.get(user_id)
            if not user:
                return jsonify({"success": False, "message": "User not found"}), 404

            # Toggle the is_active status
            user.is_active = not user.is_active
            db.session.commit()

            status = "activated" if user.is_active else "deactivated"
            return jsonify({
                "success": True,
                "message": f"User {status} successfully",
                "is_active": user.is_active
            }), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error toggling user status: {str(e)}")
            return jsonify({"success": False, "message": f"Error toggling user status: {str(e)}"}), 500

    @app.route("/admin/upload_questions", methods=["GET", "POST"])
    @admin_required
    def upload_questions():
        from models import is_permission_active
        if not is_permission_active("admins_can_upload_questions"):
            if request.method == "POST":
                return jsonify({"success": False, "message": "Question upload is disabled by the administrator"}), 403
            return redirect(url_for("admin_dashboard", denied="Question upload has been disabled by the administrator."))

        if request.method == "POST":
            try:
                data = request.get_json()
                if not data:
                    return (
                        jsonify(
                            {"success": False, "message": "No data provided"}),
                        400,
                    )

                # Get current admin user from session
                current_user = User.query.get(session["user_id"])
                if not current_user:
                    return (
                        jsonify(
                            {"success": False, "message": "Unauthorized access"}),
                        403,
                    )

                # Extract question data
                question_text = data.get("question_text", "").strip()
                question_type = data.get("question_type", "").strip()
                subject_id = data.get("subject_id")
                class_room_id = data.get("class_room_id")
                term_id = data.get("term_id")
                exam_type_id = data.get("exam_type_id")
                options_data = data.get("options", [])

                # Validation
                if not question_text:
                    return (
                        jsonify(
                            {"success": False, "message": "Question text is required"}
                        ),
                        400,
                    )

                if not question_type:
                    return (
                        jsonify(
                            {"success": False, "message": "Question type is required"}
                        ),
                        400,
                    )

                if not subject_id:
                    return (
                        jsonify(
                            {"success": False, "message": "Subject is required"}),
                        400,
                    )

                if not class_room_id:
                    return (
                        jsonify(
                            {"success": False, "message": "Class is required"}),
                        400,
                    )

                if not term_id:
                    return (
                        jsonify(
                            {"success": False, "message": "Term is required"}),
                        400,
                    )

                if not exam_type_id:
                    return (
                        jsonify(
                            {"success": False, "message": "Exam type is required"}),
                        400,
                    )

                # Validate subject exists
                subject = Subject.query.get(subject_id)
                if not subject:
                    return (
                        jsonify(
                            {"success": False, "message": "Invalid subject"}),
                        400,
                    )

                # Check if subject has a teacher assigned (mandatory requirement)
                if not subject.subject_head_id:
                    return (
                        jsonify(
                            {
                                "success": False,
                                "message": "No teacher assigned to this subject. Please assign a teacher to the subject first before uploading questions.",
                            }
                        ),
                        400,
                    )

                # Validate that the assigned teacher exists and is a staff member
                teacher = User.query.get(subject.subject_head_id)
                if not teacher or teacher.role != "staff":
                    return (
                        jsonify(
                            {
                                "success": False,
                                "message": "Assigned teacher is invalid or not a staff member. Please reassign a valid teacher to the subject.",
                            }
                        ),
                        400,
                    )

                # Validate class exists
                class_room = ClassRoom.query.get(class_room_id)
                if not class_room:
                    return jsonify({"success": False, "message": "Invalid class"}), 400

                # Validate term exists
                term = SchoolTerm.query.get(term_id)
                if not term:
                    return jsonify({"success": False, "message": "Invalid term"}), 400

                # Validate that the subject is offered in the selected class
                if subject not in class_room.subjects:
                    return (
                        jsonify(
                            {
                                "success": False,
                                "message": f"Subject '{subject.subject_name}' is not offered in class '{class_room.class_room_name}'. Please select a valid subject-class combination.",
                            }
                        ),
                        400,
                    )

                # Validate question type
                valid_types = ["mcq", "true_false", "short_answer"]
                if question_type not in valid_types:
                    return (
                        jsonify(
                            {"success": False, "message": "Invalid question type"}),
                        400,
                    )

                # For MCQ and True/False questions, validate options
                if question_type in ["mcq", "true_false"]:
                    if not options_data:
                        return (
                            jsonify(
                                {
                                    "success": False,
                                    "message": "Options are required for this question type",
                                }
                            ),
                            400,
                        )

                    # Check that there's at least one correct answer
                    correct_options = [
                        opt for opt in options_data if opt.get("is_correct", False)
                    ]
                    if not correct_options:
                        return (
                            jsonify(
                                {
                                    "success": False,
                                    "message": "At least one correct option is required",
                                }
                            ),
                            400,
                        )

                    # For true_false, ensure exactly 2 options (True and False)
                    if question_type == "true_false":
                        if len(options_data) != 2:
                            return (
                                jsonify(
                                    {
                                        "success": False,
                                        "message": "True/False questions must have exactly 2 options",
                                    }
                                ),
                                400,
                            )

                # Create the question using the assigned teacher
                from models.question import Question, Option

                new_question = Question()
                new_question.question_text = question_text
                new_question.question_type = question_type
                new_question.subject_id = subject_id
                new_question.teacher_id = subject.subject_head_id  # Use the assigned teacher
                new_question.class_room_id = class_room_id
                new_question.term_id = term_id
                new_question.exam_type_id = exam_type_id

                # For short answer questions, save the correct answer
                if question_type == "short_answer":
                    new_question.correct_answer = data.get(
                        "correct_answer", "")

                db.session.add(new_question)
                db.session.flush()  # Get the question ID without committing

                # Create options if this is an MCQ or True/False question
                if question_type in ["mcq", "true_false"]:
                    for i, option_data in enumerate(options_data):
                        option_text = option_data.get("text", "").strip()
                        is_correct = option_data.get("is_correct", False)

                        if option_text:  # Only create option if text is provided
                            option = Option()
                            option.text = option_text
                            option.is_correct = is_correct
                            option.order = i
                            option.question_id = new_question.id
                            db.session.add(option)

                db.session.commit()

                return (
                    jsonify(
                        {
                            "success": True,
                            "message": "Question created successfully",
                            "question_id": new_question.id,
                        }
                    ),
                    200,
                )

            except Exception as e:
                db.session.rollback()
                # print(f"Error creating question: {str(e)}")
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": f"Error creating question: {str(e)}",
                        }
                    ),
                    500,
                )

        # GET request - render the upload questions page
        current_user = User.query.get(session["user_id"])
        classes = db.session.query(ClassRoom).all()
        subjects = db.session.query(Subject).all()
        school_terms = db.session.query(SchoolTerm).all()
        teachers = db.session.query(User).filter_by(role="staff").all()
        assessment_types = db.session.query(AssessmentType).order_by(AssessmentType.order).all()
        active_assessment_type = [at for at in assessment_types if at.is_active][0]

        print(active_assessment_type)

        # Find the current term (if any)
        current_term = next(
            (term for term in school_terms if term.is_current), None)
        current_term_id = current_term.term_id if current_term else None

        # Get assigned subjects for teachers (for extraction dropdown)
        from models.associations import teacher_subject
        assigned_subjects = []
        for teacher in teachers:
            assignments = db.session.query(teacher_subject, Subject, ClassRoom).join(
                Subject, teacher_subject.c.subject_id == Subject.subject_id
            ).join(
                ClassRoom, teacher_subject.c.class_room_id == ClassRoom.class_room_id
            ).filter(teacher_subject.c.teacher_id == teacher.id).all()
            
            for assignment in assignments:
                assigned_subjects.append({
                    'teacher_id': teacher.id,
                    'teacher_name': f"{teacher.first_name} {teacher.last_name}",
                    'subject_id': assignment.Subject.subject_id,
                    'subject_name': assignment.Subject.subject_name,
                    'class_room_id': assignment.ClassRoom.class_room_id,
                    'class_room_name': assignment.ClassRoom.class_room_name,
                    'display_name': f"{teacher.first_name} {teacher.last_name} - {assignment.Subject.subject_name} ({assignment.ClassRoom.class_room_name})"
                })
            
        return render_template(
            "admin/upload_questions.html",
            current_user=current_user,
            classes=classes,
            subjects=subjects,
            school_terms=school_terms,
            current_term_id=current_term_id,
            teachers=teachers,
            assessment_types=assessment_types,
            assigned_subjects=assigned_subjects,
            # current
        )

    @app.route("/admin/extract_questions", methods=["POST"])
    @admin_required
    def admin_extract_questions():
        """Stream extraction milestones and final parsed questions for review."""
        from models import is_permission_active
        if not is_permission_active("admins_can_upload_questions"):
            return jsonify({"success": False, "message": "Question upload is disabled by the administrator"}), 403

        try:
            import json
            import os
            import queue
            import tempfile
            import threading
            import time

            # Verify user
            current_user = User.query.get(session["user_id"]) if session.get("user_id") else None
            if not current_user:
                return jsonify({"success": False, "message": "Unauthorized access"}), 403

            # Validate multipart - accept multiple files (up to 4)
            files = request.files.getlist('file')
            if not files or all(not f or f.filename == '' for f in files):
                return jsonify({"success": False, "message": "No file uploaded"}), 400

            # Filter out empty files and limit to 4 files
            files = [f for f in files if f and f.filename != '']
            if len(files) > 4:
                return jsonify({"success": False, "message": "Maximum 4 files allowed at once"}), 400

            if len(files) == 0:
                return jsonify({"success": False, "message": "No file selected"}), 400

            # Get metadata
            teacher_id = request.form.get('teacher_id')
            subject_id = request.form.get('subject_id')
            class_room_id = request.form.get('class_room_id')
            term_id = request.form.get('term_id')
            exam_type_id = request.form.get('exam_type_id')
            custom_prompt = request.form.get('custom_prompt')

            if not teacher_id or not subject_id or not class_room_id or not term_id or not exam_type_id:
                return jsonify({"success": False, "message": "Teacher, Subject, Class, Term, and Exam Type are required"}), 400

            # Verify teacher assignment
            from models.associations import teacher_subject
            assignment = (
                db.session.query(teacher_subject)
                .filter_by(teacher_id=teacher_id, subject_id=subject_id, class_room_id=class_room_id)
                .first()
            )
            if not assignment:
                return jsonify({"success": False, "message": "The selected teacher is not assigned to this subject-class combination"}), 403

            # Save all files to temporary files
            temp_files = []
            filenames = []
            for file in files:
                filename = file.filename or ''
                filenames.append(filename)
                suffix = ''
                if '.' in filename:
                    suffix = '.' + filename.rsplit('.', 1)[1].lower()
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                file.save(tmp.name)
                tmp.close()
                temp_files.append(tmp.name)

            from services.gemini_extractor import extract_questions_from_file

            event_queue = queue.Queue()
            result_holder = {"questions": [], "error": None, "model_name": None}
            worker_done = threading.Event()

            def stream_event(payload):
                event_queue.put(payload)

            def extractor_worker():
                try:
                    all_questions = []
                    for idx, (tmp_file, filename) in enumerate(zip(temp_files, filenames)):
                        # Add file index to progress events
                        def wrapped_stream_event(payload):
                            payload_with_file = payload.copy()
                            payload_with_file["file_index"] = idx + 1
                            payload_with_file["total_files"] = len(temp_files)
                            payload_with_file["file_name"] = filename
                            event_queue.put(payload_with_file)

                        questions, error, model_name = extract_questions_from_file(
                            tmp_file,
                            custom_prompt=custom_prompt,
                            progress_callback=wrapped_stream_event,
                        )
                        if error:
                            result_holder["error"] = f"Error processing {filename}: {error}"
                            break
                        if questions:
                            all_questions.extend(questions)
                        result_holder["model_name"] = model_name

                    result_holder["questions"] = all_questions
                except Exception as worker_error:
                    result_holder["error"] = str(worker_error)
                finally:
                    worker_done.set()

            threading.Thread(target=extractor_worker, daemon=True).start()

            def build_event(event, **payload):
                response_payload = {"event": event, **payload}
                return json.dumps(response_payload) + "\n"

            @stream_with_context
            def generate():
                last_emit_at = time.time()
                last_progress = 2
                file_list = ", ".join(filenames) if len(filenames) <= 3 else f"{filenames[0]}, {filenames[1]} and {len(filenames) - 2} more"
                yield build_event(
                    "milestone",
                    progress=2,
                    status="running",
                    stage_id="initializing-model-pipeline",
                    stage_label="Initializing model pipeline",
                    message=f"Upload received. Starting extraction workflow for {len(files)} file(s): {file_list}",
                    file_count=len(files),
                    file_names=filenames,
                )

                try:
                    while not worker_done.is_set() or not event_queue.empty():
                        try:
                            payload = event_queue.get(timeout=1.0)
                            last_emit_at = time.time()
                            if payload.get("progress") is not None:
                                last_progress = payload.get("progress")
                            yield build_event(payload.pop("event", "milestone"), **payload)
                        except queue.Empty:
                            if worker_done.is_set():
                                break
                            if time.time() - last_emit_at >= 2:
                                last_emit_at = time.time()
                                yield build_event(
                                    "heartbeat",
                                    progress=last_progress,
                                    status="waiting",
                                    stage_id="processing-input-text",
                                    stage_label="Processing input text",
                                    message="Extraction is still active. The current model is working on your document.",
                                )

                    while not event_queue.empty():
                        payload = event_queue.get_nowait()
                        yield build_event(payload.pop("event", "milestone"), **payload)

                    if result_holder["error"]:
                        yield build_event(
                            "error",
                            success=False,
                            message=result_holder["error"],
                        )
                        return

                    questions = result_holder["questions"] or []
                    if not questions:
                        yield build_event(
                            "error",
                            success=False,
                            message="No questions extracted",
                        )
                        return

                    yield build_event(
                        "complete",
                        success=True,
                        progress=100,
                        status="completed",
                        stage_id="generating-structured-question-output",
                        stage_label="Generating structured question output",
                        message=f"Extraction completed successfully with {len(questions)} question(s) from {len(files)} file(s).",
                        model_name=result_holder["model_name"],
                        questions=questions,
                        file_count=len(files),
                    )
                finally:
                    # Clean up all temporary files
                    for tmp_file in temp_files:
                        try:
                            os.unlink(tmp_file)
                        except OSError:
                            pass

            response = Response(generate(), mimetype="application/x-ndjson")
            response.headers["Cache-Control"] = "no-cache"
            response.headers["X-Accel-Buffering"] = "no"
            return response

        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/ai_chat_generate", methods=["POST"])
    @admin_required
    def admin_ai_chat_generate():
        """AI Chat endpoint for generating questions from uploaded documents or text prompts with system prompt."""
        from models import is_permission_active
        if not is_permission_active("admins_can_upload_questions"):
            return jsonify({"success": False, "message": "Question upload is disabled by the administrator"}), 403

        try:
            import os
            import tempfile
            import json

            # Verify user
            current_user = User.query.get(session["user_id"]) if session.get("user_id") else None
            if not current_user:
                return jsonify({"success": False, "error": "Unauthorized access"}), 403

            # Get files, instruction, and metadata
            files = request.files.getlist('files')
            instruction = request.form.get('instruction', '')
            teacher_id = request.form.get('teacher_id')
            conversation_history_str = request.form.get('conversation_history', '[]')

            # Filter out empty files
            files = [f for f in files if f and f.filename != '']
            
            if not files and not instruction:
                return jsonify({"success": False, "error": "Please provide files or instructions"}), 400

            # System prompt for AI
            system_prompt = """You are an academic assistant helping teachers/staff create assessment questions. 
Generate questions based on the uploaded documents and user instructions.
Format questions according to the system's question structure (question_text, question_type, options, correct_answer, has_math, context).
Support MCQ, True/False, and Short Answer question types.
For MCQ: provide at least 3-4 options with exactly one correct answer.
For True/False: represent as MCQ with exactly 2 options (True and False).
Include context field for tables, diagrams, formulas referenced in questions."""

            # Decode conversation history
            try:
                conversation_history = json.loads(conversation_history_str)
            except Exception:
                conversation_history = []

            # Format conversation history for context
            history_prompt = ""
            if conversation_history:
                history_prompt = "Here is the conversation history so far for reference:\n"
                for msg in conversation_history:
                    if msg.get("content") and "Start a conversation" not in msg.get("content"):
                        role = "User" if msg.get("type") == "user" else "Assistant"
                        history_prompt += f"{role}: {msg.get('content')}\n"
                history_prompt += "\n"

            # Combine system prompt with user instruction and history
            full_prompt = f"{system_prompt}\n\n{history_prompt}User instruction: {instruction}"

            from services.gemini_extractor import extract_questions_from_file, generate_questions_from_prompt

            all_questions = []
            model_name = None
            error = None

            if files:
                # Save files to temporary files
                temp_files = []
                for file in files:
                    filename = file.filename or ''
                    suffix = ''
                    if '.' in filename:
                        suffix = '.' + filename.rsplit('.', 1)[1].lower()
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                    file.save(tmp.name)
                    tmp.close()
                    temp_files.append(tmp.name)

                for tmp_file in temp_files:
                    questions, err, model = extract_questions_from_file(
                        tmp_file,
                        custom_prompt=full_prompt,
                        progress_callback=None  # No streaming for chat interface
                    )
                    if err:
                        error = err
                        break
                    if questions:
                        all_questions.extend(questions)
                    model_name = model

                # Clean up temporary files
                for tmp_file in temp_files:
                    try:
                        os.unlink(tmp_file)
                    except OSError:
                        pass
            else:
                # Text-only generation if no files were uploaded
                questions, err, model = generate_questions_from_prompt(
                    full_prompt,
                    progress_callback=None
                )
                if err:
                    error = err
                if questions:
                    all_questions.extend(questions)
                model_name = model

            if error:
                return jsonify({"success": False, "error": error}), 500

            if not all_questions:
                return jsonify({"success": False, "error": "No questions generated"}), 400

            return jsonify({
                "success": True,
                "message": f"Successfully generated {len(all_questions)} question(s)",
                "questions": all_questions,
                "model_used": model_name
            })

        except Exception as e:
            return jsonify({"success": False, "error": str(e)}), 500

    @app.route("/admin/create_questions_from_json", methods=["POST"])
    @admin_required
    def admin_create_questions_from_json():
        """Create questions from a JSON payload (used after AI extraction review in admin).
        Expects JSON: { teacher_id, subject_id, class_room_id, term_id, exam_type_id, questions: [...] }
        """
        from models import is_permission_active
        if not is_permission_active("admins_can_upload_questions"):
            return jsonify({"success": False, "message": "Question upload is disabled by the administrator"}), 403

        try:
            data = request.get_json()
            if not data:
                return jsonify({"success": False, "message": "No data provided"}), 400

            current_user = User.query.get(session["user_id"]) if session.get("user_id") else None
            if not current_user:
                return jsonify({"success": False, "message": "Unauthorized access"}), 403

            teacher_id = data.get('teacher_id')
            subject_id = data.get('subject_id')
            class_room_id = data.get('class_room_id')
            term_id = data.get('term_id')
            exam_type_id = data.get('exam_type_id')
            questions = data.get('questions', [])

            if not teacher_id or not subject_id or not class_room_id or not term_id or not exam_type_id:
                return jsonify({"success": False, "message": "Teacher, Subject, Class, Term, and Exam Type are required"}), 400

            # Verify that the teacher is assigned to this classroom/subject
            from models.associations import teacher_subject
            assignment = (
                db.session.query(teacher_subject)
                .filter_by(teacher_id=teacher_id, subject_id=subject_id, class_room_id=class_room_id)
                .first()
            )
            if not assignment:
                return jsonify({"success": False, "message": "The selected teacher is not assigned to this subject-class combination"}), 403

            from models.question import Question, Option

            created_questions = []
            errors = []

            for i, question_data in enumerate(questions, 1):
                try:
                    question_text = question_data.get('question_text', '').strip()
                    question_type = question_data.get('question_type', '').strip().lower()
                    options_data = question_data.get('options', [])
                    correct_answer = question_data.get('correct_answer', '')

                    if not question_text:
                        errors.append(f"Question {i}: Question text is required")
                        continue

                    if not question_type:
                        if options_data:
                            question_type = 'mcq'
                        elif correct_answer:
                            question_type = 'short_answer'
                        else:
                            errors.append(f"Question {i}: Question type is required")
                            continue

                    if question_type == 'true_false':
                        question_type = 'mcq'

                    valid_types = ["mcq", "short_answer"]
                    if question_type not in valid_types:
                        errors.append(f"Question {i}: Invalid question type '{question_type}'")
                        continue

                    if question_type == "mcq":
                        if not options_data or len(options_data) == 0:
                            errors.append(f"Question {i}: Options are required for MCQ questions")
                            continue
                        correct_options = [opt for opt in options_data if opt.get('is_correct', False)]
                        if not correct_options:
                            errors.append(f"Question {i}: At least one correct option is required")
                            continue

                    new_question = Question()
                    new_question.question_text = question_text
                    new_question.question_type = question_type
                    new_question.subject_id = subject_id
                    new_question.teacher_id = teacher_id
                    new_question.class_room_id = class_room_id
                    new_question.term_id = term_id
                    new_question.exam_type_id = exam_type_id
                    new_question.has_math = question_data.get('has_math', False)
                    new_question.question_image = question_data.get('question_image')

                    if question_type == 'short_answer':
                        new_question.correct_answer = correct_answer

                    db.session.add(new_question)
                    db.session.flush()

                    if question_type == "mcq":
                        for j, option_data in enumerate(options_data):
                            option_text = option_data.get('text', '').strip()
                            is_correct = option_data.get('is_correct', False)
                            if option_text:
                                opt = Option()
                                opt.text = option_text
                                opt.is_correct = is_correct
                                opt.order = j
                                opt.question_id = new_question.id
                                opt.has_math = option_data.get('has_math', False)
                                opt.option_image = option_data.get('option_image')
                                db.session.add(opt)

                    db.session.commit()
                    created_questions.append(new_question.id)

                except Exception as e:
                    db.session.rollback()
                    errors.append(f"Question {i}: Error creating question - {str(e)}")
                    continue

            response_data = {
                'success': True,
                'message': f"Successfully created {len(created_questions)} questions",
                'created_count': len(created_questions),
                'error_count': len(errors)
            }
            if errors:
                response_data['message'] += f" with {len(errors)} errors"
                response_data['errors'] = errors

            return jsonify(response_data), 200

        except Exception as e:
            db.session.rollback()
            import traceback
            traceback.print_exc()
            return jsonify({'success': False, 'message': str(e)}), 500


    @app.route("/admin/questions_preview")
    @admin_required
    def admin_questions_preview():
        """Endpoint to get questions preview based on selected criteria"""
        try:
            subject_id = request.args.get('subject_id')
            class_room_id = request.args.get('class_room_id')
            term_id = request.args.get('term_id')
            exam_type_id = request.args.get('exam_type_id')

            # Validate required parameters
            if not all([subject_id, class_room_id, term_id, exam_type_id]):
                return jsonify({
                    "success": False,
                    "message": "Missing required parameters"
                }), 400

            # Query questions based on criteria
            questions = Question.query.filter_by(
                subject_id=subject_id,
                class_room_id=class_room_id,
                term_id=term_id,
                exam_type_id=exam_type_id
            ).order_by(Question.created_at).all()

            # Prepare questions data with randomized options
            questions_data = []
            for question in questions:
                # Get options and randomize their order
                options = list(question.options)
                random.shuffle(options)

                # Create option data with new order indices
                options_data = []
                for i, option in enumerate(options):
                    options_data.append({
                        'id': option.id,
                        'text': option.text,
                        'is_correct': option.is_correct,
                        'order': i  # New randomized order
                    })

                questions_data.append({
                    'id': question.id,
                    'question_text': question.question_text,
                    'question_type': question.question_type,
                    'options': options_data,
                    'correct_answer': question.correct_answer
                })

            return jsonify({
                "success": True,
                "questions": questions_data,
                "total_questions": len(questions_data)
            })

        except Exception as e:
            # print(f"Error fetching questions preview: {str(e)}")
            return jsonify({
                "success": False,
                "message": "Error fetching questions preview"
            }), 500

    @app.route("/admin/bulk_upload_questions", methods=["GET", "POST"])
    @admin_required
    def bulk_upload_questions():
        from models import is_permission_active
        if not is_permission_active("admins_can_upload_questions"):
            if request.method == "POST":
                return jsonify({"success": False, "message": "Question upload is disabled by the administrator"}), 403
            return redirect(url_for("admin_dashboard", denied="Question upload has been disabled by the administrator."))

        if request.method == "POST":
            try:
                # Check if this is a file upload (Word document)
                if 'file' in request.files:
                    file = request.files['file']
                    if file and hasattr(file, 'filename') and file.filename and file.filename.endswith('.docx'):
                        # Handle Word document upload
                        from utils.docx_parser import parse_docx_questions

                        file_content = file.read()
                        questions_data = parse_docx_questions(file_content)

                        if not questions_data:
                            return (
                                jsonify(
                                    {"success": False, "message": "No questions found in the document"}),
                                400,
                            )
                    else:
                        return (
                            jsonify(
                                {"success": False, "message": "Invalid file format. Only DOCX files are supported for Word uploads."}),
                            400,
                        )
                else:
                    # Handle JSON data upload (existing functionality)
                    data = request.get_json()
                    if not data:
                        return (
                            jsonify(
                                {"success": False, "message": "No data provided"}),
                            400,
                        )

                    # Extract bulk questions data
                    questions_data = data.get("questions", [])
                    if not questions_data:
                        return (
                            jsonify(
                                {"success": False, "message": "No questions provided"}),
                            400,
                        )

                # Get current admin user from session
                current_user = User.query.get(session["user_id"])
                if not current_user:
                    return (
                        jsonify(
                            {"success": False, "message": "Unauthorized access"}),
                        403,
                    )

                created_questions = []
                errors = []

                # Process each question
                for i, question_data in enumerate(questions_data):
                    try:
                        # Extract question data
                        question_text = question_data.get(
                            "question_text", "").strip()
                        question_type = question_data.get(
                            "question_type", "").strip()
                        subject_id = question_data.get("subject_id")
                        class_room_id = question_data.get("class_room_id")
                        options_data = question_data.get("options", [])
                        correct_answer = question_data.get(
                            "correct_answer", "")

                        # Validation
                        if not question_text:
                            errors.append(
                                f"Question {i+1}: Question text is required")
                            continue

                        if not question_type:
                            errors.append(
                                f"Question {i+1}: Question type is required")
                            continue

                        if not subject_id:
                            errors.append(
                                f"Question {i+1}: Subject is required")
                            continue

                        if not class_room_id:
                            errors.append(f"Question {i+1}: Class is required")
                            continue

                        # Validate subject exists
                        subject = Subject.query.get(subject_id)
                        if not subject:
                            errors.append(f"Question {i+1}: Invalid subject")
                            continue

                        # Check if subject has a teacher assigned (mandatory requirement)
                        if not subject.subject_head_id:
                            errors.append(
                                f"Question {i+1}: No teacher assigned to this subject. Please assign a teacher to the subject first before uploading questions.")
                            continue

                        # Validate that the assigned teacher exists and is a staff member
                        teacher = User.query.get(subject.subject_head_id)
                        if not teacher or teacher.role != "staff":
                            errors.append(
                                f"Question {i+1}: Assigned teacher is invalid or not a staff member. Please reassign a valid teacher to the subject.")
                            continue

                        # Validate class exists
                        class_room = ClassRoom.query.get(class_room_id)
                        if not class_room:
                            errors.append(f"Question {i+1}: Invalid class")
                            continue

                        # Validate that the subject is offered in the selected class
                        if subject not in class_room.subjects:
                            errors.append(
                                f"Question {i+1}: Subject '{subject.subject_name}' is not offered in class '{class_room.class_room_name}'. Please select a valid subject-class combination.")
                            continue

                        # Validate question type
                        valid_types = ["mcq", "true_false", "short_answer"]
                        if question_type not in valid_types:
                            errors.append(
                                f"Question {i+1}: Invalid question type")
                            continue

                        # For MCQ and True/False questions, validate options
                        if question_type in ["mcq", "true_false"]:
                            if not options_data:
                                errors.append(
                                    f"Question {i+1}: Options are required for this question type")
                                continue

                            # Check that there's at least one correct answer
                            correct_options = [
                                opt for opt in options_data if opt.get("is_correct", False)
                            ]
                            if not correct_options:
                                errors.append(
                                    f"Question {i+1}: At least one correct option is required")
                                continue

                            # For true_false, ensure exactly 2 options (True and False)
                            if question_type == "true_false":
                                if len(options_data) != 2:
                                    errors.append(
                                        f"Question {i+1}: True/False questions must have exactly 2 options")
                                    continue

                        # Create the question using the assigned teacher
                        from models.question import Question, Option

                        new_question = Question()
                        new_question.question_text = question_text
                        new_question.question_type = question_type
                        new_question.subject_id = subject_id
                        new_question.teacher_id = subject.subject_head_id  # Use the assigned teacher
                        new_question.class_room_id = class_room_id

                        # Save math and image data if present
                        new_question.has_math = question_data.get(
                            "has_math", False)
                        new_question.question_image = question_data.get(
                            "question_image")

                        # For short answer questions, save the correct answer
                        if question_type == "short_answer":
                            new_question.correct_answer = correct_answer

                        db.session.add(new_question)
                        db.session.flush()  # Get the question ID without committing

                        # Create options if this is an MCQ or True/False question
                        if question_type in ["mcq", "true_false"]:
                            for j, option_data in enumerate(options_data):
                                option_text = option_data.get(
                                    "text", "").strip()
                                is_correct = option_data.get(
                                    "is_correct", False)

                                if option_text:  # Only create option if text is provided
                                    option = Option()
                                    option.text = option_text
                                    option.is_correct = is_correct
                                    option.order = j
                                    option.question_id = new_question.id

                                    # Save math and image data for options
                                    option.has_math = option_data.get(
                                        "has_math", False)
                                    option.option_image = option_data.get(
                                        "option_image")

                                    db.session.add(option)

                        db.session.commit()
                        created_questions.append(new_question.id)

                    except Exception as e:
                        db.session.rollback()
                        errors.append(
                            f"Question {i+1}: Error creating question - {str(e)}")

                # Prepare response
                response_data = {
                    "success": True,
                    "message": f"Successfully created {len(created_questions)} questions",
                    "created_count": len(created_questions),
                    "error_count": len(errors),
                }

                if errors:
                    response_data["message"] += f" with {len(errors)} errors"
                    response_data["errors"] = errors

                return jsonify(response_data), 200

            except Exception as e:
                db.session.rollback()
                # print(f"Error processing bulk upload: {str(e)}")
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": f"Error processing bulk upload: {str(e)}",
                        }
                    ),
                    500,
                )

        # GET request - render the bulk upload questions page
        current_user = User.query.get(session["user_id"])
        classes = db.session.query(ClassRoom).all()
        subjects = db.session.query(Subject).all()
        return render_template(
            "admin/bulk_upload_questions.html",
            current_user=current_user,
            classes=classes,
            subjects=subjects,
        )

    @app.route("/admin/download_docx_template")
    @admin_required
    def download_docx_template():
        """Generate and download a DOCX template for bulk question upload"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from io import BytesIO

            # Create a new Document
            doc = Document()

            # Add title
            title = doc.add_paragraph()
            title_run = title.add_run('Question Upload Template')
            title_run.bold = True
            title_run.font.size = Pt(16)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add instructions
            doc.add_heading('Instructions:', level=2)
            instructions = [
                '1. Each question should be separated by a blank line',
                '2. Start each question with "Question:"',
                '3. Specify the question type: MCQ, True/False, or Short Answer',
                '4. For MCQ and True/False: list options with * for correct answer',
                '5. For Short Answer: provide the correct answer',
            ]
            for instruction in instructions:
                doc.add_paragraph(instruction, style='List Bullet')

            # Add format explanation
            doc.add_heading('Format:', level=2)
            format_para = doc.add_paragraph()
            format_para.add_run('Question: ').bold = True
            format_para.add_run('[Your question text]\n\n')
            format_para.add_run('Type: ').bold = True
            format_para.add_run('[MCQ/True/False/Short Answer]\n\n')
            format_para.add_run('Options:\n\n').bold = True
            format_para.add_run('- Option 1\n')
            format_para.add_run('- Option 2\n')
            format_para.add_run(
                '- *Correct Option (use * prefix for correct answer)\n\n')
            format_para.add_run('OR for Short Answer:\n\n').bold = True
            format_para.add_run('Answer: ').bold = True
            format_para.add_run('[Correct answer text]')

            # Add examples
            doc.add_heading('Examples:', level=2)

            # Example 1: MCQ
            doc.add_heading('Example 1 - Multiple Choice Question:', level=3)
            example1 = doc.add_paragraph()
            example1.add_run('Question: ').bold = True
            example1.add_run('What is the capital of France?\n\n')
            example1.add_run('Type: ').bold = True
            example1.add_run('MCQ\n\n')
            example1.add_run('Options:\n\n').bold = True
            example1.add_run('- London\n')
            example1.add_run('- *Paris\n')
            example1.add_run('- Berlin\n')
            example1.add_run('- Madrid')

            doc.add_paragraph()  # Blank line separator

            # Example 2: True/False
            doc.add_heading('Example 2 - True/False Question:', level=3)
            example2 = doc.add_paragraph()
            example2.add_run('Question: ').bold = True
            example2.add_run('The Earth is flat.\n\n')
            example2.add_run('Type: ').bold = True
            example2.add_run('True/False\n\n')
            example2.add_run('Options:\n\n').bold = True
            example2.add_run('- True\n')
            example2.add_run('- *False')

            doc.add_paragraph()  # Blank line separator

            # Example 3: Short Answer
            doc.add_heading('Example 3 - Short Answer Question:', level=3)
            example3 = doc.add_paragraph()
            example3.add_run('Question: ').bold = True
            example3.add_run('What is the chemical symbol for water?\n\n')
            example3.add_run('Type: ').bold = True
            example3.add_run('Short Answer\n\n')
            example3.add_run('Answer: ').bold = True
            example3.add_run('H2O')

            # Save to BytesIO
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            # Send file
            from flask import send_file
            return send_file(
                file_stream,
                as_attachment=True,
                download_name='questions_template.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            # print(f"Error generating DOCX template: {str(e)}")
            return jsonify({
                "success": False,
                "message": f"Error generating template: {str(e)}"
            }), 500

    @app.route("/admin/report_generation", methods=["GET", "POST"])
    @admin_required
    def report_generation():
        current_user = User.query.get(session["user_id"])
        return render_template(
            "admin/report_generation.html", current_user=current_user
        )

    # ===============================
    # EXAM MANAGEMENT
    # ===============================
    @app.route("/admin/exams", methods=["GET", "POST"])
    @admin_required
    def exam_management():
        if request.method == "POST":
            try:
                data = request.get_json()
                is_on_the_go = data.get("is_on_the_go", False)

                # Validation
                if not data.get("exam_type"):
                    return (
                        jsonify(
                            {"success": False, "message": "Exam type is required"}),
                        400,
                    )
                if not data.get("subject_id"):
                    return (
                        jsonify(
                            {"success": False, "message": "Subject is required"}),
                        400,
                    )
                if not data.get("class_room_id"):
                    return (
                        jsonify(
                            {"success": False, "message": "Class is required"}),
                        400,
                    )
                if not is_on_the_go and not data.get("school_term_id"):
                    return (
                        jsonify(
                            {"success": False, "message": "School term is required"}
                        ),
                        400,
                    )
                if not data.get("max_score"):
                    return (
                        jsonify(
                            {"success": False, "message": "Maximum score is required"}
                        ),
                        400,
                    )
                if not data.get("duration_hours") and not data.get("duration_minutes"):
                    return (
                        jsonify(
                            {"success": False, "message": "Exam duration is required"}
                        ),
                        400,
                    )
                if not is_on_the_go and not data.get("date"):
                    return (
                        jsonify(
                            {"success": False, "message": "Exam date is required"}
                        ),
                        400,
                    )

                # Parse duration
                hours = int(data.get("duration_hours", 0))
                minutes = int(data.get("duration_minutes", 0))
                duration = timedelta(hours=hours, minutes=minutes)

                # Parse exam date (On-The-Go tests use current date if not specified)
                if data.get("date"):
                    exam_date = datetime.strptime(data.get("date"), "%Y-%m-%d")
                else:
                    exam_date = datetime.utcnow()

                # Get subject and class to auto-generate name
                subject = Subject.query.get(data.get("subject_id"))
                class_room = ClassRoom.query.get(data.get("class_room_id"))

                if not subject:
                    return jsonify({"success": False, "message": "Invalid subject"}), 400
                if not class_room:
                    return jsonify({"success": False, "message": "Invalid class"}), 400

                # Check if questions exist for this subject and class combination
                from models.question import Question
                question_count = Question.query.filter_by(
                    subject_id=data.get("subject_id"),
                    class_room_id=data.get("class_room_id")
                ).count()

                if question_count == 0:
                    return jsonify({
                        "success": False,
                        "message": f"No questions available for {subject.subject_name} in {class_room.class_room_name}. Please add questions before creating an exam."
                    }), 400

                # Auto-generate exam name: {Class Name}-{Subject Name}-{Exam Type}
                exam_name = f"{class_room.class_room_name}-{subject.subject_name}-{data.get('exam_type')}"

                # Get number of questions (optional)
                number_of_questions = data.get("number_of_questions")
                if number_of_questions:
                    number_of_questions = int(number_of_questions)
                    # Validate that requested number doesn't exceed available questions
                    if number_of_questions > question_count:
                        return jsonify({
                            "success": False,
                            "message": f"Cannot create exam with {number_of_questions} questions. Only {question_count} questions available for {subject.subject_name} in {class_room.class_room_name}."
                        }), 400

                # Check if exam already exists for this subject, class, term and exam type (skip for On-The-Go)
                if not is_on_the_go:
                    existing_exam = Exam.query.filter_by(
                        subject_id=data.get("subject_id"),
                        class_room_id=data.get("class_room_id"),
                        school_term_id=data.get("school_term_id"),
                        exam_type=data.get("exam_type")
                    ).first()

                    if existing_exam:
                        return jsonify({
                            "success": False,
                            "message": f"An exam already exists for {subject.subject_name} in {class_room.class_room_name} for this term and exam type."
                        }), 409

                # Create new exam
                new_exam = Exam()
                new_exam.name = exam_name
                new_exam.exam_type = data.get("exam_type")
                new_exam.description = data.get("description", "")
                new_exam.date = exam_date
                new_exam.duration = duration
                new_exam.subject_id = data.get("subject_id")
                new_exam.class_room_id = data.get("class_room_id")
                new_exam.school_term_id = data.get("school_term_id")
                new_exam.invigilator_id = data.get("invigilator_id")
                new_exam.max_score = float(data.get("max_score"))
                new_exam.number_of_questions = number_of_questions
                new_exam.calculator_enabled = data.get("calculator_enabled", False)
                new_exam.is_on_the_go = data.get("is_on_the_go", False)
                new_exam.save_after_completion = data.get("save_after_completion", True)
                new_exam.show_feedback = data.get("show_feedback", True)

                db.session.add(new_exam)
                db.session.commit()

                return (
                    jsonify(
                        {
                            "success": True,
                            "message": "Exam created successfully",
                            "exam_id": new_exam.id,
                            "exam_name": exam_name,
                        }
                    ),
                    200,
                )

            except Exception as e:
                db.session.rollback()
                # print(f"Error creating exam: {str(e)}")
                return (
                    jsonify(
                        {"success": False,
                            "message": f"Error creating exam: {str(e)}"}
                    ),
                    500,
                )

        # GET request
        current_user = User.query.get(session["user_id"])
        exams = Exam.query.all()
        subjects = Subject.query.all()
        school_terms = SchoolTerm.query.all()
        teachers = User.query.filter_by(role="staff").all()
        class_rooms = ClassRoom.query.all()

        # Find the current term (if any)
        current_term = next(
            (term for term in school_terms if term.is_current), None)
        current_term_id = current_term.term_id if current_term else None

        # Get assessment types for the school
        from models.assessment_type import AssessmentType
        from models.school import School
        school = School.query.first()
        assessment_types = []
        if school:
            assessment_types = AssessmentType.query.filter_by(
                school_id=school.school_id,
                is_active=True,
                is_cbt_enabled=True  # Only show CBT-enabled assessments
            ).order_by(AssessmentType.order).all()

        # Get all active exams for this term
        from models.associations import student_subject, student_exam
        active_exams = Exam.query.filter_by(
            school_term_id=current_term_id
        ).all()

        # Get all students for reset exam functionality
        students = User.query.filter_by(is_active=True, role="student").all()

        return render_template(
            "admin/exams.html",
            current_user=current_user,
            exams=exams,
            subjects=subjects,
            school_terms=school_terms,
            teachers=teachers,
            class_rooms=class_rooms,
            current_term_id=current_term_id,
            assessment_types=assessment_types,
            students=students,
        )

    @app.route("/admin/exams/classes-by-subject/<subject_id>", methods=["GET"])
    @admin_required
    def get_classes_by_subject(subject_id):
        """Get all classes that offer a specific subject"""
        try:
            # Get the subject to check if it's a general subject
            subject = Subject.query.get(subject_id)
            if not subject:
                return jsonify({"success": False, "message": "Subject not found"}), 404

            # For general subjects (Mathematics and English), return all active classes
            # These are automatically linked to all classes during initialization
            is_general_subject = subject.subject_category == "general" and subject.subject_code in [
                "MATH", "ENG"]

            if is_general_subject:
                # Return all active classes for general subjects
                classes = ClassRoom.query.filter(
                    ClassRoom.is_active == True).all()
            else:
                # For other subjects, get only classes linked via class_subject association
                from models.associations import class_subject

                # Get all class IDs linked to this subject
                class_links = db.session.execute(
                    db.select(class_subject.c.class_room_id).where(
                        class_subject.c.subject_id == subject_id
                    )
                ).fetchall()

                class_ids = [link[0] for link in class_links]

                # Get class details
                classes = ClassRoom.query.filter(
                    ClassRoom.class_room_id.in_(class_ids),
                    ClassRoom.is_active == True
                ).all()

            # Convert to serializable format
            classes_data = []
            for cls in classes:
                classes_data.append({
                    "class_room_id": cls.class_room_id,
                    "class_room_name": cls.class_room_name,
                    "level": cls.level,
                    "section": cls.section.name if cls.section else None
                })

            return jsonify({"success": True, "classes": classes_data}), 200

        except Exception as e:
            # print(f"Error getting classes by subject: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    def validate_exam_date(exam_date, term_id):
        """
        Validate that exam date is within term limits and not less than current year
        Returns: {"valid": bool, "message": str}
        """
        try:
            # Check if exam date is in the past (before current year)
            current_year = date.today().year
            if exam_date.year < current_year:
                return {
                    "valid": False,
                    "message": f"Exam date cannot be before the current year ({current_year})"
                }

            # Get the school term
            term = SchoolTerm.query.get(term_id)
            if not term:
                return {
                    "valid": False,
                    "message": "Invalid school term"
                }

            # Check if exam date is within term dates
            if exam_date.date() < term.start_date:
                return {
                    "valid": False,
                    "message": f"Exam date cannot be before term start date ({term.start_date.strftime('%Y-%m-%d')})"
                }

            if exam_date.date() > term.end_date:
                return {
                    "valid": False,
                    "message": f"Exam date cannot be after term end date ({term.end_date.strftime('%Y-%m-%d')})"
                }

            return {"valid": True, "message": "Valid exam date"}

        except Exception as e:
            return {
                "valid": False,
                "message": f"Error validating exam date: {str(e)}"
            }

    @app.route("/admin/exams/question-count", methods=["GET"])
    @admin_required
    def get_question_count():
        """Get the count of questions for a specific subject and class combination"""
        try:
            subject_id = request.args.get("subject_id")
            class_room_id = request.args.get("class_room_id")

            # Validate required parameters
            if not subject_id or not class_room_id:
                return jsonify({"success": False, "message": "Subject and class are required"}), 400

            # Count questions for this subject and class combination
            from models.question import Question
            question_count = Question.query.filter_by(
                subject_id=subject_id,
                class_room_id=class_room_id
            ).count()

            return jsonify({
                "success": True,
                "question_count": question_count
            }), 200

        except Exception as e:
            # print(f"Error getting question count: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/exams/<exam_id>", methods=["GET"])
    @admin_required
    def get_exam(exam_id):
        """Get exam details for editing"""
        try:
            # Debug: Print session info
            # print(f"Session user_id: {session.get('user_id')}")
            user = User.query.get(session["user_id"])
            # # print(f"User: {user}, Role: {user.role if user else 'None'}")

            exam = Exam.query.get(exam_id)

            if not exam:
                # print("Exam not found")
                return jsonify({"success": False, "message": "Exam not found"}), 404

            # Return exam details
            # exam_data = {
            #     "id": exam.id,
            #     "name": exam.name,
            #     "exam_type": exam.exam_type,
            #     "description": exam.description,
            #     "date": exam.date.strftime("%Y-%m-%d"),
            #     "duration_hours": exam.duration.seconds // 3600,
            #     "duration_minutes": (exam.duration.seconds % 3600) // 60,
            #     "subject_id": exam.subject_id,
            #     "class_room_id": exam.class_room_id,
            #     "school_term_id": exam.school_term_id,
            #     "invigilator_id": exam.invigilator_id,
            #     "max_score": float(exam.max_score)
            # }
            # print(exam.to_dict())
            return jsonify({"success": True, "exam": exam.to_dict()}), 200

        except Exception as e:
            # print(f"Error getting exam: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/exams/<exam_id>", methods=["PUT"])
    @admin_required
    def update_exam(exam_id):
        try:
            # Debug: Print session info
            # print(f"Session user_id: {session.get('user_id')}")
            user = User.query.get(session["user_id"])
            # print(f"User: {user}, Role: {user.role if user else 'None'}")

            data = request.get_json()
            exam = Exam.query.get(exam_id)

            if not exam:
                return jsonify({"success": False, "message": "Exam not found"}), 404

            # Update fields
            if data.get("name"):
                exam.name = data.get("name")
            if data.get("exam_type"):
                exam.exam_type = data.get("exam_type")
            if data.get("description") is not None:
                exam.description = data.get("description")
            if data.get("date"):
                exam.date = datetime.strptime(data.get("date"), "%Y-%m-%d")
            if (
                data.get("duration_hours") is not None
                or data.get("duration_minutes") is not None
            ):
                hours = int(data.get("duration_hours", 0))
                minutes = int(data.get("duration_minutes", 0))
                exam.duration = timedelta(hours=hours, minutes=minutes)
            if data.get("subject_id"):
                exam.subject_id = data.get("subject_id")
            if data.get("school_term_id"):
                exam.school_term_id = data.get("school_term_id")
            if data.get("invigilator_id") is not None:
                exam.invigilator_id = data.get("invigilator_id")
            if data.get("max_score"):
                exam.max_score = float(data.get("max_score"))

            # Handle number_of_questions update
            if "number_of_questions" in data:
                number_of_questions = data.get("number_of_questions")
                if number_of_questions:
                    number_of_questions = int(number_of_questions)
                    # Validate that requested number doesn't exceed available questions
                    from models.question import Question
                    question_count = Question.query.filter_by(
                        subject_id=exam.subject_id,
                        class_room_id=exam.class_room_id
                    ).count()

                    if number_of_questions > question_count:
                        return jsonify({
                            "success": False,
                            "message": f"Cannot set {number_of_questions} questions. Only {question_count} questions available."
                        }), 400

                    exam.number_of_questions = number_of_questions
                else:
                    # If empty string or None, set to None (use all questions)
                    exam.number_of_questions = None

            # Handle feature flag updates
            if "calculator_enabled" in data:
                exam.calculator_enabled = data.get("calculator_enabled", False)
            if "is_on_the_go" in data:
                exam.is_on_the_go = data.get("is_on_the_go", False)
            if "save_after_completion" in data:
                exam.save_after_completion = data.get("save_after_completion", True)
            if "show_feedback" in data:
                exam.show_feedback = data.get("show_feedback", True)

            db.session.commit()

            return (
                jsonify({"success": True, "message": "Exam updated successfully"}),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error updating exam: {str(e)}")
            return (
                jsonify(
                    {"success": False,
                        "message": f"Error updating exam: {str(e)}"}
                ),
                500,
            )

    @app.route("/admin/exams/<exam_id>", methods=["DELETE"])
    @admin_required
    def delete_exam(exam_id):
        try:
            # Debug: Print session info
            # print(f"Session user_id: {session.get('user_id')}")
            user = User.query.get(session["user_id"])
            # print(f"User: {user}, Role: {user.role if user else 'None'}")

            exam = Exam.query.get(exam_id)

            if not exam:
                return jsonify({"success": False, "message": "Exam not found"}), 404

            db.session.delete(exam)
            db.session.commit()

            return (
                jsonify({"success": True, "message": "Exam deleted successfully"}),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error deleting exam: {str(e)}")
            return (
                jsonify(
                    {"success": False,
                        "message": f"Error deleting exam: {str(e)}"}
                ),
                500,
            )

    @app.route("/admin/student/<student_id>/completed-exams", methods=["GET"])
    @admin_required
    def get_student_completed_exams(student_id):
        """Get all completed exams for a student"""
        try:
            from models.associations import student_exam

            # Get all completed exams for this student
            completed = db.session.execute(
                db.select(student_exam).where(
                    student_exam.c.student_id == student_id
                )
            ).fetchall()

            exams_data = []
            for record in completed:
                exam = Exam.query.get(record.exam_id)
                if exam:
                    # Ensure score is a number, default to 0 if None
                    score = float(
                        record.score) if record.score is not None else 0.0

                    exams_data.append({
                        'exam_id': exam.id,
                        'exam_name': exam.name,
                        'subject': exam.subject.subject_name if exam.subject else 'N/A',
                        'score': score,
                        'max_score': float(exam.max_score) if exam.max_score else 0.0,
                        'completed_at': record.completed_at.strftime('%Y-%m-%d %H:%M') if record.completed_at else 'N/A',
                        'status': 'Completed'
                    })

            return jsonify({
                'success': True,
                'exams': exams_data
            }), 200

        except Exception as e:
            # print(f"Error fetching completed exams: {str(e)}")
            return jsonify({
                'success': False,
                'message': str(e)
            }), 500

    # Handle exam reset for student to give chance to re-take exam
    @app.route("/admin/exam/<exam_id>/<user_id>/reset", methods=["POST"])
    @admin_required
    def reset_exam_for_student(exam_id, user_id):
        """Reset exam for student"""
        try:
            from models.associations import student_subject, student_exam

            exam = Exam.query.get(exam_id)
            if not exam:
                return jsonify({"success": False, "message": "Exam not found"}), 404

            # Check if this is a demo user
            is_demo_user = "demo" in user_id.lower()

            if not is_demo_user:
                # Regular students - apply normal checks
                # Check if student is enrolled in the exam's subject
                enrollment = db.session.execute(
                    db.select(student_subject).where(
                        student_subject.c.student_id == user_id,
                        student_subject.c.subject_id == exam.subject_id
                    )
                ).fetchone()

                if not enrollment:
                    return jsonify({"success": False, "message": "Not enrolled in this subject"}), 403

                # Check if student has already completed this exam
                completion = db.session.execute(
                    db.select(student_exam).where(
                        student_exam.c.student_id == user_id,
                        student_exam.c.exam_id == exam_id
                    )
                ).fetchone()

                # check if the user haven't already taken the exam
                if not completion:
                    return jsonify({"success": False, "message": "You haven't taken this exam yet"}), 403
                else:
                    # Reset the exam by deleting records from multiple tables

                    # 1. Delete from student_exam table (completion record)
                    db.session.execute(
                        student_exam.delete().where(
                            student_exam.c.student_id == user_id,
                            student_exam.c.exam_id == exam_id
                        )
                    )

                    # 2. Delete from exam_records table (detailed exam record with answers)
                    from models.exam_record import ExamRecord
                    exam_record = ExamRecord.query.filter_by(
                        student_id=user_id,
                        exam_id=exam_id
                    ).first()
                    if exam_record:
                        db.session.delete(exam_record)

                    # 3. Delete from exam_sessions table (if exists)
                    from models.exam_session import ExamSession
                    exam_sessions = ExamSession.query.filter_by(
                        student_id=user_id,
                        exam_id=exam_id
                    ).all()
                    for session in exam_sessions:
                        db.session.delete(session)

                    db.session.commit()
                    return jsonify({"success": True, "message": "Exam reset successfully"}), 200
            else:
                print('None')
            return jsonify({"success": False, "message": "Invalid user"}), 403

        except Exception as e:
            db.session.rollback()
            # print(f"Error resetting exam: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/get_subjects", methods=["GET"])
    @admin_required
    def get_subjects():
        """Get all subjects for dropdowns"""
        try:
            subjects = Subject.query.all()

            # Convert to serializable format
            subjects_data = []
            for subject in subjects:
                subjects_data.append({
                    "subject_id": subject.subject_id,
                    "subject_name": subject.subject_name,
                    "subject_code": subject.subject_code,
                    "subject_category": subject.subject_category
                })

            return jsonify({"success": True, "subjects": subjects_data}), 200

        except Exception as e:
            # print(f"Error getting subjects: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/classes", methods=["GET", "POST"])
    @admin_required
    def class_management():
        current_user = User.query.get(session["user_id"])
        class_rooms = db.session.query(ClassRoom).all()
        teachers = db.session.query(User).filter_by(role="staff").all()

        # Compute real student counts per class
        student_counts = {}
        total_students = 0
        for cls in class_rooms:
            count = User.query.filter_by(
                class_room_id=cls.class_room_id, role="student", is_active=True
            ).count()
            student_counts[cls.class_room_id] = count
            total_students += count

        average_class_size = total_students // len(class_rooms) if class_rooms else 0

        return render_template(
            "admin/classes.html",
            current_user=current_user,
            class_rooms=class_rooms,
            teachers=teachers,
            student_counts=student_counts,
            sections=Section.query.all(),
            average_class_size=average_class_size,
        )

    @app.route("/admin/create_class", methods=["POST"])
    @admin_required
    def create_class():
        if request.method == "POST":
            data = request.get_json()
            class_name = data.get("className")
            form_teacher = data.get("formTeacher")
            class_capacity = data.get("classCapacity")
            description = data.get("description")
            is_custom = data.get("isCustom", False)
            academic_year = data.get("academicYear")
            is_active = data.get("isActive", True)
            number_of_students = 0

            # Validation
            if not class_name:
                return jsonify({"success": False, "message": "Class name is required"}), 400

            # Check if class name already exists
            existing_class = ClassRoom.query.filter_by(
                class_room_name=class_name).first()
            if existing_class:
                return jsonify({"success": False, "message": f"Class '{class_name}' already exists"}), 409

            # For custom classes, section and level are optional
            if is_custom:
                section_id = None
                level = None
                group = data.get("group")  # Group is still optional
            else:
                # For standard classes, section and level are required
                section_id = data.get("sectionId")
                level = data.get("level")
                group = data.get("group")

                if not section_id or not level:
                    return jsonify({"success": False, "message": "Section and level are required for standard classes"}), 400

            new_class = ClassRoom(
                class_room_name=class_name,
                form_teacher_id=form_teacher,
                class_capacity=class_capacity,
                number_of_students=number_of_students,
                section_id=section_id,
                level=level,
                group=group,
                academic_year=academic_year,
                is_active=is_active,
            )
            db.session.add(new_class)
            db.session.commit()
            return (
                jsonify({"success": True, "message": "Class created successfully"}),
                200,
            )

    # Handle editing class - PUT requests
    @app.route("/admin/update/class/<class_id>", methods=["PUT"])
    @admin_required
    def edit_class(class_id):
        try:
            data = request.get_json()
            class_room = ClassRoom.query.get(class_id)

            if not class_room:
                return jsonify({"success": False, "message": "Class not found"}), 404

            # Update class fields
            class_room.class_room_name = data.get(
                "class_name", class_room.class_room_name
            )
            class_room.form_teacher_id = data.get(
                "form_teacher", class_room.form_teacher_id
            )
            class_room.class_capacity = data.get(
                "class_capacity", class_room.class_capacity
            )
            class_room.section_id = data.get(
                "section_id", class_room.section_id)
            class_room.level = data.get("level", class_room.level)
            class_room.group = data.get("group", class_room.group)
            class_room.academic_year = data.get(
                "academic_year", class_room.academic_year
            )
            class_room.is_active = data.get("is_active", class_room.is_active)

            # Handle class rep (allow clearing with empty string or null)
            if "class_rep_id" in data:
                class_rep_id = data.get("class_rep_id") or None
                class_room.class_rep_id = class_rep_id

            db.session.commit()
            return (
                jsonify({"success": True, "message": "Class updated successfully"}),
                200,
            )
        except Exception as e:
            db.session.rollback()
            # print(f"Error updating class: {str(e)}")
            return (
                jsonify(
                    {"success": False,
                        "message": f"Error updating class: {str(e)}"}
                ),
                500,
            )

    @app.route("/admin/api/students/search", methods=["GET"])
    @admin_required
    def search_students_for_class_rep():
        try:
            from models.student import Student

            query = request.args.get("q", "").strip()
            class_id = request.args.get("class_id", "").strip()

            if not class_id:
                return jsonify({"success": True, "students": []}), 200

            students = (
                User.query.outerjoin(Student, Student.user_id == User.id)
                .filter(
                    User.class_room_id == class_id,
                    User.role == "student",
                    User.is_active == True,
                    db.or_(
                        User.first_name.ilike(f"%{query}%"),
                        User.last_name.ilike(f"%{query}%"),
                        User.username.ilike(f"%{query}%"),
                        Student.admission_number.ilike(f"%{query}%"),
                    ),
                )
                .limit(20)
                .all()
            )

            students_data = []
            for s in students:
                student_profile = Student.query.filter_by(user_id=s.id).first()
                students_data.append({
                    "id": s.id,
                    "first_name": s.first_name,
                    "last_name": s.last_name,
                    "username": s.username,
                    "admission_number": student_profile.admission_number if student_profile else "",
                })

            return jsonify({"success": True, "students": students_data}), 200
        except Exception as e:
            return jsonify({"success": False, "message": str(e)}), 500

    # New endpoint to get classes that cannot write exams
    @app.route("/admin/settings/exception-classes", methods=["GET"])
    @admin_required
    def get_exception_classes():
        """
        Get classes that cannot write exams.
        For now, this returns all inactive classes as an example.
        In a real implementation, this could be based on specific criteria.
        """
        try:
            # Get all inactive classes as exceptions
            exception_classes = ClassRoom.query.filter_by(
                is_active=False).all()

            # Convert to serializable format
            classes_data = []
            for cls in exception_classes:
                class_data = {
                    "class_room_id": cls.class_room_id,
                    "class_room_name": cls.class_room_name,
                    "section": {
                        "name": cls.section.name if cls.section else None
                    } if cls.section else None
                }
                classes_data.append(class_data)

            return jsonify({
                "success": True,
                "classes": classes_data
            }), 200

        except Exception as e:
            # print(f"Error getting exception classes: {str(e)}")
            return jsonify({
                "success": False,
                "message": f"Error getting exception classes: {str(e)}"
            }), 500

    # New endpoint to get upcoming exams
    @app.route("/admin/settings/upcoming-exams", methods=["GET"])
    @admin_required
    def get_upcoming_exams():
        """
        Get upcoming exams that haven't taken place yet.
        """
        try:
            from datetime import datetime
            # Get exams that are scheduled for today or in the future
            upcoming_exams = Exam.query.filter(
                Exam.date >= datetime.utcnow().date()).all()

            # Convert to serializable format
            exams_data = []
            for exam in upcoming_exams:
                exam_data = {
                    "id": exam.id,
                    "name": exam.name,
                    "exam_type": exam.exam_type,
                    "date": exam.date.strftime("%b %d, %Y") if exam.date else "N/A",
                    "class_room_name": exam.class_room.class_room_name if exam.class_room else "N/A",
                    "subject_name": exam.subject.subject_name if exam.subject else "N/A"
                }
                exams_data.append(exam_data)

            return jsonify({
                "success": True,
                "exams": exams_data
            }), 200

        except Exception as e:
            # print(f"Error getting upcoming exams: {str(e)}")
            return jsonify({
                "success": False,
                "message": f"Error getting upcoming exams: {str(e)}"
            }), 500

    # New endpoint to save exception settings
    @app.route("/admin/settings/save-exceptions", methods=["POST"])
    @admin_required
    def save_exception_settings():
        """
        Save exception settings for exams.
        """
        try:
            data = request.get_json()
            excluded_exams = data.get("excluded_exams", [])

            # In a real implementation, you would save these settings to a database
            # For now, we'll just log them and return success
            # print(f"Excluded exams: {excluded_exams}")

            # Here you would typically:
            # 1. Save the list of excluded exams to a database table
            # 2. Associate them with the "students_can_write_exam" permission
            # 3. Use this information when determining which exams students can access

            return jsonify({
                "success": True,
                "message": f"Saved {len(excluded_exams)} exam exceptions"
            }), 200

        except Exception as e:
            # print(f"Error saving exception settings: {str(e)}")
            return jsonify({
                "success": False,
                "message": f"Error saving exception settings: {str(e)}"
            }), 500

            db.session.rollback()
            # print(f"Error updating class: {str(e)}")
            return (
                jsonify(
                    {"success": False,
                        "message": f"Error updating class: {str(e)}"}
                ),
                500,
            )

    # Handle DELETE request for a class
    @app.route("/admin/delete/class/<class_id>", methods=["DELETE"])
    @admin_required
    def delete_class(class_id):
        try:
            class_room = ClassRoom.query.get(class_id)
            if not class_room:
                return jsonify({"success": False, "message": "Class not found"}), 404

            db.session.delete(class_room)
            db.session.commit()
            return (
                jsonify({"success": True, "message": "Class deleted successfully"}),
                200,
            )
        except Exception as e:
            db.session.rollback()
            # print(f"Error deleting class: {str(e)}")
            return (
                jsonify(
                    {"success": False,
                        "message": f"Error deleting class: {str(e)}"}
                ),
                500,
            )

    @app.route("/admin/teachers", methods=["GET", "POST"])
    @admin_required
    def teacher_management():
        current_user = User.query.get(session["user_id"])
        teachers = db.session.query(User).filter_by(role="staff").all()
        active_classes = db.session.query(
            ClassRoom).filter_by(is_active=True).all()
        classes = db.session.query(ClassRoom).all()
        subjects = db.session.query(Subject).all()
        assigned_classes = [
            class_room
            for class_room in active_classes
            if class_room.form_teacher_id != None
        ]
        # print(assigned_classes)
        return render_template(
            "admin/teachers.html",
            current_user=current_user,
            teachers=[
                {
                    "id": teacher.id,
                    "first_name": teacher.first_name,
                    "last_name": teacher.last_name,
                    "email": teacher.email,
                    "phone": "+1 (555) 123-4567",
                    # "subject": teacher.subject,
                    "class_room": teacher.class_room,
                    "assigned_classes": [
                        assigned_class
                        for assigned_class in assigned_classes
                        if assigned_class.form_teacher_id == teacher.id
                    ],
                }
                for teacher in teachers
            ],
            active_classes=active_classes,
            subjects=subjects,
            classes=classes,
        )

    @app.route("/admin/profile", methods=["GET", "POST"])
    @admin_required
    def admin_profile():
        """
        Admin profile view and update endpoint.

        - GET: renders the admin profile page (template `admin/profile.html`) with `current_user`.
        - POST: expects JSON payload to update fields on the current admin user.
          Accepts: first_name, last_name, email, gender, dob (YYYY-MM-DD), image (url/path), password.
          Returns JSON with updated user data on success.
        """
        current_user = User.query.get(session["user_id"])
        # print(current_user)

        # POST: update profile
        if request.method == "POST":
            try:
                data = request.get_json(silent=True) or {}

                # Read incoming fields (only update if provided)
                first_name = data.get("first_name")
                last_name = data.get("last_name")
                email = data.get("email")
                gender = data.get("gender")
                dob_str = data.get("dob")
                image = data.get("image")
                password = data.get("password")

                if first_name is not None:
                    current_user.first_name = first_name
                if last_name is not None:
                    current_user.last_name = last_name
                if email is not None:
                    current_user.email = email
                if gender is not None:
                    current_user.gender = gender

                # Parse and set date of birth if provided
                if dob_str:
                    try:
                        current_user.dob = datetime.strptime(
                            dob_str, "%Y-%m-%d").date()
                    except ValueError:
                        return (
                            jsonify(
                                {
                                    "success": False,
                                    "message": "Invalid date format. Use YYYY-MM-DD",
                                }
                            ),
                            400,
                        )

                if image is not None:
                    current_user.image = image

                # Update password if provided
                if password:
                    current_user.set_password(password)

                db.session.commit()

                user_data = {
                    "id": current_user.id,
                    "username": current_user.username,
                    "first_name": current_user.first_name,
                    "last_name": current_user.last_name,
                    "email": current_user.email,
                    "gender": current_user.gender,
                    "dob": (
                        current_user.dob.strftime("%Y-%m-%d")
                        if current_user.dob
                        else None
                    ),
                    "image": current_user.image,
                    "role": current_user.role,
                }

                return (
                    jsonify(
                        {
                            "success": True,
                            "message": "Profile updated successfully",
                            "user": user_data,
                        }
                    ),
                    200,
                )

            except Exception as e:
                db.session.rollback()
                # print(f"Error updating admin profile: {str(e)}")
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": f"Error updating profile: {str(e)}",
                        }
                    ),
                    500,
                )

        # GET: render profile page
        return render_template("admin/profile.html", current_user=current_user)

    # ===============================
    # ===============================
    # STUDENTS MANAGEMENT
    # ===============================
    # ===============================
    # ===============================
    # ===============================
    # STUDENTS MANAGEMENT
    # ===============================
    # ===============================
    @app.route("/admin/students", methods=["GET", "POST"])
    @admin_required
    def student_management():
        current_user = User.query.get(session["user_id"])

        # Fetch all students with their associated user and class data
        # Using a different approach to ensure we get all students even if some joins fail
        students_query = db.session.query(Student, User, ClassRoom).\
            join(User, Student.user_id == User.id).\
            outerjoin(ClassRoom, User.class_room_id == ClassRoom.class_room_id)
        
        students_result = students_query.all()

        # Fetch active classes for dropdowns
        active_classes = ClassRoom.query.filter_by(is_active=True).all()

        # Calculate stats
        total_students = len(students_result)
        active_enrollments = sum(1 for s in students_result if s.User.is_active)

        # Format student data for template
        students_data = []
        for student_row in students_result:
            student, user, class_room = student_row if len(student_row) == 3 else (student_row[0], student_row[1], None)
            # Calculate attendance and performance (placeholders for now)
            attendance_rate = 95  # Placeholder
            performance_score = student.performance if student.performance else 0

            students_data.append({
                "id": student.id,
                "user_id": student.user_id,
                "first_name": user.first_name,
                "last_name": user.last_name,
                "email": user.email,
                "phone": student.parent_phone,  # Using parent phone as contact
                "admission_number": student.admission_number,
                "class_name": class_room.class_room_name if class_room else "Unassigned",
                "class_id": user.class_room_id if user.class_room_id else "",
                "performance": performance_score,
                "attendance": attendance_rate,
                "status": "Active" if user.is_active else "Inactive",
                "image": user.image,
                "gender": user.gender,
                "dob": user.dob.strftime("%Y-%m-%d") if user.dob else "",
                "parent_name": student.parent_name,
                "parent_email": student.parent_email,
                "address": student.address
            })

        return render_template(
            "admin/students.html",
            current_user=current_user,
            students=students_data,
            classes=active_classes,
            stats={
                "total_students": total_students,
                "active_enrollments": active_enrollments,
                "avg_performance": 85,  # Placeholder
                "attendance_rate": 94  # Placeholder
            }
        )

    @app.route("/admin/students/create", methods=["POST"])
    @admin_required
    def create_student():
        try:
            # Handle both JSON and Multipart/Form-Data
            if request.is_json:
                data = request.get_json()
            else:
                data = request.form.to_dict()

            # Check if email already exists
            if data.get("email") and User.query.filter_by(email=data.get("email")).first():
                return jsonify({"success": False, "message": "Email already exists"}), 400

            # Generate username
            username = User.generate_username("student")

            # Handle Image Upload
            image_path = None
            if 'image' in request.files:
                file = request.files['image']
                if file and file.filename != '':
                    original = secure_filename(file.filename)
                    if not original:
                        return jsonify({"success": False, "message": "Invalid filename"}), 400

                    if not allowed_file(original):
                        return jsonify({"success": False, "message": "Invalid file type. Only PNG, JPG, JPEG, GIF, and WEBP are allowed"}), 400

                    ext = original.rsplit('.', 1)[1].lower()
                    unique_name = f"{uuid.uuid4().hex}.{ext}"

                    from utils.paths import get_profile_images_dir
                    upload_dir = get_profile_images_dir()

                    file_path = os.path.join(upload_dir, unique_name)
                    file.save(file_path)
                    image_path = f"/uploads/profile_images/{unique_name}"

            # Create User record
            new_user = User(
                username=username,
                first_name=data.get("first_name"),
                last_name=data.get("last_name"),
                email=data.get("email"),
                gender=data.get("gender"),
                dob=datetime.strptime(data.get("dob"), "%Y-%m-%d").date(),
                role="student",
                class_room_id=data.get("class_id"),
                is_active=True,
                image=image_path
            )
            new_user.set_password("student123")  # Default password
            db.session.add(new_user)
            db.session.flush()  # Get user ID

            # Create Student record
            new_student = Student(
                user_id=new_user.id,
                admission_number=data.get("student_id") or username,
                admission_date=datetime.utcnow().date(),
                parent_name=data.get("parent_name"),  # Can be added to form
                parent_phone=data.get("phone"),
                parent_email=data.get("parent_email"),  # Can be added to form
                address=data.get("address")
            )
            db.session.add(new_student)

            db.session.commit()

            return jsonify({
                "success": True,
                "message": "Student created successfully",
                "student": {
                    "id": new_student.id,
                    "name": new_user.full_name()
                }
            }), 201

        except Exception as e:
            db.session.rollback()
            # print(f"Error creating student: {str(e)}")
            return jsonify({"success": False, "message": f"Error creating student: {str(e)}"}), 500

    @app.route("/admin/students/update/<student_id>", methods=["PUT", "POST"])
    @admin_required
    def update_student(student_id):
        try:
            # Handle both JSON and Multipart/Form-Data
            if request.is_json:
                data = request.get_json()
            else:
                data = request.form.to_dict()

            student = Student.query.get(student_id)

            if not student:
                return jsonify({"success": False, "message": "Student not found"}), 404

            user = student.user

            # Handle Image Upload
            if 'image' in request.files:
                file = request.files['image']
                if file and file.filename != '':
                    original = secure_filename(file.filename)
                    if not original:
                        return jsonify({"success": False, "message": "Invalid filename"}), 400

                    if not allowed_file(original):
                        return jsonify({"success": False, "message": "Invalid file type. Only PNG, JPG, JPEG, GIF, and WEBP are allowed"}), 400

                    ext = original.rsplit('.', 1)[1].lower()
                    unique_name = f"{uuid.uuid4().hex}.{ext}"

                    from utils.paths import get_profile_images_dir
                    upload_dir = get_profile_images_dir()

                    # Delete old image file if it exists
                    if user.image and user.image.startswith('/uploads/'):
                        from utils.paths import get_data_dir
                        old_path = os.path.join(
                            get_data_dir(), 'static', user.image.lstrip('/'))
                        if os.path.exists(old_path):
                            os.remove(old_path)

                    file_path = os.path.join(upload_dir, unique_name)
                    file.save(file_path)
                    user.image = f"/uploads/profile_images/{unique_name}"

            # Update User fields
            user.first_name = data.get("first_name", user.first_name)
            user.last_name = data.get("last_name", user.last_name)
            user.email = data.get("email", user.email)
            user.gender = data.get("gender", user.gender)
            if data.get("dob"):
                user.dob = datetime.strptime(
                    data.get("dob"), "%Y-%m-%d").date()
            user.class_room_id = data.get("class_id", user.class_room_id)

            # Update Student fields
            student.admission_number = data.get(
                "student_id", student.admission_number)
            student.parent_phone = data.get("phone", student.parent_phone)
            student.address = data.get("address", student.address)

            db.session.commit()

            return jsonify({"success": True, "message": "Student updated successfully"}), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error updating student: {str(e)}")
            return jsonify({"success": False, "message": f"Error updating student: {str(e)}"}), 500

    @app.route("/admin/students/delete/<student_id>", methods=["DELETE"])
    @admin_required
    def delete_student(student_id):
        try:
            student = Student.query.get(student_id)
            if not student:
                return jsonify({"success": False, "message": "Student not found"}), 404

            # Delete associated user (cascade should handle student record, but let's be safe)
            user = student.user
            # This should cascade delete the student record
            db.session.delete(user)

            db.session.commit()
            return jsonify({"success": True, "message": "Student deleted successfully"}), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error deleting student: {str(e)}")
            return jsonify({"success": False, "message": f"Error deleting student: {str(e)}"}), 500

    @app.route("/admin/students/enroll", methods=["POST"])
    @admin_required
    def enroll_students():
        try:
            data = request.get_json()
            class_id = data.get("class_id")
            student_ids = data.get("student_ids", [])

            if not class_id or not student_ids:
                return jsonify({"success": False, "message": "Class and students are required"}), 400

            class_room = ClassRoom.query.get(class_id)
            if not class_room:
                return jsonify({"success": False, "message": "Class not found"}), 404

            success_count = 0
            for student_id in student_ids:
                # Find student by ID and update their user's class_room_id
                student = Student.query.get(student_id)
                if student and student.user:
                    student.user.class_room_id = class_id
                    success_count += 1

            db.session.commit()

            return jsonify({
                "success": True,
                "message": f"Successfully enrolled {success_count} students in {class_room.class_room_name}"
            }), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error enrolling students: {str(e)}")
            return jsonify({"success": False, "message": f"Error enrolling students: {str(e)}"}), 500

    # ===============================
    # ===============================
    # SUBJECT MANAGEMENT
    # ===============================
    # ===============================
    @app.route("/admin/subjects", methods=["GET", "POST"])
    @admin_required
    def subject_management():
        if request.method == "POST":
            try:
                data = request.get_json()
                # print(data)

                # Handle subject_code - convert empty string to None for unique constraint
                subject_code = data.get("subject_code", "").strip()
                subject_code = subject_code if subject_code else None

                # Create subject once
                new_subject = Subject(
                    subject_name=data.get("subject_name"),
                    subject_code=subject_code,
                    description=data.get("description"),
                    grade=data.get("grade"),
                    academic_year=data.get("academic_year"),
                    academic_term=data.get("academic_term"),
                    subject_head_id=data.get("subject_head"),
                    is_active=data.get("isActive", True),
                    credit_hours=data.get("credit_hours"),
                    icon_name=data.get("icon_name"),
                    subject_category=data.get("category"),
                    category_colors=data.get("category_colors"),
                )
                db.session.add(new_subject)
                db.session.flush()  # Get the subject ID without committing

                # Record all classroom subjects - link classes to the subject via association table
                class_levels = data.get("grade_levels", [])
                for class_level in class_levels:
                    # Get classroom by class level name
                    class_room = ClassRoom.query.filter_by(
                        class_room_name=class_level
                    ).first()
                    if not class_room:
                        return (
                            jsonify(
                                {
                                    "success": False,
                                    "message": f"Invalid class: {class_level}",
                                }
                            ),
                            400,
                        )

                    # Link subject to class via association table
                    from models.associations import class_subject

                    # Check if link already exists
                    existing_link = db.session.execute(
                        db.select(class_subject).where(
                            class_subject.c.class_room_id == class_room.class_room_id,
                            class_subject.c.subject_id == new_subject.subject_id,
                        )
                    ).fetchone()

                    if not existing_link:
                        db.session.execute(
                            class_subject.insert().values(
                                class_room_id=class_room.class_room_id,
                                subject_id=new_subject.subject_id,
                            )
                        )

                db.session.commit()
                return (
                    jsonify(
                        {
                            "success": True,
                            "message": "Subject created and linked to classes successfully",
                        }
                    ),
                    200,
                )
            except Exception as e:
                db.session.rollback()
                # print(f"Error creating subject: {str(e)}")
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": f"Error creating subject: {str(e)}",
                        }
                    ),
                    500,
                )
        current_user = User.query.get(session["user_id"])
        subjects = db.session.query(Subject).all()
        classes = db.session.query(ClassRoom).all()
        users = db.session.query(User).filter_by(role="staff").all()
        sections = db.session.query(Section).all()
        icons = [
            {"calculate": "Math"},
            {"science": "Science"},
            {"biotech": "Biology"},
            {"psychology": "Psychology"},
            {"memory": "Computer Science"},
            {"code": "Programming"},
            {"menu_book": "Literature"},
            {"history_edu": "History"},
            {"language": "Languages"},
            {"translate": "Translation"},
            {"music_note": "Music"},
            {"palette": "Art"},
            {"sports_soccer": "Sports"},
            {"fitness_center": "Physical Education"},
            {"business": "Business"},
            {"account_balance": "Economics"},
            {"public": "Geography"},
            {"psychology_alt": "Advanced Psychology"},
            {"architecture": "Architecture"},
            {"engineering": "Engineering"},
            {"medical_services": "Medicine"},
            {"local_library": "Library"},
            {"history": "History"},
            {"gavel": "Law"},
            {"monetization_on": "Finance"},
            {"trending_up": "Statistics"},
            {"psychology_alt": "Philosophy"},
            {"groups": "Sociology"},
            {"public": "Political Science"},
            {"biotech": "Biotechnology"},
            {"agriculture": "Agriculture"},
            {"eco": "Environmental Science"},
            {"waves": "Physics"},
            {"science": "Chemistry"},
            {"biotech": "Biochemistry"},
            {"coronavirus": "Microbiology"},
            {"psychology": "Neuroscience"},
            {"computer": "Computer Science"},
            {"security": "Cybersecurity"},
            {"data_object": "Data Science"},
            {"analytics": "Analytics"},
            {"design_services": "Design"},
            {"brush": "Fine Arts"},
            {"theater_comedy": "Drama"},
            {"mic": "Public Speaking"},
            {"auto_stories": "Creative Writing"},
        ]

        # Add class information to subjects for the template
        subjects_with_classes = []
        for subject in subjects:
            subject_data = {
                "subject_id": subject.subject_id,
                "subject_name": subject.subject_name,
                "subject_code": subject.subject_code,
                "description": subject.description,
                "icon_name": subject.icon_name,
                "is_active": subject.is_active,
                "classes": [
                    {
                        "class_room_id": class_room.class_room_id,
                        "class_room_name": class_room.class_room_name,
                    }
                    for class_room in subject.classes
                ],
            }
            subjects_with_classes.append(subject_data)
        # print(sections)
        # Convert subject category colors to JSON object
        return render_template(
            "admin/subjects.html",
            current_user=current_user,
            subjects=subjects_with_classes,
            classes=classes,
            icons=icons,
            users=users,
            sections=sections,
        )

    # Handle editing subject - PUT requests
    @app.route("/admin/update/subjects/<subject_id>", methods=["PUT"])
    @admin_required
    def edit_subject(subject_id):
        # print("Editing subject: ", subject_id)
        try:
            data = request.get_json()
            subject = Subject.query.get(subject_id)

            if not subject:
                return jsonify({"success": False, "message": "Subject not found"}), 404

            # Update subject fields
            subject.subject_name = data.get(
                "subject_name", subject.subject_name)

            # Handle subject_code - convert empty string to None for unique constraint
            subject_code = data.get("subject_code", "").strip()
            subject.subject_code = subject_code if subject_code else None

            subject.description = data.get("description", subject.description)
            subject.icon_name = data.get("icon_name", subject.icon_name)
            subject.subject_category = data.get(
                "category", subject.subject_category)

            # Optional fields
            if data.get("grade"):
                subject.grade = data.get("grade")
            if data.get("academic_year"):
                subject.academic_year = data.get("academic_year")
            if data.get("academic_term"):
                subject.academic_term = data.get("academic_term")

            # Handle class associations - if grade_levels provided, update class associations
            if data.get("grade_levels"):
                # First, get all class rooms by their names
                class_rooms = []
                class_levels = data.get("grade_levels", [])
                for class_level in class_levels:
                    class_room = ClassRoom.query.filter_by(
                        class_room_name=class_level
                    ).first()
                    if not class_room:
                        return (
                            jsonify(
                                {
                                    "success": False,
                                    "message": f"Invalid class: {class_level}",
                                }
                            ),
                            400,
                        )
                    class_rooms.append(class_room)

                # Update class associations - clear existing and add new ones
                from models.associations import class_subject

                # Remove all existing associations for this subject
                db.session.execute(
                    db.delete(class_subject).where(
                        class_subject.c.subject_id == subject.subject_id
                    )
                )

                # Add new associations
                for class_room in class_rooms:
                    db.session.execute(
                        class_subject.insert().values(
                            class_room_id=class_room.class_room_id,
                            subject_id=subject.subject_id,
                        )
                    )

            if data.get("subject_head"):
                subject.subject_head_id = data.get("subject_head")
            if "isActive" in data:
                subject.is_active = data.get("isActive", True)
            if data.get("credit_hours"):
                subject.credit_hours = data.get("credit_hours")
            if data.get("category_colors"):
                subject.category_colors = data.get("category_colors")

            db.session.commit()
            return (
                jsonify(
                    {"success": True, "message": "Subject updated successfully"}),
                200,
            )
        except Exception as e:
            db.session.rollback()
            # print(f"Error updating subject: {str(e)}")
            return (
                jsonify(
                    {"success": False,
                        "message": f"Error updating subject: {str(e)}"}
                ),
                500,
            )

    # Handle DELETE request for a subject
    @app.route("/admin/delete/subjects/<subject_id>", methods=["DELETE"])
    @admin_required
    def delete_subject(subject_id):
        try:
            subject = Subject.query.get(subject_id)
            if not subject:
                return jsonify({"success": False, "message": "Subject not found"}), 404

            db.session.delete(subject)
            db.session.commit()
            return (
                jsonify(
                    {"success": True, "message": "Subject deleted successfully"}),
                200,
            )
        except Exception as e:
            db.session.rollback()
            # print(f"Error deleting subject: {str(e)}")
            return (
                jsonify(
                    {"success": False,
                        "message": f"Error deleting subject: {str(e)}"}
                ),
                500,
            )

    # ===============================
    # ===============================
    # SUBJECT-TEACHER ASSIGNMENT
    # ===============================
    # ===============================
    @app.route("/admin/assign_subject_teacher", methods=["POST"])
    @admin_required
    def assign_subject_teacher():
        try:
            data = request.get_json()
            subjects_ids: List[str] = data.get("subjects_ids")
            teacher_id = data.get("teacher_id")
            class_room_id = data.get("class_room_id")
            # print(data)

            # Validate inputs
            if not subjects_ids or not teacher_id or not class_room_id:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": "Subjects, teacher, and class are required",
                        }
                    ),
                    400,
                )

            # Check if teacher exists and is staff
            teacher = User.query.get(teacher_id)
            if not teacher or teacher.role != "staff":
                return (
                    jsonify(
                        {"success": False, "message": "Teacher not found or invalid"}
                    ),
                    404,
                )

            # Check if class exists
            class_room = ClassRoom.query.get(class_room_id)
            if not class_room:
                return jsonify({"success": False, "message": "Class not found"}), 404

            # Track successfully assigned subjects
            assigned_subjects = []
            already_assigned_subjects = []

            # Assign teacher to subject using the teacher_subject association table
            from models.associations import teacher_subject

            # Loop through all subject IDs and assign the same teacher to each
            for subject_id in subjects_ids:
                # Check if subject exists
                subject = Subject.query.get(subject_id)
                if not subject:
                    continue  # Skip invalid subjects

                # Check if assignment already exists for teacher-subject-class combination
                existing_assignment = db.session.execute(
                    db.select(teacher_subject).where(
                        teacher_subject.c.teacher_id == teacher_id,
                        teacher_subject.c.subject_id == subject_id,
                        teacher_subject.c.class_room_id == class_room_id,
                    )
                ).fetchone()

                if existing_assignment:
                    already_assigned_subjects.append(subject.subject_name)
                    continue

                # Insert the teacher-subject-class assignment
                stmt = teacher_subject.insert().values(
                    teacher_id=teacher_id,
                    subject_id=subject_id,
                    class_room_id=class_room_id,
                )
                db.session.execute(stmt)
                assigned_subjects.append(subject.subject_name)

            # Commit all changes
            db.session.commit()

            # Prepare response message
            if not assigned_subjects and already_assigned_subjects:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": f"Teacher {teacher.first_name} {teacher.last_name} is already assigned to: {', '.join(already_assigned_subjects)}",
                        }
                    ),
                    400,
                )
            elif already_assigned_subjects:
                return (
                    jsonify(
                        {
                            "success": True,
                            "message": f"Teacher {teacher.first_name} {teacher.last_name} assigned to: {', '.join(assigned_subjects)}. Already assigned to: {', '.join(already_assigned_subjects)}",
                        }
                    ),
                    200,
                )
            else:
                return (
                    jsonify(
                        {
                            "success": True,
                            "message": f"Teacher {teacher.first_name} {teacher.last_name} assigned to: {', '.join(assigned_subjects)}",
                        }
                    ),
                    200,
                )

        except Exception as e:
            db.session.rollback()
            # print(f"Error assigning subject teacher: {str(e)}")
            return (
                jsonify(
                    {"success": False,
                        "message": f"Error assigning teacher: {str(e)}"}
                ),
                500,
            )

    # List all subjects a classroom offers
    @app.route("/admin/subjects/<classroom_id>", methods=["POST"])
    @admin_required
    def subjects_for_classroom(classroom_id):
        # print("Getting subjects for classroom")
        try:
            classroom = ClassRoom.query.get(classroom_id)
            if not classroom:
                return jsonify({"success": False, "message": "Class not found"}), 404

            # Get subjects ids from class_subject association table
            from models.associations import class_subject

            subject_id_rows = (
                db.session.query(class_subject.c.subject_id)
                .filter_by(class_room_id=classroom_id)
                .all()
            )

            # Extract subject IDs from Row objects
            subject_ids = [row[0] for row in subject_id_rows]

            # Get subjects for these ids
            subjects = Subject.query.filter(
                Subject.subject_id.in_(subject_ids)).all()

            # print(subjects)
            return (
                jsonify(
                    {
                        "success": True,
                        "subjects": [subject.to_dict() for subject in subjects],
                    }
                ),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error getting subjects for classroom: {str(e)}")
            return (
                jsonify(
                    {
                        "success": False,
                        "message": f"Error getting subjects for classroom: {str(e)}",
                    }
                ),
                500,
            )

    @app.route("/admin/teacher_assignments_matrix", methods=["POST"])
    @admin_required
    def teacher_assignments_matrix():
        try:
            data = request.get_json()
            teacher_id = data.get("teacher_id")

            if not teacher_id:
                return jsonify({"success": False, "message": "Teacher ID is required"}), 400

            teacher = User.query.get(teacher_id)
            if not teacher or teacher.role != "staff":
                return jsonify({"success": False, "message": "Teacher not found or invalid"}), 404

            from models.associations import teacher_subject, class_subject

            # Get all active classes
            active_classes = ClassRoom.query.filter_by(is_active=True).order_by(ClassRoom.class_room_name).all()
            classes_data = [{"class_room_id": c.class_room_id, "class_room_name": c.class_room_name} for c in active_classes]

            # Get all subjects
            all_subjects = Subject.query.filter_by(is_active=True).order_by(Subject.subject_name).all()
            subjects_data = [{"subject_id": s.subject_id, "subject_name": s.subject_name} for s in all_subjects]

            # Get class_subject mapping: which subjects are offered in which classes
            cs_rows = db.session.execute(db.select(class_subject)).fetchall()
            class_subject_map = {}
            for row in cs_rows:
                cid = row.class_room_id
                sid = row.subject_id
                if cid not in class_subject_map:
                    class_subject_map[cid] = []
                class_subject_map[cid].append(sid)

            # Get teacher's existing assignments
            existing_rows = db.session.execute(
                db.select(teacher_subject).where(teacher_subject.c.teacher_id == teacher_id)
            ).fetchall()
            existing_assignments = [
                {"subject_id": row.subject_id, "class_room_id": row.class_room_id}
                for row in existing_rows
            ]

            # Get ALL teacher-subject-class assignments (for cross-teacher warnings)
            all_teacher_rows = db.session.execute(db.select(teacher_subject)).fetchall()
            # Batch-load teacher names to avoid N+1 queries
            teacher_ids = list(set(row.teacher_id for row in all_teacher_rows))
            teacher_map = {}
            if teacher_ids:
                teachers = User.query.filter(User.id.in_(teacher_ids)).all()
                teacher_map = {t.id: f"{t.first_name} {t.last_name}" for t in teachers}

            all_assignments = [
                {
                    "teacher_id": row.teacher_id,
                    "teacher_name": teacher_map.get(row.teacher_id, "Unknown"),
                    "subject_id": row.subject_id,
                    "class_room_id": row.class_room_id,
                }
                for row in all_teacher_rows
            ]

            return jsonify({
                "success": True,
                "classes": classes_data,
                "subjects": subjects_data,
                "class_subject_map": class_subject_map,
                "existing_assignments": existing_assignments,
                "all_assignments": all_assignments,
            }), 200

        except Exception as e:
            db.session.rollback()
            return jsonify({"success": False, "message": f"Error loading assignment matrix: {str(e)}"}), 500

    @app.route("/admin/assign_subjects_batch", methods=["POST"])
    @admin_required
    def assign_subjects_batch():
        try:
            data = request.get_json()
            teacher_id = data.get("teacher_id")
            assignments = data.get("assignments", [])

            if not teacher_id:
                return jsonify({"success": False, "message": "Teacher ID is required"}), 400

            teacher = User.query.get(teacher_id)
            if not teacher or teacher.role != "staff":
                return jsonify({"success": False, "message": "Teacher not found or invalid"}), 404

            from models.associations import teacher_subject

            assigned = []
            already_exists = []
            reassigned = []
            skipped = []

            for item in assignments:
                subject_id = item.get("subject_id")
                class_room_id = item.get("class_room_id")

                if not subject_id or not class_room_id:
                    continue

                subject = Subject.query.get(subject_id)
                if not subject:
                    skipped.append(subject_id)
                    continue

                class_room = ClassRoom.query.get(class_room_id)
                if not class_room:
                    skipped.append(class_room_id)
                    continue

                # Check if THIS teacher already has this assignment
                existing = db.session.execute(
                    db.select(teacher_subject).where(
                        teacher_subject.c.teacher_id == teacher_id,
                        teacher_subject.c.subject_id == subject_id,
                        teacher_subject.c.class_room_id == class_room_id,
                    )
                ).fetchone()

                if existing:
                    already_exists.append(f"{subject.subject_name} ({class_room.class_room_name})")
                    continue

                # Check if ANOTHER teacher has this subject-class assignment
                conflict = db.session.execute(
                    db.select(teacher_subject).where(
                        teacher_subject.c.subject_id == subject_id,
                        teacher_subject.c.class_room_id == class_room_id,
                    )
                ).fetchone()

                if conflict:
                    # Remove the old assignment (reassign)
                    conflicting_teacher = User.query.get(conflict.teacher_id)
                    conflict_name = f"{conflicting_teacher.first_name} {conflicting_teacher.last_name}" if conflicting_teacher else "Unknown"
                    db.session.execute(
                        db.delete(teacher_subject).where(
                            teacher_subject.c.teacher_id == conflict.teacher_id,
                            teacher_subject.c.subject_id == subject_id,
                            teacher_subject.c.class_room_id == class_room_id,
                        )
                    )
                    reassigned.append({
                        "subject": subject.subject_name,
                        "class": class_room.class_room_name,
                        "from_teacher": conflict_name,
                    })

                # Insert the new assignment
                stmt = teacher_subject.insert().values(
                    teacher_id=teacher_id,
                    subject_id=subject_id,
                    class_room_id=class_room_id,
                )
                db.session.execute(stmt)
                assigned.append(f"{subject.subject_name} ({class_room.class_room_name})")

            db.session.commit()

            parts = []
            if assigned:
                parts.append(f"Assigned {len(assigned)}: {', '.join(assigned)}")
            if reassigned:
                reassign_msgs = [f"{r['subject']} ({r['class']}) from {r['from_teacher']}" for r in reassigned]
                parts.append(f"Reassigned {len(reassigned)} from other teachers: {', '.join(reassign_msgs)}")
            if already_exists:
                parts.append(f"Already assigned to this teacher ({len(already_exists)}): {', '.join(already_exists)}")
            if skipped:
                parts.append(f"Skipped {len(skipped)} invalid entries")

            message = "; ".join(parts) if parts else "No assignments to process"
            success = len(assigned) > 0

            return jsonify({
                "success": success,
                "message": message,
                "assigned_count": len(assigned),
                "reassigned_count": len(reassigned),
                "already_exists_count": len(already_exists),
            }), 200 if success else 400

        except Exception as e:
            db.session.rollback()
            return jsonify({"success": False, "message": f"Error in batch assignment: {str(e)}"}), 500

    @app.route("/admin/remove_subject_teacher", methods=["POST"])
    @admin_required
    def remove_subject_teacher():
        try:
            data = request.get_json()
            subject_id = data.get("subject_id")

            # Validate inputs
            if not subject_id:
                return (
                    jsonify({"success": False, "message": "Subject is required"}),
                    400,
                )

            # Check if subject exists
            subject = Subject.query.get(subject_id)
            if not subject:
                return jsonify({"success": False, "message": "Subject not found"}), 404

            # Remove teacher assignment
            subject.subject_head_id = None
            db.session.commit()

            return (
                jsonify(
                    {
                        "success": True,
                        "message": f"Teacher assignment removed from {subject.subject_name} successfully",
                    }
                ),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error removing subject teacher: {str(e)}")
            return (
                jsonify(
                    {
                        "success": False,
                        "message": f"Error removing teacher assignment: {str(e)}",
                    }
                ),
                500,
            )

    # Handle assigning subject head
    @app.route("/admin/assign_subject_head", methods=["POST"])
    @admin_required
    def assign_subject_head():
        try:
            data = request.get_json()
            subject_id = data.get("subject_id")
            teacher_id = data.get("teacher_id")

            # Validate inputs
            if not subject_id or not teacher_id:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": "Subject ID and teacher ID are required",
                        }
                    ),
                    400,
                )

            # Check if teacher exists
            teacher = User.query.get(teacher_id)
            if not teacher:
                return jsonify({"success": False, "message": "Teacher not found"}), 404

            # Check if subject exists
            subject = Subject.query.get(subject_id)
            if not subject:
                return jsonify({"success": False, "message": "Subject not found"}), 404

            # Check if subject already has a teacher assigned
            if subject.subject_head_id:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": "Subject already has a teacher assigned",
                        }
                    ),
                    400,
                )

            # Assign teacher to subject
            subject.subject_head_id = teacher_id
            db.session.commit()

            return (
                jsonify(
                    {
                        "success": True,
                        "message": f"Subject {subject.subject_name} assigned to {teacher.first_name} {teacher.last_name} successfully",
                    }
                ),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error assigning subject head: {str(e)}")
            return (
                jsonify(
                    {
                        "success": False,
                        "message": f"Error assigning subject head: {str(e)}",
                    }
                ),
                500,
            )

    # ===============================
    # ===============================
    # TEACHER-CLASS ASSIGNMENT
    # ===============================
    # ===============================
    @app.route("/admin/assign_teacher_class", methods=["POST"])
    @admin_required
    def assign_teacher_class():
        try:
            data = request.get_json()
            teacher_id = data.get("teacher")
            class_id = data.get("class")
            subject_ids = data.get("subjects", [])

            # Validate inputs
            if not teacher_id or not class_id:
                return (
                    jsonify(
                        {"success": False, "message": "Teacher and class are required"}
                    ),
                    400,
                )

            # Check if teacher exists and is staff
            teacher_user = User.query.get(teacher_id)
            if not teacher_user or teacher_user.role not in ["staff", "teacher"]:
                return (
                    jsonify(
                        {"success": False, "message": "Teacher not found or invalid"}
                    ),
                    404,
                )

            # Check if class exists
            class_room = ClassRoom.query.get(class_id)
            if not class_room:
                return jsonify({"success": False, "message": "Class not found"}), 404

            # Check if subjects exist
            subjects = []
            for subject_id in subject_ids:
                subject = Subject.query.get(subject_id)
                if not subject:
                    return (
                        jsonify(
                            {
                                "success": False,
                                "message": f"Subject with ID {subject_id} not found",
                            }
                        ),
                        404,
                    )
                subjects.append(subject)

            # Check if teacher is already assigned to this class
            existing_assignment = (
                db.session.query(teacher_classroom)
                .filter_by(teacher_id=teacher_id, classroom_id=class_id)
                .first()
            )

            if not existing_assignment:
                # Assign teacher to class using the association table
                assignment = teacher_classroom.insert().values(
                    teacher_id=teacher_id, classroom_id=class_id
                )
                db.session.execute(assignment)
            else:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": f"Teacher {teacher_user.first_name} {teacher_user.last_name} is already assigned to {class_room.class_room_name}",
                        }
                    ),
                    400,
                )

            # Assign teacher to subjects
            for subject in subjects:
                subject.subject_head_id = teacher_id

            db.session.commit()

            return (
                jsonify(
                    {
                        "success": True,
                        "message": f"Teacher {teacher_user.first_name} {teacher_user.last_name} assigned to {class_room.class_room_name} successfully",
                    }
                ),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error assigning teacher to class: {str(e)}")
            return (
                jsonify(
                    {
                        "success": False,
                        "message": f"Error assigning teacher to class: {str(e)}",
                    }
                ),
                500,
            )

    # ===============================
    # ===============================
    # ADMIN SETTINGS
    # ===============================
    # ===============================
    @app.route("/admin/settings", methods=["GET", "POST"])
    @admin_required
    def settings():
        current_user = User.query.get(session["user_id"])
        # Get school information if it exists
        school = School.query.first()
        return render_template(
            "admin/settings.html", current_user=current_user, school=school
        )

    @app.route("/admin/settings/generate-demo-data", methods=["POST"])
    @admin_required
    def generate_demo_data():
        try:
            from scripts.setup.initialize_all_data import generate_demo_data as run_demo_data

            run_demo_data()
            return jsonify({"success": True, "message": "Demo data generated successfully!"}), 200
        except Exception as e:
            db.session.rollback()
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/settings/toggle-demo-questions", methods=["POST"])
    @admin_required
    def toggle_demo_questions():
        """Toggle demo questions: create on enable, delete on disable."""
        from models.demo_question import DemoQuestion, DemoOption

        data = request.get_json()
        enable = data.get("enable", False)

        if enable:
            try:
                from scripts.setup.initialize_all_data import populate_demo_questions
                populate_demo_questions()
                count = DemoQuestion.query.count()
                perm = Permission.query.filter_by(permission_name="demo_question_bank").first()
                if perm:
                    perm.is_active = True
                    db.session.commit()
                else:
                    perm = Permission(
                        permission_name="demo_question_bank",
                        permission_description="Enable demo question bank for practice",
                        is_active=True,
                        created_for="student",
                    )
                    db.session.add(perm)
                    db.session.commit()
                return jsonify({"success": True, "message": f"Demo questions created. {count} questions available.", "count": count}), 200
            except Exception as e:
                db.session.rollback()
                return jsonify({"success": False, "message": str(e)}), 500
        else:
            try:
                count = DemoQuestion.query.count()
                option_count = DemoOption.query.count()
                DemoOption.query.delete()
                DemoQuestion.query.delete()
                perm = Permission.query.filter_by(permission_name="demo_question_bank").first()
                if perm:
                    perm.is_active = False
                db.session.commit()
                return jsonify({
                    "success": True,
                    "message": f"Demo test disabled. {count} questions and {option_count} options deleted.",
                    "questions_deleted": count,
                    "options_deleted": option_count,
                }), 200
            except Exception as e:
                db.session.rollback()
                return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/settings/demo-questions-status", methods=["GET"])
    @admin_required
    def demo_questions_status():
        """Get the current count of demo questions and options."""
        from models.demo_question import DemoQuestion, DemoOption
        perm = Permission.query.filter_by(permission_name="demo_question_bank").first()
        return jsonify({
            "success": True,
            "enabled": perm.is_active if perm else False,
            "question_count": DemoQuestion.query.count(),
            "option_count": DemoOption.query.count(),
        }), 200

    @app.route("/admin/settings/data-management/status", methods=["GET"])
    @admin_required
    def data_management_status():
        """Get record counts for all data types to show status badges."""
        from models.demo_question import DemoQuestion, DemoOption
        from models.teacher import Teacher
        from models.associations import class_subject

        school = School.query.first()
        school_id = school.school_id if school else None

        status = {}
        if school_id:
            status["school"] = {"count": School.query.count()}
            status["terms"] = {"count": SchoolTerm.query.filter_by(school_id=school_id).count()}
            status["assessments"] = {"count": AssessmentType.query.filter_by(school_id=school_id).count()}
            status["sections"] = {"count": Section.query.filter_by(school_id=school_id).count()}
        else:
            status["school"] = {"count": 0}
            status["terms"] = {"count": 0}
            status["assessments"] = {"count": 0}
            status["sections"] = {"count": 0}

        status["classrooms"] = {"count": ClassRoom.query.count()}
        status["subjects"] = {"count": Subject.query.count()}
        status["teachers"] = {"count": User.query.filter_by(role="staff").count()}
        status["students"] = {"count": User.query.filter_by(role="student").count()}
        status["demo_questions"] = {"count": DemoQuestion.query.count()}
        status["permissions"] = {"count": Permission.query.count()}
        links = db.session.execute(db.select(class_subject)).fetchall()
        status["class_subjects"] = {"count": len(links)}

        return jsonify({"success": True, "status": status}), 200

    @app.route("/admin/settings/data-management", methods=["POST"])
    @admin_required
    def data_management():
        """Handle individual data init/delete and bulk operations."""
        from models.demo_question import DemoQuestion, DemoOption
        from models.teacher import Teacher
        from models.associations import class_subject

        data = request.get_json()
        action = data.get("action")
        data_type = data.get("type")

        try:
            if action == "init_all":
                from scripts.setup.initialize_all_data import generate_demo_data as run_demo_data
                run_demo_data()
                return jsonify({"success": True, "message": "All data initialized successfully!"}), 200

            if action == "delete_all":
                db.session.execute(db.delete(class_subject))
                DemoOption.query.delete()
                DemoQuestion.query.delete()
                User.query.filter_by(role="staff").delete()
                User.query.filter_by(role="student").delete()
                Teacher.query.delete()
                Student.query.delete()
                Subject.query.delete()
                ClassRoom.query.delete()
                Section.query.delete()
                AssessmentType.query.delete()
                SchoolTerm.query.delete()
                Permission.query.delete()
                School.query.delete()
                db.session.commit()
                return jsonify({"success": True, "message": "All data deleted successfully!"}), 200

            if data_type == "school":
                if action == "init":
                    from scripts.setup.initialize_all_data import create_school
                    create_school()
                    return jsonify({"success": True, "message": "School info initialized."}), 200
                elif action == "delete":
                    School.query.delete()
                    db.session.commit()
                    return jsonify({"success": True, "message": "School info deleted."}), 200

            elif data_type == "terms":
                school = School.query.first()
                if not school:
                    return jsonify({"success": False, "message": "School must be initialized first."}), 400
                if action == "init":
                    from scripts.setup.initialize_all_data import create_school_terms
                    create_school_terms(school)
                    return jsonify({"success": True, "message": "Terms initialized."}), 200
                elif action == "delete":
                    SchoolTerm.query.filter_by(school_id=school.school_id).delete()
                    db.session.commit()
                    return jsonify({"success": True, "message": "Terms deleted."}), 200

            elif data_type == "assessments":
                school = School.query.first()
                if not school:
                    return jsonify({"success": False, "message": "School must be initialized first."}), 400
                if action == "init":
                    from scripts.setup.initialize_all_data import create_assessment_types
                    create_assessment_types(school)
                    return jsonify({"success": True, "message": "Assessment types initialized."}), 200
                elif action == "delete":
                    AssessmentType.query.filter_by(school_id=school.school_id).delete()
                    db.session.commit()
                    return jsonify({"success": True, "message": "Assessment types deleted."}), 200

            elif data_type == "sections":
                school = School.query.first()
                if not school:
                    return jsonify({"success": False, "message": "School must be initialized first."}), 400
                if action == "init":
                    from scripts.setup.initialize_all_data import create_sections
                    create_sections(school)
                    return jsonify({"success": True, "message": "Sections initialized."}), 200
                elif action == "delete":
                    Section.query.filter_by(school_id=school.school_id).delete()
                    db.session.commit()
                    return jsonify({"success": True, "message": "Sections deleted."}), 200

            elif data_type == "classrooms":
                if action == "init":
                    from scripts.setup.initialize_all_data import create_sections, create_classrooms
                    school = School.query.first()
                    if not school:
                        return jsonify({"success": False, "message": "School must be initialized first."}), 400
                    sections = create_sections(school)
                    create_classrooms(sections)
                    return jsonify({"success": True, "message": "Classrooms initialized."}), 200
                elif action == "delete":
                    ClassRoom.query.delete()
                    db.session.commit()
                    return jsonify({"success": True, "message": "Classrooms deleted."}), 200

            elif data_type == "subjects":
                if action == "init":
                    from scripts.setup.initialize_all_data import create_subjects
                    create_subjects()
                    return jsonify({"success": True, "message": "Subjects initialized."}), 200
                elif action == "delete":
                    Subject.query.delete()
                    db.session.commit()
                    return jsonify({"success": True, "message": "Subjects deleted."}), 200

            elif data_type == "teachers":
                if action == "init":
                    from scripts.setup.initialize_all_data import create_classrooms, create_sections, create_teachers
                    school = School.query.first()
                    if not school:
                        return jsonify({"success": False, "message": "School must be initialized first."}), 400
                    sections = create_sections(school)
                    classrooms = create_classrooms(sections)
                    create_teachers(classrooms)
                    return jsonify({"success": True, "message": "Teachers initialized."}), 200
                elif action == "delete":
                    User.query.filter_by(role="staff").delete()
                    Teacher.query.delete()
                    db.session.commit()
                    return jsonify({"success": True, "message": "Teachers deleted."}), 200

            elif data_type == "students":
                if action == "init":
                    from scripts.setup.initialize_all_data import create_classrooms, create_sections, create_students
                    school = School.query.first()
                    if not school:
                        return jsonify({"success": False, "message": "School must be initialized first."}), 400
                    sections = create_sections(school)
                    classrooms = create_classrooms(sections)
                    create_students(classrooms)
                    return jsonify({"success": True, "message": "Students initialized."}), 200
                elif action == "delete":
                    User.query.filter_by(role="student").delete()
                    Student.query.delete()
                    db.session.commit()
                    return jsonify({"success": True, "message": "Students deleted."}), 200

            elif data_type == "demo_questions":
                if action == "init":
                    from scripts.setup.initialize_all_data import populate_demo_questions
                    populate_demo_questions()
                    count = DemoQuestion.query.count()
                    return jsonify({"success": True, "message": f"Demo questions initialized. {count} questions created."}), 200
                elif action == "delete":
                    DemoOption.query.delete()
                    DemoQuestion.query.delete()
                    db.session.commit()
                    return jsonify({"success": True, "message": "Demo questions deleted."}), 200

            elif data_type == "permissions":
                if action == "init":
                    from scripts.setup.initialize_all_data import create_permissions
                    create_permissions()
                    return jsonify({"success": True, "message": "Permissions initialized."}), 200
                elif action == "delete":
                    Permission.query.delete()
                    db.session.commit()
                    return jsonify({"success": True, "message": "Permissions deleted."}), 200

            elif data_type == "class_subjects":
                if action == "init":
                    from scripts.setup.initialize_all_data import create_subjects, create_sections, create_classrooms, link_subjects_to_classes
                    school = School.query.first()
                    if not school:
                        return jsonify({"success": False, "message": "School must be initialized first."}), 400
                    sections = create_sections(school)
                    classrooms = create_classrooms(sections)
                    subjects = create_subjects()
                    link_subjects_to_classes(subjects, classrooms)
                    return jsonify({"success": True, "message": "Class-subject links initialized."}), 200
                elif action == "delete":
                    db.session.execute(db.delete(class_subject))
                    db.session.commit()
                    return jsonify({"success": True, "message": "Class-subject links deleted."}), 200

            return jsonify({"success": False, "message": f"Unknown data type: {data_type}"}), 400

        except Exception as e:
            db.session.rollback()
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/settings/skip-setup", methods=["POST"])
    @admin_required
    def skip_setup():
        try:
            school = School.query.first()
            if school:
                school.setup_skipped = True
                db.session.commit()
            return jsonify({"success": True}), 200
        except Exception as e:
            db.session.rollback()
            return jsonify({"success": False, "message": str(e)}), 500

    # Handle school information - GET and POST
    @app.route("/admin/settings/school", methods=["GET", "POST"])
    @admin_required
    def manage_school_info():
        # GET request - retrieve school information
        if request.method == "GET":
            try:
                school = School.query.first()
                if school:
                    return (
                        jsonify(
                            {
                                "success": True,
                                "school": {
                                    "school_id": school.school_id,
                                    "school_name": school.school_name,
                                    "school_code": school.school_code,
                                    "motto": school.motto,
                                    "address": school.address,
                                    "phone": school.phone,
                                    "email": school.email,
                                    "website": school.website,
                                    "principal_name": school.principal_name,
                                    "established_date": (
                                        school.established_date.strftime(
                                            "%Y-%m-%d")
                                        if school.established_date
                                        else None
                                    ),
                                    "logo": school.logo,
                                    "current_session": school.current_session,
                                    "current_term": school.current_term,
                                },
                            }
                        ),
                        200,
                    )
                else:
                    return (
                        jsonify(
                            {
                                "success": False,
                                "message": "No school information found",
                                "data": {},
                            }
                        ),
                        404,
                    )
            except Exception as e:
                # print(f"Error getting school info: {str(e)}")
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": f"An error occurred: {str(e)}",
                        }
                    ),
                    500,
                )

        # POST request - save/update school information
        try:
            # Get form data
            school_name = request.form.get("school_name")
            school_code = request.form.get("school_code")
            motto = request.form.get("motto")
            address = request.form.get("address")
            phone = request.form.get("phone")
            email = request.form.get("email")
            website = request.form.get("website")
            principal_name = request.form.get("principal_name")
            established_date_str = request.form.get("established_date")

            # Validate required fields
            if not school_name or not address or not phone or not email:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": "Please fill all required fields (School Name, Address, Phone, Email)",
                        }
                    ),
                    400,
                )

            # Parse established date if provided
            established_date = None
            if established_date_str:
                try:
                    established_date = datetime.strptime(
                        established_date_str, "%Y-%m-%d"
                    ).date()
                except ValueError:
                    return (
                        jsonify(
                            {"success": False, "message": "Invalid date format"}),
                        400,
                    )

            # Handle logo upload
            logo_path = None
            if "logo" in request.files:
                file = request.files["logo"]
                if file and file.filename != "" and allowed_file(file.filename):
                    # Create upload directory if it doesn't exist
                    logo_folder = current_app.config["SCHOOL_LOGO_FOLDER"]
                    os.makedirs(logo_folder, exist_ok=True)

                    # Secure the filename and save
                    filename = secure_filename(file.filename)
                    # Add timestamp to avoid filename conflicts
                    filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}"
                    file_path = os.path.join(logo_folder, filename)
                    file.save(file_path)

                    # Store relative path for database
                    logo_path = os.path.join(
                        "uploads", "school_logos", filename)

            # Check if school already exists
            school = School.query.first()

            if school:
                # Update existing school
                school.school_name = school_name
                school.school_code = school_code
                school.motto = motto
                school.address = address
                school.phone = phone
                school.email = email
                school.website = website
                school.principal_name = principal_name
                school.established_date = established_date

                # Update logo only if new file was uploaded
                if logo_path:
                    # Delete old logo if exists
                    if school.logo:
                        from utils.paths import get_data_dir
                        old_logo_path = os.path.join(
                            get_data_dir(), "static", school.logo
                        )
                        if os.path.exists(old_logo_path):
                            os.remove(old_logo_path)
                    school.logo = logo_path

                db.session.commit()
                message = "School information updated successfully"
            else:
                # Create new school
                school = School(
                    school_name=school_name,
                    school_code=school_code,
                    motto=motto,
                    address=address,
                    phone=phone,
                    email=email,
                    website=website,
                    principal_name=principal_name,
                    established_date=established_date,
                    logo=logo_path,
                    current_session="2024-2025",  # Default value
                    current_term=1,  # Default value
                )
                db.session.add(school)
                db.session.commit()
                message = "School information created successfully"

            return (
                jsonify(
                    {
                        "success": True,
                        "message": message,
                        "school": {
                            "school_id": school.school_id,
                            "school_name": school.school_name,
                            "school_code": school.school_code,
                            "motto": school.motto,
                            "address": school.address,
                            "phone": school.phone,
                            "email": school.email,
                            "website": school.website,
                            "principal_name": school.principal_name,
                            "established_date": (
                                school.established_date.strftime("%Y-%m-%d")
                                if school.established_date
                                else None
                            ),
                            "logo": school.logo,
                        },
                    }
                ),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error saving school info: {str(e)}")
            return (
                jsonify(
                    {
                        "success": False,
                        "message": f"An error occurred while saving school information: {str(e)}",
                    }
                ),
                500,
            )

    # manage sections
    @app.route("/admin/settings/sections", methods=["GET", "POST"])
    @admin_required
    def section_management():
        if request.method == "POST":
            data = request.get_json()
            # # print(data)
            section_name = data["name"]
            section_abbreviation = data["abbreviation"]
            section_level = data.get("level")  # Optional field, but can be 0
            
            # Get the current school
            school = School.query.first()
            if not school:
                body = {
                    "success": False,
                    "message": "No school found. Please configure school settings first.",
                }
                return body, 400
            
            section = Section(name=section_name,
                              abbreviation=section_abbreviation,
                              level=section_level,
                              school_id=school.school_id)
            # check if section already exists
            if Section.query.filter_by(name=section_name).first():
                body = {
                    "success": False,
                    "message": "Section already exists",
                }
                return body, 400
            db.session.add(section)
            db.session.commit()

            body = {
                "success": True,
                "message": "Section created successfully",
                "section": section.to_dict(),
            }

            return body, 200

        sections = Section.query.all()
        # # print("Section:", sections)
        body = {
            "success": True,
            "sections": [section.to_dict() for section in sections],
        }
        return body, 200

    @app.route("/admin/settings/sections/<section_id>", methods=["PUT", "DELETE"])
    @admin_required
    def manage_section(section_id):
        if request.method == "PUT":
            try:
                data = request.get_json()
                section = Section.query.get(section_id)
                if not section:
                    return jsonify({"success": False, "message": "Section not found"}), 404
                
                # Update section fields
                section.name = data.get("name", section.name)
                section.abbreviation = data.get("abbreviation", section.abbreviation)
                section.level = data.get("level", section.level)
                
                db.session.commit()
                return (
                    jsonify(
                        {"success": True, "message": "Section updated successfully", "section": section.to_dict()}),
                    200,
                )
            except Exception as e:
                db.session.rollback()
                # print(f"Error updating section: {str(e)}")
                return (
                    jsonify(
                        {"success": False,
                            "message": f"Error updating section: {str(e)}"}
                    ),
                    500,
                )
        
        elif request.method == "DELETE":
            try:
                section = Section.query.get(section_id)
                if not section:
                    return jsonify({"success": False, "message": "Section not found"}), 404
                db.session.delete(section)
                db.session.commit()
                return (
                    jsonify(
                        {"success": True, "message": "Section deleted successfully"}),
                    200,
                )
            except Exception as e:
                db.session.rollback()
                # print(f"Error deleting section: {str(e)}")
                return (
                    jsonify(
                        {"success": False,
                            "message": f"Error deleting section: {str(e)}"}
                    ),
                    500,
                )

    # manage terms
    @app.route("/admin/settings/terms", methods=["GET", "POST"])
    @admin_required
    def term_management():
        if request.method == "POST":
            try:
                data = request.get_json()
                # print(data)
                term_name = data.get("term_name")
                academic_session = data.get("academic_session")

                # Get school_id from the database
                school = School.query.first()
                if not school:
                    body = {
                        "success": False,
                        "message": "Please configure your school settings first",
                    }
                    return jsonify(body), 400

                # convert first term_name to datetime
                term_start_date = datetime.strptime(
                    data.get("start_date"), "%Y-%m-%d")
                term_end_date = datetime.strptime(
                    data.get("end_date"), "%Y-%m-%d")

                term = SchoolTerm(
                    term_name=term_name,
                    start_date=term_start_date,
                    end_date=term_end_date,
                    academic_session=academic_session,
                    school_id=school.school_id,
                    open_days=data.get("open_days"),
                )
                # check if term already exists
                if SchoolTerm.query.filter_by(term_name=term_name).first():
                    body = {
                        "success": False,
                        "message": "Term already exists",
                    }
                    return jsonify(body), 404
                db.session.add(term)
                db.session.commit()

                body = {
                    "success": True,
                    "message": "Term created successfully",
                    "term": {
                        "term_id": term.term_id,
                        "term_name": term.term_name,
                        "academic_session": term.academic_session,
                        "start_date": term.start_date.strftime("%Y-%m-%d"),
                        "end_date": term.end_date.strftime("%Y-%m-%d"),
                        "is_active": term.is_active,
                        "is_current": term.is_current,
                        "open_days": term.open_days,
                    },
                }

                return jsonify(body), 200
            except Exception as e:
                # print("Errors")
                # print(e)
                body = {
                    "success": False,
                    "message": "Please configure your school settings first",
                }
                return jsonify(body), 500

        terms = SchoolTerm.query.all()
        body = {
            "success": True,
            "terms": [
                {
                    "term_id": term.term_id,
                    "term_name": term.term_name,
                    "academic_session": term.academic_session,
                    "start_date": term.start_date.strftime("%Y-%m-%d"),
                    "end_date": term.end_date.strftime("%Y-%m-%d"),
                    "is_active": term.is_active,
                    "is_current": term.is_current,
                    "open_days": term.open_days,
                }
                for term in terms
            ],
        }

        return body, 200

    # Update term
    @app.route("/admin/settings/terms/<term_id>", methods=["PUT"])
    @admin_required
    def update_term(term_id):
        try:
            data = request.get_json()

            # Find the term
            term = SchoolTerm.query.get(term_id)
            if not term:
                return jsonify({"success": False, "message": "Term not found"}), 404

            # Update fields
            term.term_name = data.get("term_name", term.term_name)
            term.academic_session = data.get(
                "academic_session", term.academic_session)
            term.is_active = data.get("is_active", term.is_active)
            term.is_current = data.get("is_current", term.is_current)

            # Update dates if provided
            if data.get("start_date"):
                term.start_date = datetime.strptime(
                    data.get("start_date"), "%Y-%m-%d")
            if data.get("end_date"):
                term.end_date = datetime.strptime(
                    data.get("end_date"), "%Y-%m-%d")

            # Update open_days if provided
            if "open_days" in data:
                term.open_days = data.get("open_days")

            # If setting as current, unset other current terms
            if data.get("is_current"):
                SchoolTerm.query.filter(SchoolTerm.term_id != term_id).update(
                    {"is_current": False}
                )

            db.session.commit()

            return (
                jsonify(
                    {
                        "success": True,
                        "message": "Term updated successfully",
                        "term": {
                            "term_id": term.term_id,
                            "term_name": term.term_name,
                            "academic_session": term.academic_session,
                            "start_date": term.start_date.strftime("%Y-%m-%d"),
                            "end_date": term.end_date.strftime("%Y-%m-%d"),
                            "is_active": term.is_active,
                            "is_current": term.is_current,
                            "open_days": term.open_days,
                        },
                    }
                ),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error updating term: {str(e)}")
            return (
                jsonify(
                    {"success": False,
                        "message": f"Error updating term: {str(e)}"}
                ),
                500,
            )

    # Delete term
    @app.route("/admin/settings/terms/<term_id>", methods=["DELETE"])
    @admin_required
    def delete_term(term_id):
        try:
            term = SchoolTerm.query.get(term_id)
            if not term:
                return jsonify({"success": False, "message": "Term not found"}), 404

            db.session.delete(term)
            db.session.commit()

            return (
                jsonify({"success": True, "message": "Term deleted successfully"}),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error deleting term: {str(e)}")
            return (
                jsonify(
                    {"success": False,
                        "message": f"Error deleting term: {str(e)}"}
                ),
                500,
            )

    # ===============================
    # Assessment Types Management
    # ===============================
    @app.route("/admin/settings/assessments", methods=["GET", "POST"])
    @admin_required
    def manage_assessments():
        """Get all assessment types or create a new one"""
        from models.assessment_type import AssessmentType

        if request.method == "POST":
            try:
                data = request.get_json()
                name = data.get("name")
                code = data.get("code")
                max_score = data.get("max_score")
                order = data.get("order", 1)
                is_cbt_enabled = data.get("is_cbt_enabled", False)
                description = data.get("description", "")

                if not name or not code or max_score is None:
                    return jsonify({"success": False, "message": "Missing required fields"}), 400

                # Get school
                school = School.query.first()

                # Check if code already exists
                existing = AssessmentType.query.filter_by(
                    code=code, school_id=school.school_id).first()
                if existing:
                    return jsonify({"success": False, "message": "Assessment code already exists"}), 400

                # Create new assessment type
                assessment = AssessmentType(
                    name=name,
                    code=code,
                    max_score=float(max_score),
                    order=int(order),
                    is_cbt_enabled=bool(is_cbt_enabled),
                    description=description,
                    school_id=school.school_id,
                    is_active=True
                )

                db.session.add(assessment)
                db.session.commit()

                return jsonify({
                    "success": True,
                    "message": "Assessment type created successfully",
                    "assessment": assessment.to_dict()
                }), 201

            except Exception as e:
                db.session.rollback()
                # print(f"Error creating assessment type: {str(e)}")
                return jsonify({"success": False, "message": str(e)}), 500

        # GET: Return all assessment types
        try:
            school = School.query.first()
            assessments = AssessmentType.query.filter_by(
                school_id=school.school_id).order_by(AssessmentType.order).all()

            return jsonify({
                "success": True,
                "assessments": [a.to_dict() for a in assessments]
            }), 200

        except Exception as e:
            # print(f"Error fetching assessment types: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/settings/assessments/<assessment_id>", methods=["PUT"])
    @admin_required
    def update_assessment(assessment_id):
        """Update an assessment type"""
        from models.assessment_type import AssessmentType

        try:
            assessment = AssessmentType.query.get(assessment_id)
            if not assessment:
                return jsonify({"success": False, "message": "Assessment type not found"}), 404

            data = request.get_json()

            # Update fields
            if "name" in data:
                assessment.name = data["name"]
            if "max_score" in data:
                assessment.max_score = float(data["max_score"])
            if "order" in data:
                assessment.order = int(data["order"])
            if "is_cbt_enabled" in data:
                assessment.is_cbt_enabled = bool(data["is_cbt_enabled"])
            if "description" in data:
                assessment.description = data["description"]
            if "is_active" in data:
                assessment.is_active = bool(data["is_active"])

            db.session.commit()

            return jsonify({
                "success": True,
                "message": "Assessment type updated successfully",
                "assessment": assessment.to_dict()
            }), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error updating assessment type: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/settings/assessments/<assessment_id>", methods=["DELETE"])
    @admin_required
    def delete_assessment(assessment_id):
        """Delete an assessment type"""
        from models.assessment_type import AssessmentType

        try:
            assessment = AssessmentType.query.get(assessment_id)
            if not assessment:
                return jsonify({"success": False, "message": "Assessment type not found"}), 404

            # Check if there are grades using this assessment
            from models.grade import Grade
            grades_count = Grade.query.filter_by(
                assessment_name=assessment.name).count()

            if grades_count > 0:
                return jsonify({
                    "success": False,
                    "message": f"Cannot delete. {grades_count} grades are using this assessment type."
                }), 400

            db.session.delete(assessment)
            db.session.commit()

            return jsonify({
                "success": True,
                "message": "Assessment type deleted successfully"
            }), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error deleting assessment type: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    # Permissions and privileges
    @app.route("/admin/settings/permissions", methods=["GET", "POST"])
    @admin_required
    def manage_permissions():
        # POST: create or update a permission
        if request.method == "POST":
            try:
                data = request.get_json() or {}
                permission_id = data.get("permission_id")
                name = data.get("permission_name")
                description = data.get("permission_description", "")
                is_active = data.get("is_active", True)
                created_for = data.get("created_for", "system")

                # Validation
                if not name:
                    return (
                        jsonify(
                            {"success": False, "message": "permission_name is required"}
                        ),
                        400,
                    )

                # Update existing permission
                if permission_id:
                    perm = Permission.query.get(permission_id)
                    if not perm:
                        return (
                            jsonify(
                                {"success": False, "message": "Permission not found"}
                            ),
                            404,
                        )

                    perm.permission_name = name or perm.permission_name
                    perm.permission_description = (
                        description or perm.permission_description
                    )
                    perm.is_active = bool(is_active)
                    perm.created_for = created_for or perm.created_for

                    db.session.commit()
                    return (
                        jsonify(
                            {
                                "success": True,
                                "message": "Permission updated successfully",
                                "permission": perm.to_dict(),
                            }
                        ),
                        200,
                    )

                # Create new permission (ensure unique name)
                existing = Permission.query.filter_by(
                    permission_name=name).first()
                if existing:
                    return (
                        jsonify(
                            {"success": False, "message": "Permission already exists"}
                        ),
                        400,
                    )

                new_perm = Permission(
                    permission_name=name,
                    permission_description=description,
                    is_active=bool(is_active),
                    created_for=created_for,
                )
                db.session.add(new_perm)
                db.session.commit()

                return (
                    jsonify(
                        {
                            "success": True,
                            "message": "Permission created successfully",
                            "permission": new_perm.to_dict(),
                        }
                    ),
                    200,
                )

            except Exception as e:
                db.session.rollback()
                # print(f"Error saving permission: {str(e)}")
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": f"Error saving permission: {str(e)}",
                        }
                    ),
                    500,
                )

        # GET: return list of permissions
        try:
            permissions = Permission.query.all()
            return (
                jsonify(
                    {
                        "success": True,
                        "message": "Permissions retrieved successfully",
                        "permissions": [
                            permission.to_dict() for permission in permissions
                        ],
                    }
                ),
                200,
            )

        except Exception as e:
            # print(f"Error retrieving permissions: {str(e)}")
            return (
                jsonify(
                    {
                        "success": False,
                        "message": f"Error retrieving permissions: {str(e)}",
                    }
                ),
                    500,
                )

    # CBT Configuration
    @app.route("/admin/settings/cbt-config", methods=["GET", "POST"])
    @admin_required
    def manage_cbt_config():
        """Get or update CBT configuration settings."""
        from models import School

        try:
            school = School.query.first()
            if not school:
                school = School(school_name="School")
                db.session.add(school)
                db.session.commit()

            if request.method == "POST":
                data = request.get_json() or {}

                # Store CBT config in school record or a dedicated config
                # For now we use a simple JSON approach via a school setting
                # We'll use the existing school fields or add a config_json field
                # Since we don't have a dedicated config table, we'll use
                # Permission-based approach for some and direct flags for others

                # Calculator enabled default
                if "calculator_enabled" in data:
                    _set_or_create_permission(
                        "cbt_calculator_enabled", bool(data["calculator_enabled"]),
                        "Allow calculator in CBT exams by default"
                    )

                # On-the-go tests allowed
                if "on_the_go_enabled" in data:
                    _set_or_create_permission(
                        "cbt_on_the_go_enabled", bool(data["on_the_go_enabled"]),
                        "Allow On-The-Go tests in CBT system"
                    )

                # Default feedback mode
                feedback_mode = data.get("default_feedback_mode", "detailed")
                _set_or_create_permission(
                    "cbt_detailed_feedback", feedback_mode == "detailed",
                    "Show detailed feedback after CBT exam submission"
                )
                _set_or_create_permission(
                    "cbt_no_feedback", feedback_mode == "none",
                    "Disable feedback after CBT exam submission"
                )

                # Timer warning threshold (stored as permission description for now)
                timer_warning = data.get("timer_warning_minutes", 5)
                _set_or_create_permission(
                    "cbt_timer_warning", True,
                    f"Timer warning threshold: {int(timer_warning)} minutes"
                )

                db.session.commit()

                return jsonify({
                    "success": True,
                    "message": "CBT configuration saved successfully"
                }), 200

            # GET: return current CBT config
            perms = {}
            for p in Permission.query.all():
                perms[p.permission_name] = p

            timer_warning_value = 5
            timer_perm = perms.get("cbt_timer_warning")
            if timer_perm and timer_perm.permission_description:
                import re
                match = re.search(r"(\d+)", timer_perm.permission_description)
                if match:
                    timer_warning_value = int(match.group(1))

            feedback_mode = "detailed"
            if perms.get("cbt_no_feedback") and perms["cbt_no_feedback"].is_active:
                feedback_mode = "none"
            elif perms.get("cbt_detailed_feedback") and not perms["cbt_detailed_feedback"].is_active:
                feedback_mode = "summary"

            config = {
                "calculator_enabled": perms.get("cbt_calculator_enabled") and perms["cbt_calculator_enabled"].is_active,
                "on_the_go_enabled": perms.get("cbt_on_the_go_enabled") and perms["cbt_on_the_go_enabled"].is_active,
                "default_feedback_mode": feedback_mode,
                "timer_warning_minutes": timer_warning_value,
            }

            return jsonify({
                "success": True,
                "config": config
            }), 200

        except Exception as e:
            db.session.rollback()
            return jsonify({
                "success": False,
                "message": f"Error managing CBT config: {str(e)}"
            }), 500

    def _set_or_create_permission(name, is_active, description=""):
        """Helper to set or create a permission by name."""
        perm = Permission.query.filter_by(permission_name=name).first()
        if perm:
            perm.is_active = bool(is_active)
            if description:
                perm.permission_description = description
        else:
            perm = Permission(
                permission_name=name,
                is_active=bool(is_active),
                permission_description=description,
                created_for="system",
            )
            db.session.add(perm)
        return perm

    @app.route("/admin/exam/<exam_id>/toggle-active", methods=["POST"])
    @admin_required
    def toggle_exam_active(exam_id):
        """Toggle exam active status"""
        try:
            # print(f"Toggle active called for exam: {exam_id}")
            exam = Exam.query.get(exam_id)
            if not exam:
                # print(f"Exam not found: {exam_id}")
                return jsonify({"success": False, "message": "Exam not found"}), 404

            # Toggle the is_active status
            old_status = exam.is_active
            exam.is_active = not exam.is_active
            db.session.commit()

            status = "activated" if exam.is_active else "deactivated"

            return jsonify({
                "success": True,
                "message": f"Exam {status} successfully",
                "is_active": exam.is_active
            }), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error toggling exam status: {str(e)}")
            import traceback
            traceback.print_exc()
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/exam/<exam_id>/finish", methods=["POST"])
    @admin_required
    def finish_exam(exam_id):
        """Mark exam as finished"""
        try:
            exam = Exam.query.get(exam_id)
            if not exam:
                return jsonify({"success": False, "message": "Exam not found"}), 404

            # Mark exam as finished
            exam.is_finished = True
            exam.is_active = False  # Also deactivate when finished
            db.session.commit()

            return jsonify({
                "success": True,
                "message": "Exam marked as finished successfully",
                "is_finished": exam.is_finished
            }), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error finishing exam: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/exam/<exam_id>/unfinish", methods=["POST"])
    @admin_required
    def unfinish_exam(exam_id):
        """Unfinish an exam (make it active again)"""
        try:
            exam = Exam.query.get(exam_id)
            if not exam:
                return jsonify({"success": False, "message": "Exam not found"}), 404

            # Unfinish the exam
            exam.is_finished = False
            exam.is_active = True  # Make it active again
            db.session.commit()

            return jsonify({
                "success": True,
                "message": "Exam unfinished successfully and is now active",
                "is_finished": exam.is_finished,
                "is_active": exam.is_active
            }), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error unfinishing exam: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    # ===============================
    # QUESTION MANAGEMENT
    # ===============================

    @app.route("/admin/questions", methods=["GET"])
    @admin_required
    def questions_management():
        """Question management page"""
        current_user = User.query.get(session["user_id"])
        subjects = Subject.query.all()
        class_rooms = ClassRoom.query.all()
        school_terms = SchoolTerm.query.all()

        return render_template(
            "admin/questions.html",
            current_user=current_user,
            subjects=subjects,
            class_rooms=class_rooms,
            school_terms=school_terms
        )

    @app.route("/admin/questions/list", methods=["GET"])
    @admin_required
    def list_questions():
        """List questions by subject, class, and term"""
        subject_id = request.args.get('subject_id')
        class_id = request.args.get('class_id')
        term_id = request.args.get('term_id')

        if not all([subject_id, class_id, term_id]):
            return jsonify({"success": False, "message": "Missing required filters"}), 400

        try:
            questions = Question.query.filter_by(
                subject_id=subject_id,
                class_room_id=class_id,
                term_id=term_id
            ).all()

            questions_data = []
            for question in questions:
                questions_data.append({
                    'id': question.id,
                    'question_text': question.question_text,
                    'question_type': question.question_type,
                    'options': [{
                        'id': opt.id,
                        'text': opt.text,
                        'is_correct': opt.is_correct
                    } for opt in question.options]
                })

            return jsonify({
                "success": True,
                "questions": questions_data,
                "count": len(questions_data)
            }), 200

        except Exception as e:
            # print(f"Error listing questions: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/questions/<question_id>", methods=["PUT"])
    @admin_required
    def update_question(question_id):
        """Update a question"""
        try:
            question = Question.query.get(question_id)
            if not question:
                return jsonify({"success": False, "message": "Question not found"}), 404

            data = request.get_json()

            # Update question text
            if 'question_text' in data:
                question.question_text = data['question_text']

            # Update options
            if 'options' in data:
                from models.question import Option
                for option_data in data['options']:
                    option = Option.query.get(option_data['id'])
                    if option:
                        option.text = option_data['text']
                        option.is_correct = option_data['is_correct']

            db.session.commit()

            return jsonify({
                "success": True,
                "message": "Question updated successfully"
            }), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error updating question: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/questions/<question_id>", methods=["DELETE"])
    @admin_required
    def delete_question(question_id):
        """Delete a question"""
        try:
            question = Question.query.get(question_id)
            if not question:
                return jsonify({"success": False, "message": "Question not found"}), 404

            db.session.delete(question)
            db.session.commit()

            return jsonify({
                "success": True,
                "message": "Question deleted successfully"
            }), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error deleting question: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/admin/questions/delete-all", methods=["POST"])
    @admin_required
    def delete_all_questions():
        """Delete all questions for a subject, class, and term"""
        try:
            data = request.get_json()
            subject_id = data.get('subjectId')
            class_id = data.get('classId')
            term_id = data.get('termId')

            if not all([subject_id, class_id, term_id]):
                return jsonify({"success": False, "message": "Missing required filters"}), 400

            questions = Question.query.filter_by(
                subject_id=subject_id,
                class_room_id=class_id,
                term_id=term_id
            ).all()

            count = len(questions)

            for question in questions:
                db.session.delete(question)

            db.session.commit()

            return jsonify({
                "success": True,
                "message": f"{count} questions deleted successfully",
                "deleted_count": count
            }), 200

        except Exception as e:
            db.session.rollback()
            # print(f"Error deleting all questions: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    # ── Trait Definitions CRUD ──────────────────────────────────
    @app.route("/admin/traits")
    @admin_required
    def admin_traits_page():
        current_user = User.query.get(session["user_id"])
        school = School.query.first()
        traits = TraitDefinition.query.filter_by(school_id=school.school_id).order_by(TraitDefinition.sort_order).all()
        return render_template("admin/traits.html", current_user=current_user, school=school, traits=traits)

    @app.route("/admin/api/traits", methods=["GET"])
    @admin_required
    def get_traits():
        school = School.query.first()
        traits = TraitDefinition.query.filter_by(school_id=school.school_id).order_by(TraitDefinition.sort_order).all()
        return jsonify({"success": True, "traits": [t.to_dict() for t in traits]})

    @app.route("/admin/api/traits", methods=["POST"])
    @admin_required
    def create_trait():
        data = request.get_json()
        school = School.query.first()
        trait = TraitDefinition(
            school_id=school.school_id,
            name=data["name"],
            max_score=data.get("max_score", 5.0),
            sort_order=data.get("sort_order", 0),
        )
        db.session.add(trait)
        db.session.commit()
        return jsonify({"success": True, "trait": trait.to_dict()})

    @app.route("/admin/api/traits/<trait_id>", methods=["PUT"])
    @admin_required
    def update_trait(trait_id):
        trait = TraitDefinition.query.get_or_404(trait_id)
        data = request.get_json()
        trait.name = data.get("name", trait.name)
        trait.max_score = data.get("max_score", trait.max_score)
        trait.sort_order = data.get("sort_order", trait.sort_order)
        trait.is_active = data.get("is_active", trait.is_active)
        db.session.commit()
        return jsonify({"success": True, "trait": trait.to_dict()})

    @app.route("/admin/api/traits/<trait_id>", methods=["DELETE"])
    @admin_required
    def delete_trait(trait_id):
        trait = TraitDefinition.query.get_or_404(trait_id)
        StudentTrait.query.filter_by(trait_id=trait_id).delete()
        db.session.delete(trait)
        db.session.commit()
        return jsonify({"success": True})

    # ── Student Traits (admin input) ────────────────────────────
    @app.route("/admin/api/students/<student_id>/traits", methods=["GET", "POST"])
    @admin_required
    def admin_student_traits(student_id):
        term_id = request.args.get("term_id")
        if request.method == "GET":
            traits = StudentTrait.query.filter_by(student_id=student_id, term_id=term_id).all()
            return jsonify({"success": True, "traits": [t.to_dict() for t in traits]})
        else:
            data = request.get_json()
            scores = data.get("scores", {})
            existing = StudentTrait.query.filter_by(student_id=student_id, term_id=term_id).all()
            existing_map = {st.trait_id: st for st in existing}
            for trait_id, score in scores.items():
                if trait_id in existing_map:
                    existing_map[trait_id].score = score
                else:
                    st = StudentTrait(student_id=student_id, term_id=term_id, trait_id=trait_id, score=score)
                    db.session.add(st)
            db.session.commit()
            updated = StudentTrait.query.filter_by(student_id=student_id, term_id=term_id).all()
            return jsonify({"success": True, "traits": [t.to_dict() for t in updated]})
