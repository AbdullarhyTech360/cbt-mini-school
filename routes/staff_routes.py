from flask import Response, jsonify, render_template, request, session, redirect, stream_with_context, url_for, flash
from datetime import datetime
from sqlalchemy import func

from sqlalchemy.sql.functions import current_user
from models import db, TraitDefinition, StudentTrait, Grade
from models.student import Student
from models.teacher import Teacher
from models.user import User
from models.subject import Subject
from models.class_room import ClassRoom
from models.school_term import SchoolTerm
from models.school import School
from routes.dashboard import staff_required


def staff_routes(app):
    # print("=" * 80)
    # print("REGISTERING STAFF ROUTES")
    # print("=" * 80)
    
    @app.route("/staff/test_icons")
    def test_icons():
        return render_template("staff/test_icons.html")

    # route for

    # Input scores
    @app.route("/staff/input_scores/<user_id>", methods=["GET", "POST"])
    @staff_required
    def input_scores(user_id):
        # Verify the logged-in user matches the user_id in the URL
        if session.get("user_id") != user_id:
            flash("Access denied. You can only access your own pages.", "error")
            return redirect(url_for("login"))

        current_user = User.query.get(user_id)
        if not current_user:
            flash("User not found.", "error")
            return redirect(url_for("login"))

        if request.method == "POST":
            try:
                # Get form data
                data = request.get_json()
                if not data:
                    return (
                        jsonify({"success": False, "message": "No data provided"}),
                        400,
                    )

                # Validate required fields
                subject_id = data.get("subject_id")
                class_id = data.get("class_id")
                term_id = data.get("term_id")
                academic_session = data.get("academic_session")
                scores_data = data.get("scores", [])
                action = data.get("action")  # New field to determine the action

                # Handle clear scores action
                if action == "clear_scores":
                    return clear_scores_for_class_subject_term(
                        user_id, subject_id, class_id, term_id
                    )

                # Handle save scores action (existing functionality)
                if not subject_id or not class_id or not scores_data:
                    return (
                        jsonify(
                            {"success": False, "message": "Missing required fields: subject, class, or scores"}
                        ),
                        400,
                    )

                # Auto-detect current term if not provided
                if not term_id or not academic_session:
                    from models.school_term import SchoolTerm
                    current_term = SchoolTerm.query.filter_by(is_current=True).first()
                    if not current_term:
                        return (
                            jsonify(
                                {"success": False, "message": "No active term set. Please set a current term in the admin panel before saving scores."}
                            ),
                            400,
                        )
                    term_id = term_id or current_term.term_id
                    academic_session = academic_session or current_term.academic_session

                # Verify teacher is assigned to this subject-class combination
                from models.associations import teacher_subject

                assignment = (
                    db.session.query(teacher_subject)
                    .filter_by(
                        teacher_id=user_id,
                        subject_id=subject_id,
                        class_room_id=class_id,
                    )
                    .first()
                )

                if not assignment:
                    return (
                        jsonify(
                            {
                                "success": False,
                                "message": "Unauthorized: You are not assigned to this subject-class combination",
                            }
                        ),
                        403,
                    )

                # Validate subject, class, and term exist
                from models.grade import Grade
                from models.school_term import SchoolTerm
                from models.grade_scale import GradeScale

                subject = Subject.query.get(subject_id)
                class_room = ClassRoom.query.get(class_id)
                term = SchoolTerm.query.get(term_id)

                if not subject or not class_room or not term:
                    return (
                        jsonify(
                            {
                                "success": False,
                                "message": "Invalid subject, class, or term",
                            }
                        ),
                        400,
                    )

                # Get assessment types for validation
                from models.assessment_type import AssessmentType
                from models.school import School
                
                school = School.query.first()
                if not school:
                    return (
                        jsonify(
                            {"success": False, "message": "School not found"}
                        ),
                        400,
                    )

                # Resolve grade scale for this class (same fallback chain as report generation)
                grade_scale = None
                if class_room.section_id:
                    section_scales = GradeScale.query.filter(
                        GradeScale.school_id == school.school_id,
                        GradeScale.is_active == True,
                        GradeScale.sections.any(section_id=class_room.section_id)
                    ).order_by(GradeScale.is_default.desc(), GradeScale.created_at.desc()).all()
                    if section_scales:
                        grade_scale = section_scales[0]
                if not grade_scale:
                    grade_scale = GradeScale.query.filter_by(
                        school_id=school.school_id, is_default=True, is_active=True).first()
                if not grade_scale:
                    grade_scale = GradeScale.query.filter_by(
                        school_id=school.school_id, is_active=True).first()
                
                assessment_types_list = AssessmentType.query.filter_by(
                    school_id=school.school_id,
                    is_active=True
                ).all()
                
                # Create a mapping of assessment codes to assessment types
                assessment_map = {at.code: at for at in assessment_types_list}
                
                # Process and validate scores
                saved_count = 0
                errors = []

                for score_entry in scores_data:
                    student_id = score_entry.get("student_id")
                    scores = score_entry.get("scores", {})

                    # Skip if no scores provided
                    if not scores:
                        continue

                    # Verify student exists and belongs to this class
                    student = User.query.filter_by(
                        id=student_id,
                        class_room_id=class_id,
                        role="student",
                        is_active=True,
                    ).first()

                    if not student:
                        errors.append(f"Invalid student ID: {student_id}")
                        continue

                    # Process each assessment score
                    for assessment_code, score_value in scores.items():
                        if score_value is None:
                            continue
                            
                        # Get assessment type details
                        assessment_type = assessment_map.get(assessment_code)
                        if not assessment_type:
                            errors.append(
                                f"Invalid assessment type: {assessment_code}"
                            )
                            continue
                        
                        # Validate score range
                        if score_value < 0 or score_value > assessment_type.max_score:
                            errors.append(
                                f"{assessment_type.name} score for {student.full_name()} must be between 0 and {assessment_type.max_score}"
                            )
                            continue

                        # Check if grade already exists for this combination
                        existing_grade = Grade.query.filter_by(
                            student_id=student_id,
                            subject_id=subject_id,
                            class_room_id=class_id,
                            term_id=term_id,
                            assessment_name=assessment_type.name,
                        ).first()

                        if existing_grade:
                            # Update existing grade
                            existing_grade.score = score_value
                            existing_grade.max_score = assessment_type.max_score
                            existing_grade.teacher_id = user_id
                            existing_grade.is_published = True
                            existing_grade.calculate_percentage()
                            existing_grade.assign_grade_letter(grade_scale)
                        else:
                            # Create new grade
                            new_grade = Grade()
                            new_grade.student_id = student_id
                            new_grade.subject_id = subject_id
                            new_grade.class_room_id = class_id
                            new_grade.teacher_id = user_id
                            new_grade.term_id = term_id
                            new_grade.assessment_type = assessment_type.code
                            new_grade.assessment_name = assessment_type.name
                            new_grade.max_score = assessment_type.max_score
                            new_grade.score = score_value
                            new_grade.academic_session = academic_session
                            new_grade.is_published = True
                            new_grade.calculate_percentage()
                            new_grade.assign_grade_letter(grade_scale)
                            db.session.add(new_grade)

                        saved_count += 1

                # Delete Grade records for assessments that were cleared (sent as null)
                deleted_count = 0
                for score_entry in scores_data:
                    student_id = score_entry.get("student_id")
                    scores = score_entry.get("scores", {})
                    for assessment_code, score_value in scores.items():
                        if score_value is None:
                            deleted = Grade.query.filter_by(
                                student_id=student_id,
                                subject_id=subject_id,
                                class_room_id=class_id,
                                term_id=term_id,
                                assessment_type=assessment_code,
                            ).delete()
                            deleted_count += deleted

                # Commit all changes
                db.session.commit()

                response_message = f"Successfully saved {saved_count} scores"
                if deleted_count > 0:
                    response_message += f", cleared {deleted_count} score(s)"
                if errors:
                    response_message += f". {len(errors)} errors occurred."

                return (
                    jsonify(
                        {
                            "success": True,
                            "message": response_message,
                            "saved_count": saved_count,
                            "errors": errors if errors else None,
                        }
                    ),
                    200,
                )

            except Exception as e:
                db.session.rollback()
                # print(f"Error saving scores: {str(e)}")
                return (
                    jsonify(
                        {"success": False, "message": f"Error saving scores: {str(e)}"}
                    ),
                    500,
                )

        # GET request - Display the form
        # Get classes assigned to this teacher
        assigned_classes = []
        if current_user:
            from models.associations import teacher_classroom

            assigned_classes = (
                db.session.query(ClassRoom)
                .join(
                    teacher_classroom,
                    ClassRoom.class_room_id == teacher_classroom.c.classroom_id,
                )
                .filter(teacher_classroom.c.teacher_id == user_id)
                .all()
            )

        # Get subject_id and class_id from query parameters
        subject_id = request.args.get("subject_id")
        class_id = request.args.get("class_id")

        # If we have both subject_id and class_id, verify the teacher is assigned to this combination
        selected_subject = None
        selected_class = None
        students = []
        current_term = None
        existing_grades = {}
        cbt_grades = {}
        assessment_types = []

        if subject_id and class_id and current_user:
            from models.associations import teacher_subject
            from models.school_term import SchoolTerm
            from models.grade import Grade

            # Check if teacher is assigned to this subject-class combination
            assignment = (
                db.session.query(teacher_subject)
                .filter_by(
                    teacher_id=user_id,
                    subject_id=subject_id,
                    class_room_id=class_id,
                )
                .first()
            )

            if assignment:
                selected_subject = Subject.query.get(subject_id)
                selected_class = ClassRoom.query.get(class_id)

                # Get current term
                current_term = SchoolTerm.query.filter_by(is_current=True).first()
                
                # Get assessment types for this school
                from models.assessment_type import AssessmentType
                from models.school import School
                from models.grade_scale import GradeScale
                school = School.query.first()
                if school:
                    assessment_types = AssessmentType.query.filter_by(
                        school_id=school.school_id,
                        is_active=True
                    ).order_by(AssessmentType.order).all()

                # Resolve grade scale for this class (same fallback chain as report generation)
                grade_scale = None
                if school:
                    if selected_class and selected_class.section_id:
                        section_scales = GradeScale.query.filter(
                            GradeScale.school_id == school.school_id,
                            GradeScale.is_active == True,
                            GradeScale.sections.any(section_id=selected_class.section_id)
                        ).order_by(GradeScale.is_default.desc(), GradeScale.created_at.desc()).all()
                        if section_scales:
                            grade_scale = section_scales[0]
                    if not grade_scale:
                        grade_scale = GradeScale.query.filter_by(
                            school_id=school.school_id, is_default=True, is_active=True).first()
                    if not grade_scale:
                        grade_scale = GradeScale.query.filter_by(
                            school_id=school.school_id, is_active=True).first()

                # Get students in this class (ordered by name)
                students = (
                    User.query.filter_by(
                        class_room_id=class_id, role="student", is_active=True
                    )
                    .order_by(User.first_name, User.last_name)
                    .all()
                )

                # Get existing grades for these students if term is selected
                if current_term:
                    # Auto-sync CBT scores if not already synced
                    from utils.grade_sync import sync_all_exam_records
                    from models.exam_record import ExamRecord
                    
                    # Check for ALL exam records for this combination
                    all_exam_records = ExamRecord.query.filter_by(
                        subject_id=subject_id,
                        class_room_id=class_id,
                        school_term_id=current_term.term_id
                    ).all()
                    
                    # print(f"Found {len(all_exam_records)} exam records for subject/class/term")
                    
                    if all_exam_records:
                        # Check if there are unsynced exam records
                        unsynced_records = ExamRecord.query.filter_by(
                            subject_id=subject_id,
                            class_room_id=class_id,
                            school_term_id=current_term.term_id
                        ).filter(
                            ~ExamRecord.id.in_(
                                db.session.query(Grade.exam_record_id).filter(
                                    Grade.exam_record_id.isnot(None)
                                )
                            )
                        ).all()
                        
                        # print(f"Found {len(unsynced_records)} unsynced exam records")
                        
                        if unsynced_records:
                            # print(f"Auto-syncing {len(unsynced_records)} CBT exam records...")
                            try:
                                sync_result = sync_all_exam_records(subject_id, class_id, current_term.term_id)
                                # print(f"Sync result: {sync_result}")
                            except Exception as e:
                                # print(f"Error during auto-sync: {str(e)}")
                                import traceback
                                traceback.print_exc()
                    
                    grades = Grade.query.filter_by(
                        subject_id=subject_id,
                        class_room_id=class_id,
                        term_id=current_term.term_id,
                    ).all()

                    # Organize existing grades by student and assessment type
                    for grade in grades:
                        key = f"{grade.student_id}_{grade.assessment_name}"
                        existing_grades[key] = grade.score
                        # Mark if it's from CBT (but still editable)
                        if grade.is_from_cbt:
                            cbt_grades[key] = True

                    # Correct stale grade_letter values using the resolved class scale
                    if grade_scale:
                        corrected = False
                        for grade in grades:
                            if grade.score is not None and grade.percentage is not None:
                                grade.assign_grade_letter(grade_scale)
                                corrected = True
                        if corrected:
                            try:
                                db.session.commit()
                            except Exception:
                                db.session.rollback()
            else:
                # Teacher is not assigned to this combination
                flash(
                    "Access denied: You are not assigned to this subject-class combination",
                    "error",
                )
                return redirect(url_for("staff_dashboard", user_id=user_id))

        return render_template(
            "staff/input_scores.html",
            current_user=current_user,
            assigned_classes=assigned_classes,
            selected_subject=selected_subject,
            selected_class=selected_class,
            students=students,
            current_term=current_term,
            existing_grades=existing_grades,
            cbt_grades=cbt_grades,
            assessment_types=assessment_types,
            grade_scale=grade_scale,
        )

    # New function to handle clearing scores
    def clear_scores_for_class_subject_term(teacher_id, subject_id, class_id, term_id):
        try:
            # Verify session user matches the teacher_id
            if session.get("user_id") != teacher_id:
                return (
                    jsonify({"success": False, "message": "Access denied"}),
                    403,
                )

            # Verify teacher is assigned to this subject-class combination
            from models.associations import teacher_subject
            from models.grade import Grade
            from models.school_term import SchoolTerm

            assignment = (
                db.session.query(teacher_subject)
                .filter_by(
                    teacher_id=teacher_id,
                    subject_id=subject_id,
                    class_room_id=class_id,
                )
                .first()
            )

            if not assignment:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": "Unauthorized: You are not assigned to this subject-class combination",
                        }
                    ),
                    403,
                )

            # Validate subject, class, and term exist
            subject = Subject.query.get(subject_id)
            class_room = ClassRoom.query.get(class_id)
            term = SchoolTerm.query.get(term_id)

            if not subject or not class_room or not term:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": "Invalid subject, class, or term",
                        }
                    ),
                    400,
                )

            # Delete all grades for this subject-class-term combination
            deleted_count = Grade.query.filter_by(
                subject_id=subject_id, class_room_id=class_id, term_id=term_id
            ).delete()

            # Commit the changes
            db.session.commit()

            return (
                jsonify(
                    {
                        "success": True,
                        "message": f"Successfully cleared {deleted_count} scores",
                        "deleted_count": deleted_count,
                    }
                ),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error clearing scores: {str(e)}")
            return (
                jsonify(
                    {"success": False, "message": f"Error clearing scores: {str(e)}"}
                ),
                500,
            )

    # Test endpoint
    # print("Registering test_moderation route...")
    @app.route("/staff/test_moderation/<user_id>", methods=["GET"])
    @staff_required
    def test_moderation(user_id):
        # print(f"Test moderation called for user: {user_id}")
        return jsonify({"success": True, "message": "Test route works!"})
    # print("✓ test_moderation route registered")
    
    # Moderate scores endpoint
    # print("Registering moderate_scores route...")
    @app.route("/staff/moderate_scores/<user_id>", methods=["POST"])
    @staff_required
    def moderate_scores(user_id):
        """Apply score moderation (bonus points) to students based on criteria"""
        # print(f"MODERATE_SCORES CALLED for user: {user_id}")
        try:
            # print(f"Session user_id: {session.get('user_id')}")
            # print(f"URL user_id: {user_id}")
            
            # Verify the logged-in user matches the user_id in the URL
            if session.get("user_id") != user_id:
                return (
                    jsonify({"success": False, "message": "Access denied"}),
                    403,
                )

            # Get form data
            data = request.get_json()
            # print(f"Received data: {data}")
            
            if not data:
                return (
                    jsonify({"success": False, "message": "No data provided"}),
                    400,
                )

            # Validate required fields
            subject_id = data.get("subject_id")
            class_id = data.get("class_id")
            term_id = data.get("term_id")
            academic_session = data.get("academic_session")
            assessment_code = data.get("assessment_code")
            bonus_value = data.get("bonus_value")
            apply_to = data.get("apply_to")  # 'all', 'range', 'username'
            threshold = data.get("threshold")  # For 'range' option
            student_id = data.get("student_id")  # For 'username' option
            include_cbt = data.get("include_cbt", False)
            require_approval = data.get("require_approval", True)
            reason = data.get("reason", "")

            if not all([subject_id, class_id, term_id, assessment_code, bonus_value, apply_to, reason]):
                return (
                    jsonify(
                        {"success": False, "message": "Missing required fields (including reason)"}
                    ),
                    400,
                )

            # Verify teacher is assigned to this subject-class combination
            from models.associations import teacher_subject

            assignment = (
                db.session.query(teacher_subject)
                .filter_by(
                    teacher_id=user_id,
                    subject_id=subject_id,
                    class_room_id=class_id,
                )
                .first()
            )

            if not assignment:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": "Unauthorized: You are not assigned to this subject-class combination",
                        }
                    ),
                    403,
                )

            # Get assessment type details
            from models.assessment_type import AssessmentType
            from models.school import School
            from models.grade import Grade
            from models.score_moderation import ScoreModeration

            school = School.query.first()
            if not school:
                return (
                    jsonify({"success": False, "message": "School not found"}),
                    400,
                )

            assessment_type = AssessmentType.query.filter_by(
                school_id=school.school_id,
                code=assessment_code,
                is_active=True
            ).first()

            if not assessment_type:
                return (
                    jsonify(
                        {"success": False, "message": "Invalid assessment type"}
                    ),
                    400,
                )

            # First, check if ANY grades exist for this combination
            all_grades = Grade.query.filter_by(
                subject_id=subject_id,
                class_room_id=class_id,
                term_id=term_id
            ).all()
            
            if all_grades:
                print(f"Sample grade assessment_names: {[g.assessment_name for g in all_grades[:5]]}")
            
            # Also check ExamRecord table for CBT scores
            from models.exam_record import ExamRecord
            exam_records = ExamRecord.query.filter_by(
                subject_id=subject_id,
                class_room_id=class_id,
                school_term_id=term_id
            ).all()
            
            # print(f"Total exam records (CBT) for subject/class/term: {len(exam_records)}")
            if exam_records:
                print(f"Sample exam types: {[e.exam_type for e in exam_records[:5]]}")
            
            # Build query based on apply_to criteria
            query = Grade.query.filter_by(
                subject_id=subject_id,
                class_room_id=class_id,
                term_id=term_id,
                assessment_name=assessment_type.name,
            )
            
            # print(f"Looking for assessment_name: '{assessment_type.name}'")

            if apply_to == "username" and student_id:
                query = query.filter_by(student_id=student_id)
                # print(f"Filtering by student_id: {student_id}")
            elif apply_to == "range" and threshold is not None:
                query = query.filter(Grade.score < threshold)
                # print(f"Filtering by score < {threshold}")
            # For 'all', no additional filter needed

            # Get grades to moderate from Grade table
            grades_to_moderate = query.all()
            
            if grades_to_moderate:
                print(f"Sample grades: {[(g.student_id[:8], g.score) for g in grades_to_moderate[:3]]}")
            
            # Also get matching exam records if include_cbt is True
            exam_records_to_moderate = []
            if include_cbt:
                # Try to match exam_type with assessment name
                # Common patterns: "First CA" might be stored as "CA1", "First CA", etc.
                exam_query = ExamRecord.query.filter_by(
                    subject_id=subject_id,
                    class_room_id=class_id,
                    school_term_id=term_id
                )
                
                # Filter by exam_type matching assessment
                all_exam_records = exam_query.all()
                for record in all_exam_records:
                    # Check if exam_type matches assessment name
                    if (assessment_type.name.lower() in record.exam_type.lower() or 
                        record.exam_type.lower() in assessment_type.name.lower()):
                        
                        # Apply criteria filters
                        should_include = False
                        if apply_to == "all":
                            should_include = True
                        elif apply_to == "range" and threshold is not None:
                            should_include = record.raw_score < threshold
                        elif apply_to == "username" and student_id:
                            should_include = record.student_id == student_id
                        
                        if should_include:
                            exam_records_to_moderate.append(record)
                
                # print(f"Found {len(exam_records_to_moderate)} CBT exam records to moderate")

            if not grades_to_moderate and not exam_records_to_moderate:
                # print("WARNING: No grades or exam records found matching criteria")
                
                # Provide helpful message
                if len(all_grades) == 0 and len(exam_records) == 0:
                    message = "No saved grades or CBT scores found. Please enter scores and save them first, or ensure students have completed CBT exams."
                elif len(exam_records) > 0:
                    message = f"Found {len(exam_records)} CBT scores but none match your criteria. Try checking 'Include CBT Scores' or adjusting your threshold. CBT exam types: {', '.join(set(e.exam_type for e in exam_records[:5]))}"
                else:
                    message = f"No grades found for '{assessment_type.name}' matching your criteria. Available assessments: {', '.join(set(g.assessment_name for g in all_grades[:10]))}"
                
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": message,
                        }
                    ),
                    200,
                )

            # If approval is required, create moderation request
            if require_approval:
                moderation_request = ScoreModeration()
                moderation_request.teacher_id = user_id
                moderation_request.subject_id = subject_id
                moderation_request.class_room_id = class_id
                moderation_request.term_id = term_id
                moderation_request.assessment_code = assessment_code
                moderation_request.assessment_name = assessment_type.name
                moderation_request.bonus_value = float(bonus_value)
                moderation_request.apply_to = apply_to
                moderation_request.threshold = float(threshold) if threshold else None
                moderation_request.target_student_id = student_id
                moderation_request.include_cbt = include_cbt
                moderation_request.reason = reason
                moderation_request.status = 'pending'
                moderation_request.affected_count = len(grades_to_moderate)
                moderation_request.academic_session = academic_session
                
                db.session.add(moderation_request)
                db.session.commit()
                
                return (
                    jsonify(
                        {
                            "success": True,
                            "message": f"Moderation request submitted for {len(grades_to_moderate)} score(s). Pending admin approval.",
                            "moderation_id": moderation_request.moderation_id,
                            "affected_count": len(grades_to_moderate),
                            "status": "pending"
                        }
                    ),
                    200,
                )
            
            # If no approval required, apply immediately
            moderated_count = 0
            
            # Moderate Grade table entries
            for grade in grades_to_moderate:
                # Add bonus but don't exceed max_score
                new_score = min(
                    grade.score + float(bonus_value),
                    assessment_type.max_score
                )
                # print(f"Moderating Grade: {grade.student_id[:8]} from {grade.score} to {new_score}")
                grade.score = new_score
                grade.calculate_percentage()
                grade.assign_grade_letter()
                moderated_count += 1
            
            # Moderate ExamRecord entries (CBT scores)
            for record in exam_records_to_moderate:
                new_score = min(
                    record.raw_score + float(bonus_value),
                    record.max_score
                )
                record.raw_score = new_score
                record.score_percentage = (new_score / record.max_score) * 100 if record.max_score > 0 else 0
                # Use configurable grade scale instead of hardcoded boundaries
                from models.grade_scale import GradeScale
                default_scale = GradeScale.query.filter_by(school_id=school.school_id, is_default=True).first()
                if default_scale:
                    record.letter_grade, _ = default_scale.get_grade_for_percentage(record.score_percentage)
                else:
                    record.letter_grade = 'F'
                moderated_count += 1

            # Log the moderation (even without approval)
            moderation_log = ScoreModeration()
            moderation_log.teacher_id = user_id
            moderation_log.subject_id = subject_id
            moderation_log.class_room_id = class_id
            moderation_log.term_id = term_id
            moderation_log.assessment_code = assessment_code
            moderation_log.assessment_name = assessment_type.name
            moderation_log.bonus_value = float(bonus_value)
            moderation_log.apply_to = apply_to
            moderation_log.threshold = float(threshold) if threshold else None
            moderation_log.target_student_id = student_id
            moderation_log.include_cbt = include_cbt
            moderation_log.reason = reason
            moderation_log.status = 'approved'  # Auto-approved
            moderation_log.approved_by = user_id  # Self-approved
            moderation_log.approval_date = datetime.utcnow()
            moderation_log.affected_count = moderated_count
            moderation_log.academic_session = academic_session
            
            db.session.add(moderation_log)
            db.session.commit()

            return (
                jsonify(
                    {
                        "success": True,
                        "message": f"Successfully moderated {moderated_count} score(s)",
                        "moderated_count": moderated_count,
                        "status": "applied"
                    }
                ),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print("=" * 80)
            # print(f"ERROR IN MODERATE_SCORES: {str(e)}")
            # print(f"Error type: {type(e).__name__}")
            import traceback
            traceback.print_exc()
            # print("=" * 80)
            return (
                jsonify(
                    {"success": False, "message": f"Error moderating scores: {str(e)}"}
                ),
                500,
            )
    
    # print("✓ moderate_scores route registered")
    
    # Sync exam records to grades
    @app.route("/staff/sync_cbt_scores/<user_id>", methods=["POST"])
    @staff_required
    def sync_cbt_scores(user_id):
        """Manually sync CBT exam records to grades table"""
        try:
            from utils.grade_sync import sync_all_exam_records
            
            data = request.get_json() or {}
            subject_id = data.get("subject_id")
            class_id = data.get("class_id")
            term_id = data.get("term_id")
            
            result = sync_all_exam_records(subject_id, class_id, term_id)
            
            return jsonify({
                "success": True,
                "message": f"Synced {result['synced']} new grades, updated {result['updated']} existing grades",
                "stats": result
            }), 200
            
        except Exception as e:
            # print(f"Error syncing CBT scores: {str(e)}")
            import traceback
            traceback.print_exc()
            return jsonify({
                "success": False,
                "message": f"Error syncing: {str(e)}"
            }), 500

    @app.route("/staff/subjects/<user_id>")
    @staff_required
    def get_assigned_subjects(user_id):
        # Verify the logged-in user matches the user_id in the URL but allow admin

        # Note: User ID is not necessarily the same as the session ID, since admin can log in as a different user to upload questions
        user = User.query.get(session.get("user_id"))
        is_admin = user and user.role == "admin"
        if session.get("user_id") != user_id and not is_admin:
            return jsonify({"success": False, "message": "Access denied"}), 403

        user = User.query.get(user_id)
        if not user:
            print("User not present")
            return jsonify({"success": False, "message": "User not found"}), 404

        return jsonify({"success": True, "subjects": grab_staff_subjects(user_id)})

    def grab_staff_subjects(user_id):
        # Grab the existing assignment
        from models.associations import teacher_subject

        existing_assignment = db.session.execute(
            db.select(teacher_subject).where(
                teacher_subject.c.teacher_id == user_id,
            )
        ).fetchall()
        # print("Existing Assignment:", len(existing_assignment))
        subject_teacher = []
        for assignment in existing_assignment:
            class_room = ClassRoom.query.get(assignment[2])
            subject = Subject.query.get(assignment[1])

            # Skip if class or subject not found
            if not class_room or not subject:
                # print(class_room, subject)
                continue

            subject_teacher.append(
                {
                    "class_room_name": class_room.class_room_name,
                    "class_room_id": class_room.class_room_id,
                    "subject_id": subject.subject_id,
                    "subject_name": subject.subject_name,
                    "display_name": f"{subject.subject_name.title()} - {class_room.class_room_name.title()}",
                }
            )

        return subject_teacher

    @app.route("/staff/attendance/<user_id>")
    @staff_required
    def attendance(user_id):
        # Verify the logged-in user matches the user_id in the URL
        if session.get("user_id") != user_id:
            flash("Access denied. You can only access your own pages.", "error")
            return redirect(url_for("login"))

        current_user = User.query.get(user_id)
        if not current_user:
            flash("User not found.", "error")
            return redirect(url_for("login"))

        # Get assigned classes for this teacher (from teacher_classroom association)
        from models.associations import teacher_classroom

        assigned_classes = (
            db.session.query(ClassRoom)
            .join(
                teacher_classroom,
                ClassRoom.class_room_id == teacher_classroom.c.classroom_id,
            )
            .filter(teacher_classroom.c.teacher_id == user_id)
            .all()
        )

        # Also include classes where this teacher is form master
        form_master_classes = ClassRoom.query.filter_by(
            form_teacher_id=user_id, is_active=True
        ).all()
        form_master_ids = {c.class_room_id for c in form_master_classes}
        for cls in assigned_classes:
            form_master_ids.discard(cls.class_room_id)
        assigned_classes.extend(form_master_classes)

        # Get current term
        current_term = SchoolTerm.query.filter_by(is_current=True).first()

        current_date = datetime.now().strftime("%B %d, %Y")
        return render_template(
            "staff/attendance.html",
            current_date=current_date,
            current_user=current_user,
            assigned_classes=assigned_classes,
            current_term=current_term,
        )

    @app.route("/staff/attendance/get_students/<user_id>", methods=["POST"])
    @staff_required
    def get_attendance_students(user_id):
        # Verify the logged-in user matches the user_id in the URL
        if session.get("user_id") != user_id:
            return jsonify({"success": False, "message": "Access denied"}), 403

        try:
            data = request.get_json()
            class_id = data.get("class_id")
            attendance_date = data.get("attendance_date")

            if not class_id:
                return (
                    jsonify({"success": False, "message": "Class ID is required"}),
                    400,
                )

            # Verify teacher is assigned to this class (via teacher_classroom or form_teacher_id)
            from models.associations import teacher_classroom
            from models.class_room import ClassRoom

            assignment = (
                db.session.query(teacher_classroom)
                .filter_by(teacher_id=user_id, classroom_id=class_id)
                .first()
            )

            if not assignment:
                # Fallback: check if teacher is form master of this class
                form_master_class = ClassRoom.query.filter_by(
                    form_teacher_id=user_id, class_room_id=class_id
                ).first()
                if not form_master_class:
                    return (
                        jsonify(
                            {
                                "success": False,
                                "message": "You are not assigned to this class",
                            }
                        ),
                        403,
                    )

            # Get students in this class
            students = (
                User.query.filter_by(
                    class_room_id=class_id, role="student", is_active=True
                )
                .order_by(User.first_name, User.last_name)
                .all()
            )

            # Get existing attendance records for this date
            from models.attendance import Attendance
            from models.school_term import SchoolTerm
            from datetime import datetime as dt

            if attendance_date:
                date_obj = dt.strptime(attendance_date, "%Y-%m-%d").date()
            else:
                date_obj = dt.now().date()

            current_term = SchoolTerm.query.filter_by(is_current=True).first()
            existing_attendance = Attendance.query.filter_by(
                class_room_id=class_id, attendance_date=date_obj
            )
            if current_term:
                existing_attendance = existing_attendance.filter_by(term_id=current_term.term_id)
            existing_attendance = existing_attendance.all()

            # Create a dictionary of existing attendance
            attendance_dict = {att.student_id: att for att in existing_attendance}

            # Prepare student data with attendance status
            students_data = []
            for student in students:
                att_record = attendance_dict.get(student.id)
                students_data.append(
                    {
                        "id": student.id,
                        "first_name": student.first_name,
                        "last_name": student.last_name,
                        "full_name": student.full_name(),
                        "email": student.email,
                        "status": att_record.status if att_record else "present",
                        "remarks": att_record.remarks if att_record else "",
                        "attendance_id": (
                            att_record.attendance_id if att_record else None
                        ),
                    }
                )

            return jsonify(
                {
                    "success": True,
                    "students": students_data,
                    "total_students": len(students_data),
                }
            )

        except Exception as e:
            app.logger.error(f"Error getting attendance students: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/staff/attendance/mark/<user_id>", methods=["POST"])
    @staff_required
    def mark_attendance(user_id):
        # Verify the logged-in user matches the user_id in the URL
        if session.get("user_id") != user_id:
            return jsonify({"success": False, "message": "Access denied"}), 403

        try:
            data = request.get_json()
            class_id = data.get("class_id")
            attendance_date = data.get("attendance_date")
            attendance_records = data.get("attendance_records", [])

            if not class_id or not attendance_records:
                return (
                    jsonify({"success": False, "message": "Missing required fields"}),
                    400,
                )

            # Verify teacher is assigned to this class (via teacher_classroom or form_teacher_id)
            from models.associations import teacher_classroom
            from models.class_room import ClassRoom

            assignment = (
                db.session.query(teacher_classroom)
                .filter_by(teacher_id=user_id, classroom_id=class_id)
                .first()
            )

            if not assignment:
                form_master_class = ClassRoom.query.filter_by(
                    form_teacher_id=user_id, class_room_id=class_id
                ).first()
                if not form_master_class:
                    return (
                        jsonify(
                            {
                                "success": False,
                                "message": "You are not assigned to this class",
                            }
                        ),
                        403,
                    )

            # Get current term
            current_term = SchoolTerm.query.filter_by(is_current=True).first()

            # Parse date
            from datetime import datetime as dt
            from models.attendance import Attendance

            if attendance_date:
                date_obj = dt.strptime(attendance_date, "%Y-%m-%d").date()
            else:
                date_obj = dt.now().date()

            # Validate date is within current term
            if current_term and current_term.start_date and current_term.end_date:
                if date_obj < current_term.start_date or date_obj > current_term.end_date:
                    return (
                        jsonify({"success": False, "message": f"Date must be within the current term ({current_term.start_date.strftime('%b %d')} - {current_term.end_date.strftime('%b %d, %Y')})"}),
                        400,
                    )

            # Reject future dates
            from datetime import date as _date
            today = _date.today()
            if date_obj > today:
                return (
                    jsonify({"success": False, "message": "Cannot mark attendance for a future date"}),
                    400,
                )

            saved_count = 0
            updated_count = 0

            # Check for bulk records on this date — daily cannot overwrite bulk
            # Exception: when editing from history, allow converting bulk to daily
            is_editing = data.get("is_editing", False)
            if not is_editing:
                bulk_conflict = Attendance.query.filter_by(
                    class_room_id=class_id, attendance_date=date_obj, source="bulk"
                ).first()
                if bulk_conflict:
                    return (
                        jsonify({"success": False, "message": "This date was bulk-marked. Edit the bulk record instead."}),
                        400,
                    )

            # Validate all students belong to this class
            student_ids = [r.get("student_id") for r in attendance_records if r.get("student_id")]
            valid_students = (
                User.query.filter(
                    User.id.in_(student_ids),
                    User.class_room_id == class_id,
                    User.role == "student",
                    User.is_active == True,
                ).with_entities(User.id).all()
            )
            valid_ids = {s.id for s in valid_students}
            invalid_ids = set(student_ids) - valid_ids
            if invalid_ids:
                return (
                    jsonify({"success": False, "message": f"Some students do not belong to this class: {', '.join(invalid_ids)}"}),
                    400,
                )

            for record in attendance_records:
                student_id = record.get("student_id")
                status = record.get("status", "present")
                remarks = record.get("remarks", "")

                # Check if attendance already exists
                existing = Attendance.query.filter_by(
                    student_id=student_id,
                    class_room_id=class_id,
                    attendance_date=date_obj,
                ).first()

                if existing:
                    # Update existing record
                    existing.status = status
                    existing.source = "daily"
                    existing.remarks = remarks
                    existing.marked_by_id = user_id
                    updated_count += 1
                else:
                    # Create new record
                    new_attendance = Attendance(
                        student_id=student_id,
                        class_room_id=class_id,
                        attendance_date=date_obj,
                        status=status,
                        source="daily",
                        remarks=remarks,
                        marked_by_id=user_id,
                        term_id=current_term.term_id if current_term else None,
                        academic_session=(
                            current_term.academic_session if current_term else None
                        ),
                    )
                    db.session.add(new_attendance)
                    saved_count += 1

            db.session.commit()

            return jsonify(
                {
                    "success": True,
                    "message": f"Attendance marked successfully. {saved_count} new records, {updated_count} updated.",
                    "saved_count": saved_count,
                    "updated_count": updated_count,
                }
            )

        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Error marking attendance: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/staff/attendance/mark-bulk/<user_id>", methods=["POST"])
    @staff_required
    def mark_attendance_bulk(user_id):
        if session.get("user_id") != user_id:
            return jsonify({"success": False, "message": "Access denied"}), 403

        try:
            from datetime import datetime as dt, timedelta
            from models.attendance import Attendance

            data = request.get_json()
            class_id = data.get("class_id")
            start_date_str = data.get("start_date")
            end_date_str = data.get("end_date")
            records = data.get("records", [])

            if not class_id or not start_date_str or not end_date_str or not records:
                return jsonify({"success": False, "message": "Missing required fields"}), 400

            from models.associations import teacher_classroom
            from models.class_room import ClassRoom
            assignment = db.session.query(teacher_classroom).filter_by(
                teacher_id=user_id, classroom_id=class_id
            ).first()
            if not assignment:
                form_master_class = ClassRoom.query.filter_by(
                    form_teacher_id=user_id, class_room_id=class_id
                ).first()
                if not form_master_class:
                    return jsonify({"success": False, "message": "You are not assigned to this class"}), 403

            start_date = dt.strptime(start_date_str, "%Y-%m-%d").date()
            end_date = dt.strptime(end_date_str, "%Y-%m-%d").date()

            if end_date < start_date:
                return jsonify({"success": False, "message": "End date must be after start date"}), 400

            # Validate dates are within current term
            current_term = SchoolTerm.query.filter_by(is_current=True).first()
            if current_term and current_term.start_date and current_term.end_date:
                if start_date < current_term.start_date or end_date > current_term.end_date:
                    return jsonify({"success": False, "message": f"Date range must be within the current term ({current_term.start_date.strftime('%b %d')} - {current_term.end_date.strftime('%b %d, %Y')})"}), 400

            # Reject future end dates
            from datetime import date as _date
            today = _date.today()
            if end_date > today:
                return jsonify({"success": False, "message": "Cannot mark attendance for future dates"}), 400

            # Build the list of school days (exclude weekends, exclude future)
            all_days = []
            current = start_date
            while current <= end_date:
                if current.weekday() < 5 and current <= today:  # Mon-Fri, not future
                    all_days.append(current)
                current += timedelta(days=1)

            total_days = len(all_days)
            if total_days == 0:
                return jsonify({"success": False, "message": "No school days in the selected range"}), 400

            # Check for daily records in the range — bulk cannot overwrite daily
            if all_days:
                daily_conflicts = (
                    db.session.query(Attendance.attendance_date)
                    .filter(
                        Attendance.class_room_id == class_id,
                        Attendance.attendance_date.in_(all_days),
                        Attendance.source == "daily",
                    )
                    .distinct()
                    .all()
                )
                if daily_conflicts:
                    conflict_dates = [r[0].strftime('%b %d') for r in daily_conflicts]
                    return jsonify({"success": False, "message": f"These dates already have daily records: {', '.join(conflict_dates)}. Clear them first or adjust the range."}), 400

            import random as _random
            saved_count = 0
            updated_count = 0

            for rec in records:
                student_id = rec.get("student_id")
                days_present = min(int(rec.get("days_present", 0)), total_days)

                # Shuffle day indices then pick first N as present
                day_indices = list(range(total_days))
                _random.shuffle(day_indices)
                present_set = set(day_indices[:days_present])

                for i, day in enumerate(all_days):
                    existing = Attendance.query.filter_by(
                        student_id=student_id,
                        class_room_id=class_id,
                        attendance_date=day,
                    ).first()

                    # Distribute: randomly assign present/absent across the range
                    status = "present" if i in present_set else "absent"

                    if existing:
                        existing.status = status
                        existing.source = "bulk"
                        existing.marked_by_id = user_id
                        updated_count += 1
                    else:
                        new_rec = Attendance(
                            student_id=student_id,
                            class_room_id=class_id,
                            attendance_date=day,
                            status=status,
                            source="bulk",
                            marked_by_id=user_id,
                            term_id=current_term.term_id if current_term else None,
                            academic_session=current_term.academic_session if current_term else None,
                        )
                        db.session.add(new_rec)
                        saved_count += 1

            db.session.commit()

            return jsonify({
                "success": True,
                "message": f"Bulk attendance saved. {saved_count} new, {updated_count} updated across {total_days} days.",
                "total_days": total_days,
                "saved_count": saved_count,
                "updated_count": updated_count,
            })

        except Exception as e:
            db.session.rollback()
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/staff/attendance/history/<user_id>", methods=["GET"])
    @staff_required
    def get_attendance_history(user_id):
        # Verify the logged-in user matches the user_id in the URL
        if session.get("user_id") != user_id:
            return jsonify({"success": False, "message": "Access denied"}), 403

        try:
            # Get recent attendance records marked by this teacher
            from models.attendance import Attendance
            from sqlalchemy import func, distinct, case

            # Get unique attendance dates for classes taught by this teacher
            from models.associations import teacher_classroom

            assigned_class_ids = [
                str(row[0])
                for row in db.session.query(teacher_classroom.c.classroom_id)
                .filter(teacher_classroom.c.teacher_id == user_id)
                .all()
            ]

            # Also include form-master classes
            form_master_class_ids = [
                str(cls.class_room_id)
                for cls in ClassRoom.query.filter_by(form_teacher_id=user_id, is_active=True).all()
            ]
            for fid in form_master_class_ids:
                if fid not in assigned_class_ids:
                    assigned_class_ids.append(fid)

            if not assigned_class_ids:
                return jsonify({"success": True, "history": []})

            # Optional class filter
            class_id_filter = request.args.get("class_id")
            if class_id_filter:
                if class_id_filter not in assigned_class_ids:
                    return jsonify({"success": True, "history": []})
                assigned_class_ids = [class_id_filter]

            # Get attendance summary (past/today only)
            from datetime import date as _today_date
            today = _today_date.today()
            history_data = []

            # Daily records: grouped by date + class
            daily_query = (
                db.session.query(
                    Attendance.attendance_date,
                    Attendance.class_room_id,
                    ClassRoom.class_room_name,
                    func.count(Attendance.attendance_id).label("total_marked"),
                    func.sum(case((Attendance.status == "present", 1), else_=0)).label("present_count"),
                    func.sum(case((Attendance.status == "absent", 1), else_=0)).label("absent_count"),
                )
                .join(ClassRoom, Attendance.class_room_id == ClassRoom.class_room_id)
                .filter(Attendance.class_room_id.in_(assigned_class_ids))
                .filter(Attendance.attendance_date <= today)
                .filter(Attendance.source == "daily")
                .group_by(Attendance.attendance_date, Attendance.class_room_id, ClassRoom.class_room_name)
                .all()
            )
            for record in daily_query:
                rate = (record.present_count / record.total_marked * 100) if record.total_marked > 0 else 0
                history_data.append({
                    "date": record.attendance_date.strftime("%Y-%m-%d"),
                    "date_formatted": record.attendance_date.strftime("%B %d, %Y"),
                    "end_date": None,
                    "end_date_formatted": None,
                    "class_id": record.class_room_id,
                    "class_name": record.class_room_name,
                    "total_marked": record.total_marked,
                    "present_count": record.present_count,
                    "absent_count": record.absent_count,
                    "attendance_rate": round(rate, 1),
                    "source": "daily",
                })

            # Bulk records: grouped by class only, show date range
            bulk_query = (
                db.session.query(
                    func.min(Attendance.attendance_date).label("start_date"),
                    func.max(Attendance.attendance_date).label("end_date"),
                    Attendance.class_room_id,
                    ClassRoom.class_room_name,
                    func.count(Attendance.attendance_id).label("total_marked"),
                    func.sum(case((Attendance.status == "present", 1), else_=0)).label("present_count"),
                    func.sum(case((Attendance.status == "absent", 1), else_=0)).label("absent_count"),
                )
                .join(ClassRoom, Attendance.class_room_id == ClassRoom.class_room_id)
                .filter(Attendance.class_room_id.in_(assigned_class_ids))
                .filter(Attendance.attendance_date <= today)
                .filter(Attendance.source == "bulk")
                .group_by(Attendance.class_room_id, ClassRoom.class_room_name)
                .all()
            )
            for record in bulk_query:
                rate = (record.present_count / record.total_marked * 100) if record.total_marked > 0 else 0
                history_data.append({
                    "date": record.start_date.strftime("%Y-%m-%d"),
                    "date_formatted": record.start_date.strftime("%b %d"),
                    "end_date": record.end_date.strftime("%Y-%m-%d"),
                    "end_date_formatted": record.end_date.strftime("%b %d, %Y"),
                    "class_id": record.class_room_id,
                    "class_name": record.class_room_name,
                    "total_marked": record.total_marked,
                    "present_count": record.present_count,
                    "absent_count": record.absent_count,
                    "attendance_rate": round(rate, 1),
                    "source": "bulk",
                })

            # Sort by date descending, limit 10
            history_data.sort(key=lambda x: x["date"], reverse=True)
            history_data = history_data[:10]

            return jsonify({"success": True, "history": history_data})

        except Exception as e:
            app.logger.error(f"Error getting attendance history: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/staff/attendance/history/<user_id>/<class_id>/<date>", methods=["GET"])
    @staff_required
    def get_attendance_history_detail(user_id, class_id, date):
        if session.get("user_id") != user_id:
            return jsonify({"success": False, "message": "Access denied"}), 403

        try:
            from models.attendance import Attendance
            from datetime import datetime as dt

            date_obj = dt.strptime(date, "%Y-%m-%d").date()

            records = (
                db.session.query(
                    Attendance.student_id,
                    User.first_name,
                    User.last_name,
                    Attendance.status,
                    Attendance.remarks,
                )
                .join(User, Attendance.student_id == User.id)
                .filter(
                    Attendance.class_room_id == class_id,
                    Attendance.attendance_date == date_obj,
                )
                .order_by(User.last_name, User.first_name)
                .all()
            )

            records_data = [
                {
                    "student_id": r.student_id,
                    "student_name": f"{r.first_name} {r.last_name}",
                    "status": r.status,
                    "remarks": r.remarks or "",
                }
                for r in records
            ]

            return jsonify({"success": True, "records": records_data})

        except Exception as e:
            app.logger.error(f"Error getting attendance history detail: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/staff/attendance/delete/<user_id>", methods=["DELETE"])
    @staff_required
    def delete_attendance(user_id):
        if session.get("user_id") != user_id:
            return jsonify({"success": False, "message": "Access denied"}), 403

        try:
            from models.attendance import Attendance

            data = request.get_json()
            class_id = data.get("class_id")
            attendance_date = data.get("attendance_date")
            scope = data.get("scope", "date")  # "date" or "all"

            # For "date" scope, class_id is required
            if scope == "date" and not class_id:
                return jsonify({"success": False, "message": "Class ID is required"}), 400

            # Build list of allowed class IDs for this teacher
            from models.associations import teacher_classroom
            from models.class_room import ClassRoom

            allowed_class_ids = [
                str(row[0])
                for row in db.session.query(teacher_classroom.c.classroom_id)
                .filter(teacher_classroom.c.teacher_id == user_id)
                .all()
            ]
            form_master_ids = [
                str(cls.class_room_id)
                for cls in ClassRoom.query.filter_by(form_teacher_id=user_id, is_active=True).all()
            ]
            for fid in form_master_ids:
                if fid not in allowed_class_ids:
                    allowed_class_ids.append(fid)

            if not allowed_class_ids:
                return jsonify({"success": False, "message": "You are not assigned to any classes"}), 403

            # Determine which classes to delete from
            if class_id:
                if class_id not in allowed_class_ids:
                    return jsonify({"success": False, "message": "You are not assigned to this class"}), 403
                target_class_ids = [class_id]
            else:
                target_class_ids = allowed_class_ids

            query = Attendance.query.filter(
                Attendance.class_room_id.in_(target_class_ids)
            )

            if scope == "date":
                if not attendance_date:
                    return jsonify({"success": False, "message": "Date is required"}), 400
                from datetime import datetime as dt
                date_obj = dt.strptime(attendance_date, "%Y-%m-%d").date()
                query = query.filter_by(attendance_date=date_obj)

            deleted_count = query.delete()
            db.session.commit()

            label = f"for {attendance_date}" if scope == "date" else f"across {len(target_class_ids)} class(es)"
            return jsonify({
                "success": True,
                "message": f"Deleted {deleted_count} attendance record(s) {label}.",
                "deleted_count": deleted_count,
            })

        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Error deleting attendance: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/staff/attendance/get_bulk_students/<class_id>", methods=["GET"])
    @staff_required
    def get_bulk_students(class_id):
        """Return students with aggregated days_present across a date range (for bulk edit)."""
        try:
            user_id = session.get("user_id")
            start_date_str = request.args.get("start_date")
            end_date_str = request.args.get("end_date")
            if not start_date_str or not end_date_str:
                return jsonify({"success": False, "message": "start_date and end_date are required"}), 400

            from datetime import datetime as dt
            start_date = dt.strptime(start_date_str, "%Y-%m-%d").date()
            end_date = dt.strptime(end_date_str, "%Y-%m-%d").date()

            # Permission check
            from models.class_room import ClassRoom
            from models.associations import teacher_classroom
            assignment = db.session.query(teacher_classroom).filter_by(teacher_id=user_id, classroom_id=class_id).first()
            if not assignment:
                form_master_class = ClassRoom.query.filter_by(form_teacher_id=user_id, class_room_id=class_id).first()
                if not form_master_class:
                    return jsonify({"success": False, "message": "You are not assigned to this class"}), 403

            students = (
                User.query.filter_by(class_room_id=class_id, role="student", is_active=True)
                .order_by(User.first_name, User.last_name)
                .all()
            )

            # Aggregate days_present per student across the date range
            from models.attendance import Attendance
            from sqlalchemy import func as safunc
            current_term = SchoolTerm.query.filter_by(is_current=True).first()
            present_counts = (
                db.session.query(
                    Attendance.student_id,
                    safunc.count(Attendance.attendance_id).label("days_present"),
                )
                .filter(
                    Attendance.class_room_id == class_id,
                    Attendance.attendance_date >= start_date,
                    Attendance.attendance_date <= end_date,
                    Attendance.status == "present",
                )
                .group_by(Attendance.student_id)
                .all()
            )
            present_map = {str(row.student_id): row.days_present for row in present_counts}

            students_data = []
            for student in students:
                students_data.append({
                    "id": student.id,
                    "first_name": student.first_name,
                    "last_name": student.last_name,
                    "full_name": student.full_name(),
                    "email": student.email,
                    "days_present": present_map.get(str(student.id), 0),
                })

            return jsonify({"success": True, "students": students_data, "total_students": len(students_data)})

        except Exception as e:
            app.logger.error(f"Error getting bulk students: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/staff/upload_questions/<user_id>")
    @staff_required
    def staff_upload_questions(user_id):
        from models import is_permission_active
        if not is_permission_active("teachers_can_upload_questions"):
            return redirect(url_for("staff_dashboard", denied="Question upload has been disabled by the administrator."))

        # Verify the logged-in user matches the user_id in the URL
        if session.get("user_id") != user_id:
            flash("Access denied. You can only access your own pages.", "error")
            return redirect(url_for("login"))

        current_user = User.query.get(user_id)
        if not current_user:
            flash("User not found.", "error")
            return redirect(url_for("login"))

        # Get subjects assigned to this teacher
        from models.associations import teacher_subject, teacher_classroom

        assigned_subjects = (
            db.session.query(Subject)
            .join(teacher_subject, Subject.subject_id == teacher_subject.c.subject_id)
            .filter(teacher_subject.c.teacher_id == current_user.id)
            .all()
        )

        # Get classes assigned to this teacher
        assigned_classes = (
            db.session.query(ClassRoom)
            .join(
                teacher_classroom,
                ClassRoom.class_room_id == teacher_classroom.c.classroom_id,
            )
            .filter(teacher_classroom.c.teacher_id == current_user.id)
            .all()
        )

        # Get school terms
        school_terms = db.session.query(SchoolTerm).all()
        
        # Find the current term (if any)
        current_term = next((term for term in school_terms if term.is_current), None)
        current_term_id = current_term.term_id if current_term else None
        
        # Get assessment types for the school
        from models.assessment_type import AssessmentType
        from models.school import School
        school = School.query.first()
        assessment_types = []
        if school:
            assessment_types = AssessmentType.query.filter_by(
                school_id=school.school_id,
                is_active=True,
                is_cbt_enabled=True  # Only show CBT-enabled assessments
            ).order_by(AssessmentType.order).all()

        return render_template(
            "staff/upload_questions.html",
            current_user=current_user,
            assigned_subjects=grab_staff_subjects(user_id),
            subjects=assigned_subjects,
            classes=assigned_classes,
            school_terms=school_terms,
            current_term_id=current_term_id,
            assessment_types=assessment_types,
        )

    @app.route("/staff/upload_questions", methods=["POST"])
    @staff_required
    def staff_upload_question():
        from models import is_permission_active
        if not is_permission_active("teachers_can_upload_questions"):
            return jsonify({"success": False, "message": "Question upload is disabled by the administrator"}), 403

        # print("Uploading Questions")
        try:
            # print("Uploading Questions")
            data = request.get_json()
            if not data:
                return (
                    jsonify({"success": False, "message": "No data provided"}),
                    400,
                )

            # Get current staff user from session
            current_user = User.query.get(session["user_id"])
            if not current_user:
                return (
                    jsonify({"success": False, "message": "Unauthorized access"}),
                    403,
                )

            # Extract question data
            question_text = data.get("question_text", "").strip()
            question_type = data.get("question_type", "").strip()
            subject_id = data.get("subject_id")
            class_room_id = data.get("class_room_id")
            term_id = data.get("term_id")
            exam_type_id = data.get("exam_type_id")
            options_data = data.get("options", [])

            # Validation
            if not question_text:
                return (
                    jsonify({"success": False, "message": "Question text is required"}),
                    400,
                )

            if not question_type:
                return (
                    jsonify({"success": False, "message": "Question type is required"}),
                    400,
                )

            if not subject_id:
                return (
                    jsonify({"success": False, "message": "Subject is required"}),
                    400,
                )

            if not class_room_id:
                return (
                    jsonify({"success": False, "message": "Class is required"}),
                    400,
                )

            if not term_id:
                return (
                    jsonify({"success": False, "message": "Term is required"}),
                    400,
                )

            if not exam_type_id:
                return (
                    jsonify({"success": False, "message": "Exam type is required"}),
                    400,
                )

            # Validate subject exists and is assigned to this teacher
            from models.associations import teacher_subject

            subject_assignment = (
                db.session.query(teacher_subject)
                .filter_by(teacher_id=current_user.id, subject_id=subject_id)
                .first()
            )

            if not subject_assignment:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": "You are not assigned to this subject",
                        }
                    ),
                    403,
                )

            subject = Subject.query.get(subject_id)
            if not subject:
                return (
                    jsonify({"success": False, "message": "Invalid subject"}),
                    400,
                )

            # Validate class exists and is assigned to this teacher
            from models.associations import teacher_classroom

            # print(current_user.id, class_room_id)

            class_room = ClassRoom.query.get(class_room_id)
            if not class_room:
                return jsonify({"success": False, "message": "Invalid class"}), 400

            # Validate term exists
            term = SchoolTerm.query.get(term_id)
            if not term:
                return jsonify({"success": False, "message": "Invalid term"}), 400

            # Validate that the subject is offered in the selected class
            if subject not in class_room.subjects:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": f"Subject '{subject.subject_name}' is not offered in class '{class_room.class_room_name}'. Please select a valid subject-class combination.",
                        }
                    ),
                    400,
                )

            # Validate question type
            valid_types = ["mcq", "true_false", "short_answer"]
            if question_type not in valid_types:
                return (
                    jsonify({"success": False, "message": "Invalid question type"}),
                    400,
                )

            # For MCQ and True/False questions, validate options
            if question_type in ["mcq", "true_false"]:
                if not options_data:
                    return (
                        jsonify(
                            {
                                "success": False,
                                "message": "Options are required for this question type",
                            }
                        ),
                        400,
                    )

                # Check that there's at least one correct answer
                correct_options = [
                    opt for opt in options_data if opt.get("is_correct", False)
                ]
                if not correct_options:
                    return (
                        jsonify(
                            {
                                "success": False,
                                "message": "At least one correct option is required",
                            }
                        ),
                        400,
                    )

                # For true_false, ensure exactly 2 options (True and False)
                if question_type == "true_false":
                    if len(options_data) != 2:
                        return (
                            jsonify(
                                {
                                    "success": False,
                                    "message": "True/False questions must have exactly 2 options",
                                }
                            ),
                            400,
                        )

            # Create the question using the current teacher
            from models.question import Question, Option

            new_question = Question()
            new_question.question_text = question_text
            new_question.question_type = question_type
            new_question.subject_id = subject_id
            new_question.teacher_id = current_user.id  # Use the current teacher
            new_question.class_room_id = class_room_id
            new_question.term_id = term_id
            new_question.exam_type_id = exam_type_id

            # For short answer questions, save the correct answer
            if question_type == "short_answer":
                new_question.correct_answer = data.get("correct_answer", "")

            db.session.add(new_question)
            db.session.flush()  # Get the question ID without committing

            # Create options if this is an MCQ or True/False question
            if question_type in ["mcq", "true_false"]:
                for i, option_data in enumerate(options_data):
                    option_text = option_data.get("text", "").strip()
                    is_correct = option_data.get("is_correct", False)

                    if option_text:  # Only create option if text is provided
                        option = Option()
                        option.text = option_text
                        option.is_correct = is_correct
                        option.order = i
                        option.question_id = new_question.id
                        db.session.add(option)

            db.session.commit()

            return (
                jsonify(
                    {
                        "success": True,
                        "message": "Question created successfully",
                        "question_id": new_question.id,
                    }
                ),
                200,
            )

        except Exception as e:
            db.session.rollback()
            # print(f"Error creating question: {str(e)}")
            return (
                jsonify(
                    {
                        "success": False,
                        "message": f"Error creating question: {str(e)}",
                    }
                ),
                500,
            )

    @app.route("/staff/questions_preview")
    @staff_required
    def staff_questions_preview():
        """Endpoint to get questions preview based on selected criteria"""
        # print("Previewing Questions")
        try:
            subject_id = request.args.get("subject_id")
            class_room_id = request.args.get("class_room_id")
            term_id = request.args.get("term_id")
            exam_type_id = request.args.get("exam_type_id")

            # Validate required parameters
            if not all([subject_id, class_room_id, term_id, exam_type_id]):
                return (
                    jsonify(
                        {"success": False, "message": "Missing required parameters"}
                    ),
                    400,
                )

            # Query questions based on criteria
            from models.question import Question

            questions = (
                Question.query.filter_by(
                    subject_id=subject_id,
                    class_room_id=class_room_id,
                    term_id=term_id,
                    exam_type_id=exam_type_id,
                )
                .order_by(Question.created_at)
                .all()
            )

            # Format questions for response
            questions_data = []
            for question in questions:
                question_data = {
                    "id": question.id,
                    "question_text": question.question_text,
                    "question_type": question.question_type,
                    "correct_answer": question.correct_answer,
                    "options": [],
                }

                # Add options for MCQ and True/False questions
                if question.question_type in ["mcq", "true_false"]:
                    # Sort options by order
                    options = sorted(question.options, key=lambda x: x.order)
                    for option in options:
                        question_data["options"].append(
                            {"text": option.text, "is_correct": option.is_correct}
                        )

                questions_data.append(question_data)

            return jsonify({"success": True, "questions": questions_data}), 200

        except Exception as e:
            # print(f"Error fetching questions preview: {str(e)}")
            return (
                jsonify(
                    {"success": False, "message": "Error fetching questions preview"}
                ),
                500,
            )

    @app.route("/staff/download_docx_template")
    @staff_required
    def staff_download_docx_template():
        """Generate and download a DOCX template for bulk question upload"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from io import BytesIO

            # Create a new Document
            doc = Document()

            # Add title
            title = doc.add_paragraph()
            title_run = title.add_run("Question Upload Template")
            title_run.bold = True
            title_run.font.size = Pt(16)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add instructions
            doc.add_heading("Instructions:", level=2)
            instructions = [
                "1. Each question should be separated by a blank line",
                '2. Start each question with "Question:"',
                "3. Specify the question type: MCQ, True/False, or Short Answer",
                "4. For MCQ and True/False: list options with * for correct answer",
                "5. For Short Answer: provide the correct answer",
            ]
            for instruction in instructions:
                doc.add_paragraph(instruction, style="List Bullet")

            # Add format explanation
            doc.add_heading("Format:", level=2)
            format_para = doc.add_paragraph()
            format_para.add_run("Question: ").bold = True
            format_para.add_run("[Your question text]\n\n")
            format_para.add_run("Type: ").bold = True
            format_para.add_run("[MCQ/True/False/Short Answer]\n\n")
            format_para.add_run("Options:\n\n").bold = True
            format_para.add_run("- Option 1\n")
            format_para.add_run("- Option 2\n")
            format_para.add_run(
                "- *Correct Option (use * prefix for correct answer)\n\n"
            )
            format_para.add_run("OR for Short Answer:\n\n").bold = True
            format_para.add_run("Answer: ").bold = True
            format_para.add_run("[Correct answer text]")

            # Add examples
            doc.add_heading("Examples:", level=2)

            # Example 1: MCQ
            doc.add_heading("Example 1 - Multiple Choice Question:", level=3)
            example1 = doc.add_paragraph()
            example1.add_run("Question: ").bold = True
            example1.add_run("What is the capital of France?\n\n")
            example1.add_run("Type: ").bold = True
            example1.add_run("MCQ\n\n")
            example1.add_run("Options:\n\n").bold = True
            example1.add_run("- London\n")
            example1.add_run("- *Paris\n")
            example1.add_run("- Berlin\n")
            example1.add_run("- Madrid")

            doc.add_paragraph()  # Blank line separator

            # Example 2: True/False
            doc.add_heading("Example 2 - True/False Question:", level=3)
            example2 = doc.add_paragraph()
            example2.add_run("Question: ").bold = True
            example2.add_run("The Earth is flat.\n\n")
            example2.add_run("Type: ").bold = True
            example2.add_run("True/False\n\n")
            example2.add_run("Options:\n\n").bold = True
            example2.add_run("- True\n")
            example2.add_run("- *False")

            doc.add_paragraph()  # Blank line separator

            # Example 3: Short Answer
            doc.add_heading("Example 3 - Short Answer Question:", level=3)
            example3 = doc.add_paragraph()
            example3.add_run("Question: ").bold = True
            example3.add_run("What is the chemical symbol for water?\n\n")
            example3.add_run("Type: ").bold = True
            example3.add_run("Short Answer\n\n")
            example3.add_run("Answer: ").bold = True
            example3.add_run("H2O")

            # Save to BytesIO
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            # Send file
            from flask import send_file

            return send_file(
                file_stream,
                as_attachment=True,
                download_name="questions_template.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        except Exception as e:
            # print(f"Error generating DOCX template: {str(e)}")
            return (
                jsonify(
                    {
                        "success": False,
                        "message": f"Error generating template: {str(e)}",
                    }
                ),
                500,
            )

    @app.route("/staff/bulk_upload_questions/<user_id>", methods=["GET", "POST"])
    @staff_required
    def staff_bulk_upload_questions(user_id):
        from models import is_permission_active
        if not is_permission_active("teachers_can_upload_questions"):
            if request.method == "POST":
                return jsonify({"success": False, "message": "Question upload is disabled by the administrator"}), 403
            return redirect(url_for("staff_dashboard", denied="Question upload has been disabled by the administrator."))

        # print("Bulk Upload Questions")
        # Verify the logged-in user matches the user_id in the URL
        if session.get("user_id") != user_id:
            return jsonify({"success": False, "message": "Access denied"}), 403

        if request.method == "POST":
            try:
                from utils.question_parser import parse_questions_file
                from models.question import Question, Option
                from models.associations import teacher_subject, teacher_classroom
                from models.school_term import SchoolTerm
                
                # Get current staff user from session
                current_user = User.query.get(session["user_id"])
                if not current_user:
                    return jsonify({"success": False, "message": "Unauthorized access"}), 403
                
                # Get form data (from FormData)
                subject_id = request.form.get("subject_id")
                class_room_id = request.form.get("class_room_id")
                term_id = request.form.get("term_id")
                exam_type_id = request.form.get("exam_type_id")
                
                # Validate required fields
                if not subject_id or not class_room_id or not term_id or not exam_type_id:
                    return jsonify({
                        "success": False,
                        "message": "Subject, Class, Term, and Exam Type are required"
                    }), 400
                
                # Check if file was uploaded
                if "file" not in request.files:
                    return jsonify({
                        "success": False,
                        "message": "No file uploaded"
                    }), 400
                
                file = request.files["file"]
                if not file or file.filename == '':
                    return jsonify({
                        "success": False,
                        "message": "No file selected"
                    }), 400
                
                # Validate subject-class combination is assigned to this teacher
                assignment = (
                    db.session.query(teacher_subject)
                    .filter_by(
                        teacher_id=current_user.id,
                        subject_id=subject_id,
                        class_room_id=class_room_id
                    )
                    .first()
                )
                if not assignment:
                    return jsonify({
                        "success": False,
                        "message": "You are not assigned to this subject-class combination"
                    }), 403

                subject = Subject.query.get(subject_id)
                if not subject:
                    return jsonify({"success": False, "message": "Invalid subject"}), 400

                class_room = ClassRoom.query.get(class_room_id)
                if not class_room:
                    return jsonify({"success": False, "message": "Invalid class"}), 400

                # Validate term exists
                term = SchoolTerm.query.get(term_id)
                if not term:
                    return jsonify({"success": False, "message": "Invalid term"}), 400
                
                # Parse the file using our universal parser
                questions_data, error = parse_questions_file(file)
                
                if error:
                    return jsonify({"success": False, "message": error}), 400
                
                if not questions_data:
                    return jsonify({
                        "success": False,
                        "message": "No valid questions found in the file"
                    }), 400

                created_questions = []
                errors = []

                # Process each question
                for i, question_data in enumerate(questions_data, 1):
                    try:
                        # Extract question data
                        question_text = question_data.get("question_text", "").strip()
                        question_type = question_data.get("question_type", "").strip().lower()
                        options_data = question_data.get("options", [])
                        correct_answer = question_data.get("correct_answer", "")

                        # Validation
                        if not question_text:
                            errors.append(f"Question {i}: Question text is required")
                            continue

                        if not question_type:
                            errors.append(f"Question {i}: Question type is required")
                            continue

                        # Validate question type
                        valid_types = ["mcq", "true_false", "short_answer"]
                        if question_type not in valid_types:
                            errors.append(f"Question {i}: Invalid question type '{question_type}'. Must be one of: {', '.join(valid_types)}")
                            continue

                        # For MCQ and True/False questions, validate options
                        if question_type in ["mcq", "true_false"]:
                            if not options_data or len(options_data) == 0:
                                errors.append(f"Question {i}: Options are required for {question_type} questions")
                                continue

                            # Check that there's at least one correct answer
                            correct_options = [opt for opt in options_data if opt.get("is_correct", False)]
                            if not correct_options:
                                errors.append(f"Question {i}: At least one correct option is required")
                                continue

                            # For true_false, ensure exactly 2 options
                            if question_type == "true_false" and len(options_data) != 2:
                                errors.append(f"Question {i}: True/False questions must have exactly 2 options")
                                continue

                        # Create the question
                        new_question = Question()
                        new_question.question_text = question_text
                        new_question.question_type = question_type
                        new_question.subject_id = subject_id
                        new_question.teacher_id = current_user.id
                        new_question.class_room_id = class_room_id
                        new_question.term_id = term_id
                        new_question.exam_type_id = exam_type_id
                        
                        # Save math and image data if present
                        new_question.has_math = question_data.get("has_math", False)
                        new_question.question_image = question_data.get("question_image")

                        # For short answer questions, save the correct answer
                        if question_type == "short_answer":
                            new_question.correct_answer = correct_answer

                        db.session.add(new_question)
                        db.session.flush()  # Get the question ID

                        # Create options if this is an MCQ or True/False question
                        if question_type in ["mcq", "true_false"]:
                            for j, option_data in enumerate(options_data):
                                option_text = option_data.get("text", "").strip()
                                is_correct = option_data.get("is_correct", False)

                                if option_text:
                                    option = Option()
                                    option.text = option_text
                                    option.is_correct = is_correct
                                    option.order = j
                                    option.question_id = new_question.id
                                    
                                    # Save math and image data for options
                                    option.has_math = option_data.get("has_math", False)
                                    option.option_image = option_data.get("option_image")
                                    
                                    db.session.add(option)

                        db.session.commit()
                        created_questions.append(new_question.id)

                    except Exception as e:
                        db.session.rollback()
                        errors.append(f"Question {i}: Error creating question - {str(e)}")
                        continue

                # Prepare response
                response_data = {
                    "success": True,
                    "message": f"Successfully created {len(created_questions)} questions",
                    "created_count": len(created_questions),
                    "error_count": len(errors),
                }

                if errors:
                    response_data["message"] += f" with {len(errors)} errors"
                    response_data["errors"] = errors

                return jsonify(response_data), 200

            except Exception as e:
                db.session.rollback()
                # print(f"Error processing bulk upload: {str(e)}")
                import traceback
                traceback.print_exc()
                return jsonify({
                    "success": False,
                    "message": f"Error processing bulk upload: {str(e)}"
                }), 500

        # GET request - render the bulk upload questions page
        current_user = User.query.get(session["user_id"])
        if not current_user:
            return redirect(url_for("login"))

        # Get subjects assigned to this teacher
        from models.associations import teacher_subject

        assigned_subjects = (
            db.session.query(Subject)
            .join(teacher_subject, Subject.subject_id == teacher_subject.c.subject_id)
            .filter(teacher_subject.c.teacher_id == current_user.id)
            .all()
        )

        # Get classes assigned to this teacher
        from models.associations import teacher_classroom

        assigned_classes = (
            db.session.query(ClassRoom)
            .join(
                teacher_classroom,
                ClassRoom.class_room_id == teacher_classroom.c.classroom_id,
            )
            .filter(teacher_classroom.c.teacher_id == current_user.id)
            .all()
        )

        # Get school terms
        from models.school_term import SchoolTerm
        school_terms = SchoolTerm.query.order_by(SchoolTerm.term_name).all()
        
        # Get current term
        current_term = SchoolTerm.query.filter_by(is_current=True).first()
        
        # Get assessment types for the school
        from models.assessment_type import AssessmentType
        from models.school import School
        school = School.query.first()
        assessment_types = []
        if school:
            assessment_types = AssessmentType.query.filter_by(
                school_id=school.school_id,
                is_active=True,
                is_cbt_enabled=True  # Only show CBT-enabled assessments
            ).order_by(AssessmentType.order).all()

        return render_template(
            "staff/bulk_upload_questions.html",
            current_user=current_user,
            assigned_subjects=assigned_subjects,
            assigned_classes=assigned_classes,
            school_terms=school_terms,
            current_term=current_term,
            assessment_types=assessment_types,
        )

    @app.route("/staff/extract_questions", methods=["POST"])
    @staff_required
    def staff_extract_questions():
        """Stream extraction milestones and final parsed questions for review."""
        from models import is_permission_active
        if not is_permission_active("teachers_can_upload_questions"):
            return jsonify({"success": False, "message": "Question upload is disabled by the administrator"}), 403

        try:
            import json
            import os
            import queue
            import tempfile
            import threading
            import time

            # Verify user
            current_user = User.query.get(session["user_id"]) if session.get("user_id") else None
            if not current_user:
                return jsonify({"success": False, "message": "Unauthorized access"}), 403

            # Validate multipart - accept multiple files (up to 4)
            files = request.files.getlist('file')
            if not files or all(not f or f.filename == '' for f in files):
                return jsonify({"success": False, "message": "No file uploaded"}), 400

            # Filter out empty files and limit to 4 files
            files = [f for f in files if f and f.filename != '']
            if len(files) > 4:
                return jsonify({"success": False, "message": "Maximum 4 files allowed at once"}), 400

            if len(files) == 0:
                return jsonify({"success": False, "message": "No file selected"}), 400

            # Get metadata
            subject_id = request.form.get('subject_id')
            class_room_id = request.form.get('class_room_id')
            term_id = request.form.get('term_id')
            exam_type_id = request.form.get('exam_type_id')
            custom_prompt = request.form.get('custom_prompt')

            if not subject_id or not class_room_id or not term_id or not exam_type_id:
                return jsonify({"success": False, "message": "Subject, Class, Term, and Exam Type are required"}), 400

            # Verify teacher assignment
            from models.associations import teacher_subject
            assignment = (
                db.session.query(teacher_subject)
                .filter_by(teacher_id=current_user.id, subject_id=subject_id, class_room_id=class_room_id)
                .first()
            )
            if not assignment:
                return jsonify({"success": False, "message": "You are not assigned to this subject-class combination"}), 403

            # Save all files to temporary files
            temp_files = []
            filenames = []
            for file in files:
                filename = file.filename or ''
                filenames.append(filename)
                suffix = ''
                if '.' in filename:
                    suffix = '.' + filename.rsplit('.', 1)[1].lower()
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                file.save(tmp.name)
                tmp.close()
                temp_files.append(tmp.name)

            from services.gemini_extractor import extract_questions_from_file

            event_queue = queue.Queue()
            result_holder = {"questions": [], "error": None, "model_name": None}
            worker_done = threading.Event()

            def stream_event(payload):
                event_queue.put(payload)

            def extractor_worker():
                try:
                    all_questions = []
                    for idx, (tmp_file, filename) in enumerate(zip(temp_files, filenames)):
                        # Add file index to progress events
                        def wrapped_stream_event(payload):
                            payload_with_file = payload.copy()
                            payload_with_file["file_index"] = idx + 1
                            payload_with_file["total_files"] = len(temp_files)
                            payload_with_file["file_name"] = filename
                            event_queue.put(payload_with_file)

                        questions, error, model_name = extract_questions_from_file(
                            tmp_file,
                            custom_prompt=custom_prompt,
                            progress_callback=wrapped_stream_event,
                        )
                        if error:
                            result_holder["error"] = f"Error processing {filename}: {error}"
                            break
                        if questions:
                            all_questions.extend(questions)
                        result_holder["model_name"] = model_name

                    result_holder["questions"] = all_questions
                except Exception as worker_error:
                    result_holder["error"] = str(worker_error)
                finally:
                    worker_done.set()

            threading.Thread(target=extractor_worker, daemon=True).start()

            def build_event(event, **payload):
                response_payload = {"event": event, **payload}
                return json.dumps(response_payload) + "\n"

            @stream_with_context
            def generate():
                last_emit_at = time.time()
                last_progress = 2
                file_list = ", ".join(filenames) if len(filenames) <= 3 else f"{filenames[0]}, {filenames[1]} and {len(filenames) - 2} more"
                yield build_event(
                    "milestone",
                    progress=2,
                    status="running",
                    stage_id="initializing-model-pipeline",
                    stage_label="Initializing model pipeline",
                    message=f"Upload received. Starting extraction workflow for {len(files)} file(s): {file_list}",
                    file_count=len(files),
                    file_names=filenames,
                )

                try:
                    while not worker_done.is_set() or not event_queue.empty():
                        try:
                            payload = event_queue.get(timeout=1.0)
                            last_emit_at = time.time()
                            if payload.get("progress") is not None:
                                last_progress = payload.get("progress")
                            yield build_event(payload.pop("event", "milestone"), **payload)
                        except queue.Empty:
                            if worker_done.is_set():
                                break
                            if time.time() - last_emit_at >= 2:
                                last_emit_at = time.time()
                                yield build_event(
                                    "heartbeat",
                                    progress=last_progress,
                                    status="waiting",
                                    stage_id="processing-input-text",
                                    stage_label="Processing input text",
                                    message="Extraction is still active. The current model is working on your document.",
                                )

                    while not event_queue.empty():
                        payload = event_queue.get_nowait()
                        yield build_event(payload.pop("event", "milestone"), **payload)

                    if result_holder["error"]:
                        yield build_event(
                            "error",
                            success=False,
                            message=result_holder["error"],
                        )
                        return

                    questions = result_holder["questions"] or []
                    if not questions:
                        yield build_event(
                            "error",
                            success=False,
                            message="No questions extracted",
                        )
                        return

                    yield build_event(
                        "complete",
                        success=True,
                        progress=100,
                        status="completed",
                        stage_id="generating-structured-question-output",
                        stage_label="Generating structured question output",
                        message=f"Extraction completed successfully with {len(questions)} question(s) from {len(files)} file(s).",
                        model_name=result_holder["model_name"],
                        questions=questions,
                        file_count=len(files),
                    )
                finally:
                    # Clean up all temporary files
                    for tmp_file in temp_files:
                        try:
                            os.unlink(tmp_file)
                        except OSError:
                            pass

            response = Response(generate(), mimetype="application/x-ndjson")
            response.headers["Cache-Control"] = "no-cache"
            response.headers["X-Accel-Buffering"] = "no"
            return response

        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"success": False, "message": str(e)}), 500

    @app.route("/staff/ai_chat_generate", methods=["POST"])
    @staff_required
    def staff_ai_chat_generate():
        """AI Chat endpoint for generating questions from uploaded documents or text prompts with system prompt."""
        from models import is_permission_active
        if not is_permission_active("teachers_can_upload_questions"):
            return jsonify({"success": False, "message": "Question upload is disabled by the administrator"}), 403

        try:
            import os
            import tempfile
            import json

            # Verify user
            current_user = User.query.get(session["user_id"]) if session.get("user_id") else None
            if not current_user:
                return jsonify({"success": False, "error": "Unauthorized access"}), 403

            # Get files and instruction
            files = request.files.getlist('files')
            instruction = request.form.get('instruction', '')
            conversation_history_str = request.form.get('conversation_history', '[]')

            # Filter out empty files
            files = [f for f in files if f and f.filename != '']
            
            if not files and not instruction:
                return jsonify({"success": False, "error": "Please provide files or instructions"}), 400

            # System prompt for AI
            system_prompt = """You are an academic assistant helping teachers create assessment questions. 
Generate questions based on the uploaded documents and user instructions.
Format questions according to the system's question structure (question_text, question_type, options, correct_answer, has_math, context).
Support MCQ, True/False, and Short Answer question types.
For MCQ: provide at least 3-4 options with exactly one correct answer.
For True/False: represent as MCQ with exactly 2 options (True and False).
Include context field for tables, diagrams, formulas referenced in questions."""

            # Decode conversation history
            try:
                conversation_history = json.loads(conversation_history_str)
            except Exception:
                conversation_history = []

            # Format conversation history for context
            history_prompt = ""
            if conversation_history:
                history_prompt = "Here is the conversation history so far for reference:\n"
                for msg in conversation_history:
                    # Ignore the initial empty state message
                    if msg.get("content") and "Start a conversation" not in msg.get("content"):
                        role = "User" if msg.get("type") == "user" else "Assistant"
                        history_prompt += f"{role}: {msg.get('content')}\n"
                history_prompt += "\n"

            # Combine system prompt with user instruction and history
            full_prompt = f"{system_prompt}\n\n{history_prompt}User instruction: {instruction}"

            from services.gemini_extractor import extract_questions_from_file, generate_questions_from_prompt

            all_questions = []
            model_name = None
            error = None

            if files:
                # Save files to temporary files
                temp_files = []
                for file in files:
                    filename = file.filename or ''
                    suffix = ''
                    if '.' in filename:
                        suffix = '.' + filename.rsplit('.', 1)[1].lower()
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                    file.save(tmp.name)
                    tmp.close()
                    temp_files.append(tmp.name)

                for tmp_file in temp_files:
                    questions, err, model = extract_questions_from_file(
                        tmp_file,
                        custom_prompt=full_prompt,
                        progress_callback=None  # No streaming for chat interface
                    )
                    if err:
                        error = err
                        break
                    if questions:
                        all_questions.extend(questions)
                    model_name = model

                # Clean up temporary files
                for tmp_file in temp_files:
                    try:
                        os.unlink(tmp_file)
                    except OSError:
                        pass
            else:
                # Text-only generation if no files were uploaded
                questions, err, model = generate_questions_from_prompt(
                    full_prompt,
                    progress_callback=None
                )
                if err:
                    error = err
                if questions:
                    all_questions.extend(questions)
                model_name = model

            if error:
                return jsonify({"success": False, "error": error}), 500

            if not all_questions:
                return jsonify({"success": False, "error": "No questions generated"}), 400

            return jsonify({
                "success": True,
                "message": f"Successfully generated {len(all_questions)} question(s)",
                "questions": all_questions,
                "model_used": model_name
            })

        except Exception as e:
            return jsonify({"success": False, "error": str(e)}), 500


    @app.route("/staff/create_questions_from_json", methods=["POST"])
    @staff_required
    def staff_create_questions_from_json():
        """Create questions from a JSON payload (used after AI extraction review).
        Expects JSON: { subject_id, class_room_id, term_id, exam_type_id, questions: [...] }
        """
        from models import is_permission_active
        if not is_permission_active("teachers_can_upload_questions"):
            return jsonify({"success": False, "message": "Question upload is disabled by the administrator"}), 403

        try:
            data = request.get_json()
            if not data:
                return jsonify({"success": False, "message": "No data provided"}), 400

            current_user = User.query.get(session["user_id"]) if session.get("user_id") else None
            if not current_user:
                return jsonify({"success": False, "message": "Unauthorized access"}), 403

            subject_id = data.get('subject_id')
            class_room_id = data.get('class_room_id')
            term_id = data.get('term_id')
            exam_type_id = data.get('exam_type_id')
            questions = data.get('questions', [])

            if not subject_id or not class_room_id or not term_id or not exam_type_id:
                return jsonify({"success": False, "message": "Subject, Class, Term, and Exam Type are required"}), 400

            from models.associations import teacher_subject
            assignment = (
                db.session.query(teacher_subject)
                .filter_by(teacher_id=current_user.id, subject_id=subject_id, class_room_id=class_room_id)
                .first()
            )
            if not assignment:
                return jsonify({"success": False, "message": "You are not assigned to this subject-class combination"}), 403

            from models.question import Question, Option
            from models.school_term import SchoolTerm

            created_questions = []
            errors = []

            for i, question_data in enumerate(questions, 1):
                try:
                    question_text = question_data.get('question_text', '').strip()
                    question_type = question_data.get('question_type', '').strip().lower()
                    options_data = question_data.get('options', [])
                    correct_answer = question_data.get('correct_answer', '')

                    if not question_text:
                        errors.append(f"Question {i}: Question text is required")
                        continue

                    # Infer type if omitted; default to MCQ for questions with options.
                    if not question_type:
                        if options_data:
                            question_type = 'mcq'
                        elif correct_answer:
                            question_type = 'short_answer'
                        else:
                            errors.append(f"Question {i}: Question type is required")
                            continue

                    if question_type == 'true_false':
                        question_type = 'mcq'

                    valid_types = ["mcq", "short_answer"]
                    if question_type not in valid_types:
                        errors.append(f"Question {i}: Invalid question type '{question_type}'")
                        continue

                    if question_type == "mcq":
                        if not options_data or len(options_data) == 0:
                            errors.append(f"Question {i}: Options are required for MCQ questions")
                            continue
                        correct_options = [opt for opt in options_data if opt.get('is_correct', False)]
                        if not correct_options:
                            errors.append(f"Question {i}: At least one correct option is required")
                            continue

                    new_question = Question()
                    new_question.question_text = question_text
                    new_question.question_type = question_type
                    new_question.subject_id = subject_id
                    new_question.teacher_id = current_user.id
                    new_question.class_room_id = class_room_id
                    new_question.term_id = term_id
                    new_question.exam_type_id = exam_type_id
                    new_question.has_math = question_data.get('has_math', False)
                    new_question.question_image = question_data.get('question_image')

                    if question_type == 'short_answer':
                        new_question.correct_answer = correct_answer

                    db.session.add(new_question)
                    db.session.flush()

                    if question_type == "mcq":
                        for j, option_data in enumerate(options_data):
                            option_text = option_data.get('text', '').strip()
                            is_correct = option_data.get('is_correct', False)
                            if option_text:
                                opt = Option()
                                opt.text = option_text
                                opt.is_correct = is_correct
                                opt.order = j
                                opt.question_id = new_question.id
                                opt.has_math = option_data.get('has_math', False)
                                opt.option_image = option_data.get('option_image')
                                db.session.add(opt)

                    db.session.commit()
                    created_questions.append(new_question.id)

                except Exception as e:
                    db.session.rollback()
                    errors.append(f"Question {i}: Error creating question - {str(e)}")
                    continue

            response_data = {
                'success': True,
                'message': f"Successfully created {len(created_questions)} questions",
                'created_count': len(created_questions),
                'error_count': len(errors)
            }
            if errors:
                response_data['message'] += f" with {len(errors)} errors"
                response_data['errors'] = errors

            return jsonify(response_data), 200

        except Exception as e:
            db.session.rollback()
            import traceback
            traceback.print_exc()
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route("/staff/profile/<user_id>")
    @staff_required
    def staff_profile(user_id):
        # Verify the logged-in user matches the user_id in the URL
        if session.get("user_id") != user_id:
            flash("Access denied. You can only access your own pages.", "error")
            return redirect(url_for("login"))

        current_user = User.query.get(user_id)
        if not current_user:
            flash("User not found.", "error")
            return redirect(url_for("login"))

        return render_template("staff/profile.html", current_user=current_user)

    @app.route("/staff/classes/<user_id>")
    @staff_required
    def staff_classes(user_id):
        # Verify the logged-in user matches the user_id in the URL
        if session.get("user_id") != user_id:
            flash("Access denied. You can only access your own pages.", "error")
            return redirect(url_for("login"))

        current_user = User.query.get(user_id)
        if not current_user:
            flash("User not found.", "error")
            return redirect(url_for("login"))

        return render_template("staff/classes.html", current_user=current_user)

    # ── Staff Trait Input ───────────────────────────────────────
    def _seed_default_traits(school):
        """Create default behavioral and psychomotor trait definitions if none exist."""
        if not school:
            return
        existing = TraitDefinition.query.filter_by(school_id=school.school_id).first()
        if existing:
            return
        default_behavioral = [
            ("Punctuality", 1), ("Discipline", 2), ("Respect", 3),
            ("Honesty", 4), ("Neatness", 5), ("Leadership", 6),
            ("Teamwork", 7), ("Responsibility", 8),
        ]
        default_psychomotor = [
            ("Handwriting", 1), ("Drawing/Illustration", 2), ("Verbal Fluency", 3),
            ("Physical Development", 4), ("Retention", 5), ("Creative Ability", 6),
            ("Practical Skills", 7), ("Handling of Tools/Equipment", 8),
        ]
        for name, order in default_behavioral + default_psychomotor:
            trait = TraitDefinition(
                school_id=school.school_id, name=name,
                max_score=5.0, sort_order=order, is_active=True,
            )
            db.session.add(trait)
        db.session.commit()

    @app.route("/staff/api/traits/class/<class_id>", methods=["GET"])
    @staff_required
    def staff_class_traits(class_id):
        term_id = request.args.get("term_id")
        school = School.query.first()
        trait_defs = TraitDefinition.query.filter_by(school_id=school.school_id, is_active=True).order_by(TraitDefinition.sort_order).all()
        users = User.query.filter_by(class_room_id=class_id, role="student", is_active=True).order_by(User.first_name, User.last_name).all()
        result = []
        for u in users:
            student = u.student
            if not student:
                continue
            existing = []
            if term_id:
                existing = StudentTrait.query.filter_by(student_id=u.id, term_id=term_id).all()
            score_map = {st.trait_id: st.score for st in existing}
            result.append({
                "student_id": u.id,
                "name": f"{u.first_name} {u.last_name}".strip(),
                "scores": {t.id: score_map.get(t.id, None) for t in trait_defs},
            })
        return jsonify({
            "success": True,
            "students": result,
            "traits": [t.to_dict() for t in trait_defs],
        })

    @app.route("/staff/api/traits/save", methods=["POST"])
    @staff_required
    def staff_save_traits():
        data = request.get_json()
        student_id = data.get("student_id")
        term_id = data.get("term_id")
        scores = data.get("scores", {})
        existing = StudentTrait.query.filter_by(student_id=student_id, term_id=term_id).all()
        existing_map = {st.trait_id: st for st in existing}
        for trait_id, score in scores.items():
            if score is None:
                continue
            if trait_id in existing_map:
                existing_map[trait_id].score = score
            else:
                st = StudentTrait(student_id=student_id, term_id=term_id, trait_id=trait_id, score=score)
                db.session.add(st)
        db.session.commit()
        return jsonify({"success": True})

    # ── Staff Psychomotor Domain ────────────────────────────────
    @app.route("/staff/psychomotor/<user_id>")
    @staff_required
    def staff_psychomotor_page(user_id):
        if session.get("user_id") != user_id:
            flash("Access denied.", "error")
            return redirect(url_for("login"))
        current_user = User.query.get(user_id)
        if not current_user:
            flash("User not found.", "error")
            return redirect(url_for("login"))
        teacher = Teacher.query.filter_by(user_id=user_id).first()
        school = School.query.first()
        _seed_default_traits(school)
        terms = SchoolTerm.query.filter_by(school_id=school.school_id).all() if school else []
        form_master_classes = ClassRoom.query.filter_by(form_teacher_id=user_id, is_active=True).all()
        return render_template("staff/psychomotor.html", current_user=current_user, teacher=teacher, terms=terms, form_master_classes=form_master_classes)

    @app.route("/staff/api/psychomotor/class/<class_id>", methods=["GET"])
    @staff_required
    def staff_class_psychomotor(class_id):
        term_id = request.args.get("term_id")
        if not term_id:
            return jsonify({"success": False, "message": "term_id is required"}), 400

        school = School.query.first()
        trait_defs = TraitDefinition.query.filter_by(
            school_id=school.school_id, is_active=True
        ).order_by(TraitDefinition.sort_order).all()

        users = User.query.filter_by(
            class_room_id=class_id, role="student", is_active=True
        ).order_by(User.first_name, User.last_name).all()

        student_ids = [u.id for u in users]

        totals = db.session.query(
            Grade.student_id,
            func.sum(Grade.score).label("total")
        ).filter(
            Grade.class_room_id == class_id,
            Grade.term_id == term_id,
            Grade.student_id.in_(student_ids),
            db.or_(Grade.is_published == True, Grade.is_from_cbt == True)
        ).group_by(Grade.student_id).all()

        total_map = {sid: t for sid, t in totals}
        all_totals = sorted(total_map.values(), reverse=True)

        def get_position(score):
            try:
                return all_totals.index(score) + 1
            except ValueError:
                return None

        result = []
        for u in users:
            student = u.student
            if not student:
                continue
            total_score = total_map.get(u.id, 0) or 0
            position = get_position(total_score)

            existing_traits = StudentTrait.query.filter_by(
                student_id=u.id, term_id=term_id
            ).all() if term_id else []
            score_map = {st.trait_id: st.score for st in existing_traits}

            result.append({
                "student_id": u.id,
                "name": f"{u.first_name} {u.last_name}".strip(),
                "reg_no": student.admission_number or "",
                "total_score": total_score,
                "position": position,
                "trait_scores": {t.id: score_map.get(t.id, None) for t in trait_defs},
            })

        return jsonify({
            "success": True,
            "students": result,
            "traits": [t.to_dict() for t in trait_defs],
        })

    # List all registered routes
    # print("=" * 80)
    # print("REGISTERED STAFF ROUTES:")
    for rule in app.url_map.iter_rules():
        if rule.endpoint and 'staff' in rule.rule:
            print(f"  {rule.rule} -> {rule.endpoint} {list(rule.methods)}")
