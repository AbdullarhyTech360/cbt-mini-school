"""
Utility functions for parsing Word documents containing question data.
"""
import json
import re
from docx import Document
from io import BytesIO

def parse_docx_questions(file_content):
    """
    Parse questions from a Word document.
    
    The Word document should be structured with:
    - Each question in a separate paragraph or table row
    - Question text followed by options (for MCQ/True-False)
    - Clear indication of correct answers
    - Metadata like subject_id, class_room_id, question_type
    
    This implementation assumes a simple structure where questions are 
    separated by double newlines and follow a specific format.
    
    Args:
        file_content (bytes): The content of the .docx file
        
    Returns:
        list: List of question dictionaries
    """
    # Load the document
    doc = Document(BytesIO(file_content))
    
    questions = []
    
    # Collect all paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # Group paragraphs into questions based on patterns
    current_question = {}
    in_options = False
    question_complete = False
    
    i = 0
    while i < len(paragraphs):
        paragraph = paragraphs[i]
        
        # Check if this is the start of a new question
        if not current_question and not paragraph.lower().startswith(('type:', 'subject:', 'class:', 'options:', 'correct answer:')):
            # This is likely a question text
            current_question = {
                'question_text': paragraph,
                'question_type': '',
                'subject_id': '',
                'class_room_id': '',
                'options': [],
                'correct_answer': ''
            }
            in_options = False
            question_complete = False
        elif current_question:
            # Process metadata
            if paragraph.lower().startswith('type:'):
                current_question['question_type'] = paragraph.split(':', 1)[1].strip().lower()
            elif paragraph.lower().startswith('subject:'):
                current_question['subject_id'] = paragraph.split(':', 1)[1].strip()
            elif paragraph.lower().startswith('class:'):
                current_question['class_room_id'] = paragraph.split(':', 1)[1].strip()
            elif paragraph.lower().startswith('correct answer:'):
                current_question['correct_answer'] = paragraph.split(':', 1)[1].strip()
            elif paragraph.lower() == 'options:':
                in_options = True
            elif in_options and paragraph.startswith('-'):
                # Parse option
                option_text = paragraph[1:].strip()
                is_correct = '[correct]' in option_text.lower()
                if is_correct:
                    option_text = re.sub(r'\[correct\]', '', option_text, flags=re.IGNORECASE).strip()
                
                option = {
                    'text': option_text,
                    'is_correct': is_correct
                }
                current_question['options'].append(option)
            elif not paragraph.lower().startswith(('type:', 'subject:', 'class:', 'options:', 'correct answer:')) and not paragraph.startswith('-'):
                # This might be the start of a new question
                # Add the current question to the list if it's valid
                if is_valid_question(current_question):
                    questions.append(current_question)
                
                # Start a new question
                current_question = {
                    'question_text': paragraph,
                    'question_type': '',
                    'subject_id': '',
                    'class_room_id': '',
                    'options': [],
                    'correct_answer': ''
                }
                in_options = False
                question_complete = False
        
        i += 1
    
    # Add the last question if it's valid
    if current_question and is_valid_question(current_question):
        questions.append(current_question)
    
    return questions

def is_valid_question(question):
    """
    Check if a question has all required fields.
    
    Args:
        question (dict): Question dictionary
        
    Returns:
        bool: True if question is valid, False otherwise
    """
    if not question.get('question_text') or not question.get('question_type'):
        return False
        
    # For MCQ and True/False, ensure we have options
    if question['question_type'] in ['mcq', 'true_false'] and not question.get('options'):
        return False
        
    # For short answer, ensure we have a correct answer
    if question['question_type'] == 'short_answer' and not question.get('correct_answer'):
        return False
        
    return True

def parse_question_block(block):
    """
    Parse a block of text as a question.
    
    Expected format:
    Question Text
    Type: mcq|true_false|short_answer
    Subject: SUBJECT_ID
    Class: CLASS_ROOM_ID
    Options:
    - Option 1 [correct]
    - Option 2
    Correct Answer: Answer text (for short answer)
    
    Args:
        block (str): Text block containing question data
        
    Returns:
        dict: Question dictionary or None if parsing fails
    """
    lines = block.strip().split('\n')
    if not lines:
        return None
        
    question = {
        'question_text': '',
        'question_type': '',
        'subject_id': '',
        'class_room_id': '',
        'options': [],
        'correct_answer': ''
    }
    
    # First line is the question text
    question['question_text'] = lines[0].strip()
    
    # Parse the rest of the lines
    options_started = False
    for line in lines[1:]:
        line = line.strip()
        if not line:
            continue
            
        if line.lower().startswith('type:'):
            question['question_type'] = line.split(':', 1)[1].strip().lower()
        elif line.lower().startswith('subject:'):
            question['subject_id'] = line.split(':', 1)[1].strip()
        elif line.lower().startswith('class:'):
            question['class_room_id'] = line.split(':', 1)[1].strip()
        elif line.lower().startswith('correct answer:'):
            question['correct_answer'] = line.split(':', 1)[1].strip()
        elif line.lower() == 'options:':
            options_started = True
        elif options_started and line.startswith('-'):
            # Parse option
            option_text = line[1:].strip()
            is_correct = '[correct]' in option_text.lower()
            if is_correct:
                option_text = re.sub(r'\[correct\]', '', option_text, flags=re.IGNORECASE).strip()
            
            option = {
                'text': option_text,
                'is_correct': is_correct
            }
            question['options'].append(option)
    
    # Validate required fields
    if not question['question_text'] or not question['question_type']:
        return None
        
    # For MCQ and True/False, ensure we have options
    if question['question_type'] in ['mcq', 'true_false'] and not question['options']:
        return None
        
    # For short answer, ensure we have a correct answer
    if question['question_type'] == 'short_answer' and not question['correct_answer']:
        return None
        
    return question

# Alternative approach for more structured documents (e.g., tables)
def parse_docx_questions_from_tables(file_content):
    """
    Parse questions from a Word document with questions in tables.
    
    Args:
        file_content (bytes): The content of the .docx file
        
    Returns:
        list: List of question dictionaries
    """
    doc = Document(BytesIO(file_content))
    
    questions = []
    
    # Look for tables in the document
    for table in doc.tables:
        # Assume first row is header
        if len(table.rows) < 2:
            continue
            
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        # Process each data row
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) != len(headers):
                continue
                
            # Create a dictionary from headers and cell values
            row_data = dict(zip(headers, cells))
            
            # Convert to question format
            question = {
                'question_text': row_data.get('Question', ''),
                'question_type': row_data.get('Type', '').lower(),
                'subject_id': row_data.get('Subject', ''),
                'class_room_id': row_data.get('Class', ''),
                'correct_answer': row_data.get('Correct Answer', ''),
                'options': []
            }
            
            # Parse options if they exist in a JSON format
            options_str = row_data.get('Options', '')
            if options_str:
                try:
                    question['options'] = json.loads(options_str)
                except json.JSONDecodeError:
                    # If not valid JSON, try to parse as a simple list
                    if ';' in options_str:
                        option_texts = options_str.split(';')
                        for i, opt_text in enumerate(option_texts):
                            is_correct = '[correct]' in opt_text.lower()
                            if is_correct:
                                opt_text = re.sub(r'\[correct\]', '', opt_text, flags=re.IGNORECASE).strip()
                            question['options'].append({
                                'text': opt_text.strip(),
                                'is_correct': is_correct
                            })
            
            # Validate and add question
            if (question['question_text'] and 
                question['question_type'] and 
                question['subject_id'] and 
                question['class_room_id']):
                
                # Additional validation based on question type
                if question['question_type'] in ['mcq', 'true_false']:
                    if question['options']:
                        questions.append(question)
                elif question['question_type'] == 'short_answer':
                    if question['correct_answer']:
                        questions.append(question)
    
    return questions